# -*- coding: utf-8 -*-
import os
import logging

# 환경 변수
os.environ['PYTHONIOENCODING'] = 'utf-8'

# 로깅 비활성화
logging.getLogger('anthropic').setLevel(logging.ERROR)
logging.getLogger('httpx').setLevel(logging.ERROR)

import streamlit as st
import streamlit.components.v1 as components
import re
import json
import html
import base64
import urllib.parse
import uuid
import platform

try:
    from PIL import Image, ImageDraw, ImageFont
    PILLOW_AVAILABLE = True
except Exception:
    PILLOW_AVAILABLE = False
import hashlib
import requests
from datetime import datetime, timedelta
from pathlib import Path

# Claude API
try:
    import anthropic
    CLAUDE_AVAILABLE = True
except ImportError:
    CLAUDE_AVAILABLE = False

# Gemini (이미지 생성용으로만 사용)
try:
    import google.generativeai as genai
    from google import genai as google_genai
    from google.genai import types as genai_types
    GEMINI_AVAILABLE = True
    IMAGEN_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False
    IMAGEN_AVAILABLE = False

# YouTube 자막 추출용
try:
    from youtube_transcript_api import YouTubeTranscriptApi
    YOUTUBE_TRANSCRIPT_AVAILABLE = True
except ImportError:
    YOUTUBE_TRANSCRIPT_AVAILABLE = False

# 브라우저 ID용 (클라우드 배포 시 필요)
try:
    from streamlit_javascript import st_javascript
    BROWSER_ID_AVAILABLE = True
except ImportError:
    BROWSER_ID_AVAILABLE = False

# 쿠키 매니저 (데이터 저장용)
try:
    import extra_streamlit_components as stx
    COOKIE_AVAILABLE = True
except ImportError:
    COOKIE_AVAILABLE = False

# Word 문서 생성용
try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_SECTION
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import io
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ==========================================
# 설정
# ==========================================
def get_config_path():
    return Path.home() / ".ebook_app_config.json"

def load_config():
    try:
        if get_config_path().exists():
            with open(get_config_path(), 'r') as f:
                return json.load(f)
    except:
        pass
    return {}

def save_config(data):
    try:
        config = load_config()
        config.update(data)
        with open(get_config_path(), 'w') as f:
            json.dump(config, f)
    except:
        pass

def load_saved_api_key():
    return load_config().get('api_key', '')

def save_api_key(api_key):
    save_config({'api_key': api_key})

def is_authenticated():
    return load_config().get('authenticated', False)

def save_authenticated():
    save_config({'authenticated': True})

# ==========================================
# 간단 비밀번호 인증
# ==========================================
CORRECT_PASSWORD = "cashmaker2024"  # ← 비밀번호 변경하려면 여기만 수정


# 비디오 배경용 base64 인코딩
@st.cache_data
def get_video_base64(video_path):
    try:
        with open(video_path, "rb") as f:
            return base64.b64encode(f.read()).decode()
    except:
        return None

st.set_page_config(page_title="Writey", layout="wide", page_icon="✍")

# 쿠키 매니저 초기화 및 데이터 불러오기/저장
# 승인 유효기간: 365일 + 롤링 만료 (방문할 때마다 자동 갱신)
COOKIE_LIFETIME_DAYS = 365  # 한 번 로그인 후 365일 유지. 방문할 때마다 자동 연장됨

if COOKIE_AVAILABLE:
    import time
    cookie_manager = stx.CookieManager(key="writey_cookies")
    cookies = cookie_manager.get_all()

    # [중요] CookieManager는 새 세션의 '첫 실행'에서는 브라우저와의 통신이 끝나기 전이라
    # 쿠키가 실제로 있어도 빈 dict({})를 반환한다. 그 상태로 자동 로그인을 판정하면
    # 항상 로그인 화면이 떠서 "로그인 정보가 기억되지 않는" 증상이 발생한다.
    # → 첫 실행에서 쿠키가 비어 있으면 딱 한 번만 재실행해서 쿠키를 읽을 기회를 준다.
    if not cookies and not st.session_state.get('_cookie_loaded'):
        st.session_state['_cookie_loaded'] = True
        time.sleep(0.3)
        st.rerun()
    st.session_state['_cookie_loaded'] = True

    # 쿠키에서 비밀번호/API키 복원
    if cookies:
        if 'writey_password' in cookies and cookies['writey_password']:
            if 'saved_password' not in st.session_state:
                st.session_state['saved_password'] = cookies['writey_password']
        if 'writey_api_key' in cookies and cookies['writey_api_key']:
            if 'saved_api_key' not in st.session_state:
                st.session_state['saved_api_key'] = cookies['writey_api_key']

    # 롤링 만료: 매 세션마다 만료일 연장
    if 'cookie_rolling_refreshed' not in st.session_state:
        if st.session_state.get('saved_password'):
            st.session_state['pending_save_password'] = st.session_state['saved_password']
        if st.session_state.get('saved_api_key'):
            st.session_state['pending_save_api'] = st.session_state['saved_api_key']
        st.session_state['cookie_rolling_refreshed'] = True

    # pending 값을 쿠키에 저장
    if 'pending_save_password' in st.session_state:
        cookie_manager.set('writey_password', st.session_state['pending_save_password'], expires_at=datetime.now() + timedelta(days=COOKIE_LIFETIME_DAYS))
        del st.session_state['pending_save_password']
    if 'pending_save_api' in st.session_state:
        cookie_manager.set('writey_api_key', st.session_state['pending_save_api'], expires_at=datetime.now() + timedelta(days=COOKIE_LIFETIME_DAYS))
        del st.session_state['pending_save_api']
else:
    cookie_manager = None

def save_password_to_browser(password):
    """비밀번호를 쿠키에 저장"""
    st.session_state['saved_password'] = password
    st.session_state['pending_save_password'] = password

def save_api_key_to_browser(api_key):
    """API 키를 쿠키에 저장"""
    st.session_state['saved_api_key'] = api_key
    st.session_state['pending_save_api'] = api_key

def get_saved_api_key():
    """저장된 API 키 반환"""
    return st.session_state.get('saved_api_key', None)

# ==========================================
# APPLE STYLE CSS
# ==========================================
st.markdown("""
<style>
@import url('https://cdn.jsdelivr.net/gh/orioncactus/pretendard/dist/web/static/pretendard.css');
@import url('https://fonts.googleapis.com/css2?family=Noto+Serif+KR:wght@400;500;600;700;900&display=swap');
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@400;500;600;700;800;900&display=swap');
@import url('https://fonts.googleapis.com/css2?family=Cormorant+Garamond:wght@400;500;600;700&display=swap');
@import url('https://fonts.googleapis.com/css2?family=Cinzel:wght@400;500;600;700&display=swap');

/* S-Core Dream 폰트 */
@font-face {
    font-family: 'S-CoreDream';
    src: url('https://cdn.jsdelivr.net/gh/projectnoonnu/noonfonts_six@1.2/S-CoreDream-1Thin.woff') format('woff');
    font-weight: 100;
}
@font-face {
    font-family: 'S-CoreDream';
    src: url('https://cdn.jsdelivr.net/gh/projectnoonnu/noonfonts_six@1.2/S-CoreDream-2ExtraLight.woff') format('woff');
    font-weight: 200;
}
@font-face {
    font-family: 'S-CoreDream';
    src: url('https://cdn.jsdelivr.net/gh/projectnoonnu/noonfonts_six@1.2/S-CoreDream-3Light.woff') format('woff');
    font-weight: 300;
}
@font-face {
    font-family: 'S-CoreDream';
    src: url('https://cdn.jsdelivr.net/gh/projectnoonnu/noonfonts_six@1.2/S-CoreDream-4Regular.woff') format('woff');
    font-weight: 400;
}
@font-face {
    font-family: 'S-CoreDream';
    src: url('https://cdn.jsdelivr.net/gh/projectnoonnu/noonfonts_six@1.2/S-CoreDream-5Medium.woff') format('woff');
    font-weight: 500;
}
@font-face {
    font-family: 'S-CoreDream';
    src: url('https://cdn.jsdelivr.net/gh/projectnoonnu/noonfonts_six@1.2/S-CoreDream-6Bold.woff') format('woff');
    font-weight: 600;
}
@font-face {
    font-family: 'S-CoreDream';
    src: url('https://cdn.jsdelivr.net/gh/projectnoonnu/noonfonts_six@1.2/S-CoreDream-7ExtraBold.woff') format('woff');
    font-weight: 700;
}
@font-face {
    font-family: 'S-CoreDream';
    src: url('https://cdn.jsdelivr.net/gh/projectnoonnu/noonfonts_six@1.2/S-CoreDream-8Heavy.woff') format('woff');
    font-weight: 800;
}
@font-face {
    font-family: 'S-CoreDream';
    src: url('https://cdn.jsdelivr.net/gh/projectnoonnu/noonfonts_six@1.2/S-CoreDream-9Black.woff') format('woff');
    font-weight: 900;
}

:root {
    --gold: #C9A24B;
    --gold-light: #E0C074;
    --gold-dark: #A8852F;
    --rose-gold: #b76e79;
    --cream: #FAF8F4;
    --charcoal: #141416;
    --dark: #0B0B0D;
    --card: rgba(255,255,255,0.025);
    --card2: rgba(255,255,255,0.05);
    --text: #F5F3EF;
    --text2: #8A8780;
    --text3: #7A776F;
    --line: rgba(201,162,75,0.18);
    --line2: rgba(255,255,255,0.06);
    --glow: rgba(201,162,75,0.32);
    --success: #6FA86F;
    --warning: #E0C074;
    --danger: #C97A6F;
}

/* 애니메이션 정의 */
@keyframes fadeInUp {
    from { opacity: 0; transform: translateY(30px); }
    to { opacity: 1; transform: translateY(0); }
}
@keyframes shimmer {
    0% { background-position: -200% 0; }
    100% { background-position: 200% 0; }
}
@keyframes pulse {
    0%, 100% { opacity: 1; transform: scale(1); }
    50% { opacity: 0.8; transform: scale(1.02); }
}
@keyframes borderGlow {
    0%, 100% { box-shadow: 0 0 5px var(--glow), inset 0 0 5px rgba(201,162,75,0.1); }
    50% { box-shadow: 0 0 20px var(--glow), inset 0 0 10px rgba(201,162,75,0.2); }
}
@keyframes float {
    0%, 100% { transform: translateY(0); }
    50% { transform: translateY(-5px); }
}
@keyframes goldShine {
    0% { background-position: -100% 0; }
    100% { background-position: 200% 0; }
}

*:not([data-testid*="Icon"]):not(.material-icons):not([class*="icon"]):not(span[aria-hidden="true"]) {
    font-family: 'S-CoreDream', 'Pretendard', -apple-system, sans-serif !important;
}
/* 아이콘 폰트 복원 */
[data-testid*="Icon"], .material-icons, span[aria-hidden="true"], button[kind="header"] span {
    font-family: 'Material Symbols Rounded', 'Material Icons', sans-serif !important;
}
.stDeployButton, footer, #MainMenu { display: none !important; }
/* 헤더 투명하게 (사이드바 버튼은 보임) */
header[data-testid="stHeader"] {
    background: transparent !important;
}

/* 럭셔리 배경 - 미세한 그라데이션 */
.stApp {
    background:
        radial-gradient(ellipse at 20% 0%, rgba(201,162,75,0.04) 0%, transparent 55%),
        radial-gradient(ellipse at 80% 100%, rgba(201,162,75,0.025) 0%, transparent 55%),
        linear-gradient(180deg, #0B0B0D 0%, #08080A 50%, #0B0B0D 100%) !important;
    background-attachment: fixed;
}

.main .block-container { max-width: 1000px; padding: 3rem 2rem; }

/* 사이드바 - 미니멀 */
[data-testid="stSidebar"] {
    background: var(--charcoal) !important;
    border-right: 1px solid var(--line);
}
[data-testid="stSidebar"] * { color: var(--text2) !important; }

/* 타이포그래피 - 가독성 향상 */
h1, h2, h3 { color: var(--text) !important; font-weight: 300 !important; letter-spacing: 0.5px; }
h1 { font-size: 34px !important; color: var(--cream) !important; font-weight: 300 !important; }
h2 { font-size: 26px !important; margin-bottom: 20px !important; font-weight: 300 !important; }
h3 { font-size: 21px !important; color: var(--gold) !important; font-weight: 400 !important; }
p, span, label, div { color: var(--text) !important; font-size: 16px !important; line-height: 1.7 !important; }
li { font-size: 16px !important; line-height: 1.8 !important; }

/* 버튼 - 채워진 골드 그라데이션 (첨부 디자인) */
.stButton > button {
    background: linear-gradient(135deg, #E0C074 0%, #C9A24B 100%) !important;
    color: #0B0B0D !important;
    -webkit-text-fill-color: #0B0B0D !important;
    border: none !important;
    border-radius: 12px;
    font-weight: 600;
    font-size: 15px !important;
    padding: 15px 36px;
    letter-spacing: 0.4px;
    text-transform: none;
    transition: all 0.35s cubic-bezier(0.4, 0, 0.2, 1);
    position: relative;
    overflow: hidden;
    box-shadow: 0 6px 20px rgba(201,162,75,0.22);
}
.stButton > button * {
    color: #0B0B0D !important;
    -webkit-text-fill-color: #0B0B0D !important;
}
.stButton > button::before {
    content: '';
    position: absolute;
    top: 0;
    left: -100%;
    width: 100%;
    height: 100%;
    background: linear-gradient(90deg, transparent, rgba(255,255,255,0.45), transparent);
    transition: left 0.6s ease;
}
.stButton > button:hover::before {
    left: 100%;
}
.stButton > button:hover {
    background: linear-gradient(135deg, #EBCE86 0%, #D4AC56 100%) !important;
    box-shadow: 0 10px 32px rgba(201,162,75,0.4);
    transform: translateY(-2px);
}
.stButton > button:active {
    transform: translateY(0);
    box-shadow: 0 4px 15px rgba(201,162,75,0.3);
}

/* 입력 필드 - 밝은 배경 + 검은 글씨 */
.stTextInput input, .stTextArea textarea, .stNumberInput input {
    background: #ffffff !important;
    background-color: #ffffff !important;
    border: 0.5px solid var(--line) !important;
    border-radius: 10px !important;
    color: #000000 !important;
    -webkit-text-fill-color: #000000 !important;
    caret-color: #000000 !important;
    padding: 18px !important;
    font-size: 17px !important;
}
.stTextInput input:focus, .stTextArea textarea:focus, .stNumberInput input:focus {
    border-color: var(--gold) !important;
    box-shadow: 0 0 0 2px rgba(201,162,75,0.2) !important;
}

/* 셀렉트박스 컨테이너 */
.stSelectbox > div > div {
    background: var(--card) !important;
    border: 0.5px solid var(--line) !important;
    border-radius: 10px;
}
/* 셀렉트박스 선택된 값 - 흰색 */
.stSelectbox [data-baseweb="select"] > div {
    color: #ffffff !important;
    -webkit-text-fill-color: #ffffff !important;
}

/* 스코어 카드 - 럭셔리 */
.score-card {
    background: linear-gradient(145deg, var(--card) 0%, rgba(30,30,30,0.95) 100%) !important;
    border: 0.5px solid var(--gold);
    border-radius: 20px;
    padding: 50px 40px;
    text-align: center;
    animation: fadeInUp 0.6s ease-out;
    transition: all 0.5s cubic-bezier(0.4, 0, 0.2, 1);
    position: relative;
    overflow: hidden;
    box-shadow: 0 10px 40px rgba(201,162,75,0.15);
}
.score-card::before {
    content: '';
    position: absolute;
    top: 0;
    left: 0;
    right: 0;
    height: 3px;
    background: linear-gradient(90deg, transparent, var(--gold), transparent);
    opacity: 1;
}
.score-card:hover {
    border-color: var(--gold);
    box-shadow: 0 20px 60px rgba(201,162,75,0.3), inset 0 1px 0 rgba(201,162,75,0.1);
    transform: translateY(-5px);
}
.score-card:hover::before {
    opacity: 1;
}
.score-number {
    font-size: 140px;
    font-weight: 300;
    background: linear-gradient(135deg, var(--gold-light) 0%, var(--gold) 50%, var(--gold-dark) 100%);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    background-clip: text;
    line-height: 1;
    letter-spacing: -4px;
    animation: fadeInUp 0.8s ease-out;
    filter: drop-shadow(0 2px 4px rgba(201,162,75,0.3));
}

/* 정보 카드 + 애니메이션 */
.info-card {
    background: transparent !important;
    border: none;
    border-left: 2px solid var(--gold);
    padding: 20px 24px;
    margin: 20px 0;
    animation: fadeInUp 0.5s ease-out;
    transition: all 0.3s ease;
}
.info-card:hover {
    background: rgba(201,169,98,0.05) !important;
    border-left-width: 4px;
    padding-left: 22px;
}

/* 스탯 박스 + 애니메이션 */
.stat-box {
    background: var(--card) !important;
    border: 0.5px solid var(--line2);
    border-radius: 14px;
    padding: 32px;
    text-align: center;
    animation: fadeInUp 0.5s ease-out;
    transition: all 0.4s ease;
}
.stat-box:hover {
    transform: translateY(-4px);
    box-shadow: 0 8px 25px rgba(0,0,0,0.3);
    border-color: var(--gold);
}
.stat-value {
    font-size: 42px;
    font-weight: 200;
    color: var(--gold) !important;
    letter-spacing: -2px;
    transition: transform 0.3s ease;
}
.stat-box:hover .stat-value {
    transform: scale(1.05);
}
.stat-label {
    font-size: 11px;
    color: var(--text2) !important;
    margin-top: 12px;
    text-transform: uppercase;
    letter-spacing: 3px;
}

/* 데이터 카드 + 애니메이션 */
.data-card {
    background: var(--card) !important;
    border-left: 2px solid var(--gold);
    padding: 20px 24px;
    margin: 16px 0;
    animation: fadeInUp 0.4s ease-out;
    transition: all 0.3s ease;
}
.data-card:hover {
    border-left-width: 4px;
    background: var(--card2) !important;
}

/* 서머리 허브 + 애니메이션 */
.summary-hub {
    background: var(--card) !important;
    border: 0.5px solid var(--line2);
    border-radius: 14px;
    padding: 40px;
    animation: fadeInUp 0.5s ease-out;
    transition: all 0.4s ease;
}
.summary-hub:hover {
    border-color: var(--gold);
}

/* 배지 - 미니멀 + 펄스 */
.verdict-go {
    background: transparent !important;
    color: var(--success) !important;
    border: 1px solid var(--success);
    padding: 12px 32px;
    border-radius: 20px;
    font-weight: 400;
    font-size: 12px;
    letter-spacing: 3px;
    text-transform: uppercase;
    animation: fadeInUp 0.6s ease-out;
}
.verdict-wait {
    background: transparent !important;
    color: var(--warning) !important;
    border: 1px solid var(--warning);
    padding: 12px 32px;
    border-radius: 20px;
    font-weight: 400;
    font-size: 12px;
    letter-spacing: 3px;
    animation: fadeInUp 0.6s ease-out;
}
.verdict-no {
    background: transparent !important;
    color: var(--danger) !important;
    border: 1px solid var(--danger);
    padding: 12px 32px;
    border-radius: 20px;
    font-weight: 400;
    font-size: 12px;
    letter-spacing: 3px;
    animation: fadeInUp 0.6s ease-out;
}

/* 네비게이션 */
.premium-nav-container {
    background: transparent;
    border-top: 1px solid var(--line);
    border-bottom: 1px solid var(--line);
    padding: 0;
    margin-bottom: 48px;
}
.nav-item {
    padding: 18px 12px;
    text-align: center;
    font-size: 14px;
    color: var(--text2);
    letter-spacing: 1px;
    transition: all 0.3s ease;
}
.nav-item.active {
    background: linear-gradient(135deg, rgba(201,162,75,0.2) 0%, rgba(201,162,75,0.1) 100%);
    color: var(--gold) !important;
    font-weight: 600;
    border-bottom: 3px solid var(--gold);
    box-shadow: 0 4px 15px rgba(201,162,75,0.2);
}

/* 섹션 타이틀 - 미니멀 (첨부 디자인) */
.section-title-box {
    background: rgba(255,255,255,0.025);
    border: 0.5px solid var(--line2);
    border-radius: 14px;
    padding: 30px 36px;
    margin-bottom: 35px;
    text-align: left;
    position: relative;
    overflow: hidden;
    animation: fadeInUp 0.5s ease-out;
}
.section-title-box h2 {
    font-size: 26px !important;
    color: #FAF8F4 !important;
    margin: 0 0 8px 0 !important;
    font-weight: 300 !important;
    letter-spacing: 0.5px;
}
.section-title-box p {
    color: var(--text2) !important;
    font-size: 14px !important;
    margin: 0 !important;
}
.section-step {
    display: inline-block;
    background: transparent;
    color: var(--gold) !important;
    font-size: 11px;
    font-weight: 500;
    padding: 0;
    margin-bottom: 12px;
    letter-spacing: 0.22em;
}

/* 제목 카드 + 애니메이션 */
.title-card {
    background: var(--card);
    border: 0.5px solid var(--line2);
    border-radius: 12px;
    padding: 28px;
    margin: 16px 0;
    transition: all 0.4s cubic-bezier(0.4, 0, 0.2, 1);
    animation: fadeInUp 0.5s ease-out;
}
.title-card:hover {
    border-color: var(--gold);
    background: rgba(201,169,98,0.05);
    transform: translateX(8px);
    box-shadow: -4px 0 20px rgba(201,169,98,0.15);
}
.title-main {
    font-size: 18px;
    font-weight: 400;
    color: var(--text) !important;
    letter-spacing: 1px;
    transition: color 0.3s ease;
}
.title-card:hover .title-main {
    color: var(--gold) !important;
}
.title-sub {
    font-size: 13px;
    color: var(--text2) !important;
    margin-top: 8px;
}

/* 로그인 - 럭셔리 */
.login-card {
    max-width: 420px;
    margin: 100px auto;
    padding: 70px 50px;
    background: linear-gradient(145deg, rgba(26,26,31,0.98) 0%, rgba(11,11,13,0.98) 100%);
    border: 0.5px solid var(--line);
    border-radius: 16px;
    text-align: center;
    animation: fadeInUp 0.8s ease-out;
    position: relative;
    box-shadow: 0 25px 80px rgba(0,0,0,0.5), 0 0 40px rgba(201,162,75,0.05);
}
.login-card::before {
    content: '';
    position: absolute;
    top: -1px;
    left: 20%;
    right: 20%;
    height: 2px;
    background: linear-gradient(90deg, transparent, var(--gold), transparent);
}
.login-card::after {
    content: '';
    position: absolute;
    bottom: -1px;
    left: 20%;
    right: 20%;
    height: 1px;
    background: linear-gradient(90deg, transparent, var(--gold-dark), transparent);
}
.login-title {
    font-size: 32px;
    font-weight: 300;
    background: linear-gradient(135deg, var(--cream) 0%, var(--gold-light) 50%, var(--gold) 100%);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    background-clip: text;
    letter-spacing: 7px;
    animation: fadeInUp 1s ease-out;
}
.login-subtitle {
    font-size: 11px;
    color: var(--gold) !important;
    margin-top: 20px;
    letter-spacing: 4px;
    text-transform: uppercase;
    animation: fadeInUp 1.2s ease-out;
    opacity: 0.8;
}

/* 헤더 - 럭셔리 */
.main-header {
    text-align: center;
    padding: 80px 20px 60px;
    margin-bottom: 50px;
    border-bottom: 1px solid var(--line);
    animation: fadeInUp 0.6s ease-out;
    position: relative;
    background: linear-gradient(180deg, rgba(201,162,75,0.02) 0%, transparent 100%);
}
.main-header::after {
    content: '';
    position: absolute;
    bottom: -1px;
    left: 10%;
    right: 10%;
    height: 1px;
    background: linear-gradient(90deg, transparent, var(--gold), transparent);
}
.main-header-brand {
    font-size: 11px;
    color: var(--gold) !important;
    letter-spacing: 10px;
    text-transform: uppercase;
    animation: fadeInUp 0.8s ease-out;
    text-shadow: 0 0 20px rgba(201,162,75,0.3);
}
.main-header-title {
    font-size: 42px;
    font-weight: 200;
    background: linear-gradient(135deg, var(--cream) 0%, var(--gold-light) 100%);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    background-clip: text;
    letter-spacing: 8px;
    margin-top: 24px;
    animation: fadeInUp 1s ease-out;
}
.header-tagline {
    font-size: 13px;
    color: var(--text2) !important;
    margin-top: 24px;
    letter-spacing: 3px;
    animation: fadeInUp 1.2s ease-out;
}

/* Expander + 애니메이션 */
.stExpander {
    background: var(--card) !important;
    border: 0.5px solid var(--line2) !important;
    border-radius: 12px !important;
    animation: fadeInUp 0.4s ease-out;
    transition: border-color 0.3s ease;
}
.stExpander:hover {
    border-color: var(--gold) !important;
}
.stProgress > div > div > div {
    background: linear-gradient(90deg, var(--gold-dark), var(--gold), var(--gold-light), var(--gold), var(--gold-dark)) !important;
    background-size: 300% 100%;
    animation: goldShine 3s ease infinite;
    border-radius: 4px;
    box-shadow: 0 0 15px rgba(201,162,75,0.4);
}
.stProgress > div > div {
    background: rgba(20,20,20,0.8);
    border-radius: 4px;
    border: 1px solid var(--line);
}

/* 라디오 & 탭 */
.stRadio > div { background: transparent; border: 1px solid var(--line); padding: 16px; }
.stTabs [data-baseweb="tab-list"] { background: transparent; border-bottom: 1px solid var(--line); }
.stTabs [aria-selected="true"] {
    background: transparent !important;
    color: var(--gold) !important;
    border-bottom: 2px solid var(--gold) !important;
}

/* 알림 */
.stSuccess > div { background: rgba(111,168,111,0.1) !important; border: 0.5px solid rgba(111,168,111,0.3) !important; border-radius: 10px; }
.stWarning > div { background: rgba(224,192,116,0.1) !important; border: 0.5px solid rgba(224,192,116,0.3) !important; border-radius: 10px; }
.stError > div { background: rgba(201,122,111,0.1) !important; border: 0.5px solid rgba(201,122,111,0.3) !important; border-radius: 10px; }
.stInfo > div { background: rgba(201,162,75,0.08) !important; border: 0.5px solid var(--line) !important; border-radius: 10px; }

/* 스크롤바 */
::-webkit-scrollbar { width: 6px; }
::-webkit-scrollbar-track { background: var(--dark); }
::-webkit-scrollbar-thumb { background: var(--gold-dark); }

/* 다운로드 버튼 - 럭셔리 골드 */
.stDownloadButton button {
    background: linear-gradient(135deg, var(--gold) 0%, var(--gold-dark) 50%, var(--gold) 100%) !important;
    background-size: 200% 100%;
    color: var(--dark) !important;
    border: none !important;
    border-radius: 12px;
    font-weight: 600;
    letter-spacing: 1.5px;
    box-shadow: 0 4px 20px rgba(201,162,75,0.3);
    transition: all 0.4s ease;
    text-shadow: 0 1px 1px rgba(255,255,255,0.2);
}
.stDownloadButton button:hover {
    background-position: 100% 0 !important;
    box-shadow: 0 8px 35px rgba(201,162,75,0.5);
    transform: translateY(-2px);
}

/* 구분선 */
hr { border: none; height: 1px; background: var(--line); margin: 40px 0; }

/* 표지 미리보기 - 실제 책처럼 */
.book-wrapper {
    perspective: 1000px;
    display: flex;
    justify-content: center;
    padding: 30px;
    background: linear-gradient(135deg, #1a1a1a 0%, #2d2d2d 100%);
    border-radius: 8px;
}
.ebook-cover {
    font-family: 'Pretendard', sans-serif !important;
    box-shadow:
        0 0 5px rgba(0,0,0,0.3),
        5px 5px 15px rgba(0,0,0,0.4),
        10px 10px 30px rgba(0,0,0,0.3),
        15px 15px 50px rgba(0,0,0,0.2),
        inset -3px 0 10px rgba(0,0,0,0.2);
    transform: rotateY(-3deg);
    border-radius: 0 3px 3px 0;
    position: relative;
}
.ebook-cover::before {
    content: '';
    position: absolute;
    left: 0;
    top: 0;
    bottom: 0;
    width: 25px;
    background: linear-gradient(90deg,
        rgba(0,0,0,0.4) 0%,
        rgba(0,0,0,0.1) 30%,
        rgba(255,255,255,0.05) 50%,
        rgba(0,0,0,0.1) 70%,
        rgba(0,0,0,0.3) 100%);
    border-radius: 3px 0 0 3px;
}
.ebook-cover::after {
    content: '';
    position: absolute;
    right: 0;
    top: 2px;
    bottom: 2px;
    width: 8px;
    background: linear-gradient(90deg,
        rgba(255,255,255,0.03) 0%,
        rgba(255,255,255,0.08) 50%,
        rgba(0,0,0,0.1) 100%);
}
.ebook-cover * {
    color: inherit !important;
    -webkit-text-fill-color: inherit !important;
}

/* ============================================
   입력 필드 텍스트 색상 - 최우선 적용
   ============================================ */

/* 모든 입력 필드 - 흰 배경 + 검은 글씨 */
.stTextInput input,
.stTextArea textarea,
.stNumberInput input,
[data-testid="stTextInput"] input,
[data-testid="stTextArea"] textarea,
[data-testid="stNumberInput"] input {
    background: #ffffff !important;
    background-color: #ffffff !important;
    color: #000000 !important;
    -webkit-text-fill-color: #000000 !important;
    caret-color: #000000 !important;
}

/* Placeholder 색상 */
input::placeholder,
textarea::placeholder {
    color: #888888 !important;
    -webkit-text-fill-color: #888888 !important;
}

/* 셀렉트박스 - 선택된 값 (어두운 배경에 흰 글씨) */
.stSelectbox [data-baseweb="select"] > div,
.stSelectbox [data-baseweb="select"] span,
.stSelectbox > div > div > div {
    color: #ffffff !important;
    -webkit-text-fill-color: #ffffff !important;
}

/* ============================================
   드롭다운/팝오버 - 검은 글씨 (흰 배경)
   ============================================ */
[data-baseweb="popover"],
[data-baseweb="popover"] *,
[data-baseweb="menu"],
[data-baseweb="menu"] *,
[data-baseweb="list"],
[data-baseweb="list"] *,
[role="listbox"],
[role="listbox"] *,
[role="option"],
[role="option"] *,
.stSelectbox ul,
.stSelectbox ul *,
.stSelectbox li,
.stSelectbox li * {
    background: #ffffff !important;
    background-color: #ffffff !important;
    color: #000000 !important;
    -webkit-text-fill-color: #000000 !important;
}

/* 드롭다운 옵션 호버 */
[role="option"]:hover,
[data-baseweb="menu"] li:hover,
.stSelectbox li:hover {
    background: #f0f0f0 !important;
    background-color: #f0f0f0 !important;
}

/* select 요소 */
select,
select option {
    color: #000000 !important;
    background: #ffffff !important;
}

/* Expander 스타일 정리 */
.stExpander details summary {
    background: var(--card) !important;
    overflow: hidden !important;
}
/* 모든 텍스트 숨기기 (keyboard_arrow 등 영어 텍스트 포함) */
.stExpander details summary * {
    font-size: 0 !important;
    color: transparent !important;
    -webkit-text-fill-color: transparent !important;
}
/* 한국어 제목만 보이게 */
.stExpander details summary p {
    font-size: 15px !important;
    color: var(--text) !important;
    -webkit-text-fill-color: var(--text) !important;
}
/* 화살표 아이콘만 보이게 */
.stExpander details summary svg {
    width: 20px !important;
    height: 20px !important;
    color: var(--gold) !important;
    fill: var(--gold) !important;
}

/* 버튼 앞 불필요한 라벨 숨기기 */
.stButton > div:not([data-testid="baseButton-secondary"]):not([data-testid="baseButton-primary"]) > p,
.stButton > div > div > p:first-child:not(:last-child),
.stButton label,
.stExpander .stButton > div:first-child > p {
    display: none !important;
}
/* 링크버튼 라벨 숨기기 */
.stLinkButton > div:first-child > p {
    display: none !important;
}
</style>
""", unsafe_allow_html=True)



# ==========================================
# 비밀번호 인증 (단순)
# ==========================================
if 'authenticated' not in st.session_state:
    st.session_state['authenticated'] = False

# 저장된 비밀번호로 자동 로그인 (매 실행마다 재확인)
# 쿠키는 첫 실행 뒤 재실행 시점에 늦게 로드될 수 있으므로, 미인증 상태에서는
# 매번 saved_password를 다시 확인해야 쿠키가 늦게 들어와도 자동 로그인이 된다.
if not st.session_state['authenticated']:
    if st.session_state.get('saved_password', '') == CORRECT_PASSWORD:
        st.session_state['authenticated'] = True

if not st.session_state['authenticated']:
    st.markdown("""
    <div class="login-card">
        <div class="login-subtitle">CASHMAKER</div>
        <div class="login-title">WRITEY</div>
        <div class="login-subtitle">Premium E-Book Studio</div>
    </div>
    """, unsafe_allow_html=True)

    col1, col2, col3 = st.columns([1, 1, 1])
    with col2:
        pw = st.text_input("비밀번호", type="password", key="pw_login", placeholder="Enter password...")
        remember_pw = st.checkbox("비밀번호 저장 (다음 접속 시 자동 로그인)", value=True, key="remember_pw")
        if st.button("입장", key="btn_login", use_container_width=True):
            if pw == CORRECT_PASSWORD:
                st.session_state['authenticated'] = True
                if remember_pw:
                    save_password_to_browser(pw)
                st.rerun()
            else:
                st.error("비밀번호가 올바르지 않습니다")
    st.stop()


# 세션 초기화
defaults = {
    'topic': '', 'target_persona': '', 'pain_points': '',
    'outline': [], 'chapters': {}, 'book_title': '', 'subtitle': '',
    'score_details': None, 'generated_titles': None, 'suggested_targets': None,
    'analyzed_pains': None, 'review_analysis': None, 'market_gaps': None,
    'knowledge_hub': [], 'study_summary': None, 'current_page': 0,
    'recommended_refs': None, 'generated_ideas': None,
    # 인터뷰 관련 변수
    'interview_completed': False,
    'interview_data': {},
    'author_name': '',
}
for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

# 사이드바
with st.sidebar:
    # API 키 섹션 (접기/펼치기 가능)
    if 'show_api_section' not in st.session_state:
        st.session_state['show_api_section'] = True

    if 'api_key' not in st.session_state:
        # 쿠키에서 API 키 불러오기
        st.session_state['api_key'] = st.session_state.get('saved_api_key', '') or ''
    # api_key가 비어 있는데 쿠키에 저장된 키가 (늦게) 들어왔다면 복원
    elif not st.session_state['api_key'] and st.session_state.get('saved_api_key'):
        st.session_state['api_key'] = st.session_state['saved_api_key']

    # API 키가 입력되어 있으면 기본적으로 접힌 상태로
    api_key_exists = bool(st.session_state['api_key'])

    col_title, col_toggle = st.columns([4, 1])
    with col_title:
        st.markdown("### 🔑 Claude API 키")
    with col_toggle:
        toggle_label = "▼" if st.session_state['show_api_section'] else "▶"
        if st.button(toggle_label, key="toggle_api_section", help="접기/펼치기"):
            st.session_state['show_api_section'] = not st.session_state['show_api_section']
            st.rerun()

    if st.session_state['show_api_section']:
        api_key = st.text_input("키 입력", value=st.session_state['api_key'], type="password", key="api_sidebar", label_visibility="collapsed", placeholder="sk-ant-api03-... 형식")
        if api_key != st.session_state['api_key']:
            st.session_state['api_key'] = api_key
            # 쿠키에 저장
            if api_key:
                save_api_key_to_browser(api_key)
                # 비밀번호 흐름과 동일하게 즉시 재실행해 pending 값을 쿠키에 바로 기록
                # (재실행이 없으면 다음 상호작용 전까지 쿠키 기록이 미뤄져 저장이 누락될 수 있음)
                st.rerun()

        if api_key:
            st.success("✅ Claude 키 입력 완료!")
        else:
            st.error("⚠️ Claude API 키를 입력하세요")
    else:
        # 접힌 상태에서 간단한 상태 표시
        if st.session_state['api_key']:
            st.caption("✅ API 키 설정됨")
        else:
            st.caption("⚠️ API 키 필요")

    # 모델 선택
    st.markdown("### 🤖 모델 선택")
    if 'claude_model' not in st.session_state:
        st.session_state['claude_model'] = "claude-sonnet-4-5"

    model_options = {
        "Claude Sonnet 4.5 (추천)": "claude-sonnet-4-5",
        "Claude Opus 4.5 (최고 품질)": "claude-opus-4-5",
        "Claude Haiku 4.5 (저렴)": "claude-haiku-4-5"
    }
    selected_model = st.selectbox(
        "모델 선택",
        options=list(model_options.keys()),
        index=0,
        label_visibility="collapsed"
    )
    st.session_state['claude_model'] = model_options[selected_model]

    if "Haiku" in selected_model:
        st.info("💰 가장 저렴하고 빠름.\n📌 단, 목차·본문·프롤로그·에필로그·컨셉·제목 생성은 품질 보장을 위해 자동으로 Sonnet 4.5 사용")
    elif "Opus" in selected_model:
        st.info("💎 최고 품질, 단가가 가장 높음")
    else:
        st.info("⚡ 균형잡힌 품질/가격, 일반적으로 가장 추천")

    # API 키 발급 방법 안내
    with st.expander("📖 Claude API 키 발급 방법 (상세)", expanded=False):
        st.markdown("""
        ### 🟣 1단계: Anthropic 회원가입

        1. 아래 버튼을 클릭하세요
        2. **"Sign up"** 클릭
        3. Google 계정 또는 이메일로 가입
        """)
        st.link_button("🔗 Anthropic 가입 페이지", "https://console.anthropic.com/", use_container_width=True)

        st.markdown("""
        ---
        ### 💳 2단계: 결제 수단 등록

        1. 로그인 후 왼쪽 메뉴에서 **"Settings"** 클릭
        2. **"Billing"** 클릭
        3. **"Add payment method"** 클릭
        4. 카드 정보 입력 후 저장
        5. **"Add credits"**로 크레딧 충전 ($5~10 추천)
        """)
        st.link_button("🔗 Billing 페이지 바로가기", "https://console.anthropic.com/settings/billing", use_container_width=True)

        st.markdown("""
        ---
        ### 🔑 3단계: API 키 발급

        1. 왼쪽 메뉴에서 **"API Keys"** 클릭
        2. **"Create Key"** 버튼 클릭
        3. 이름 입력 (예: ebook)
        4. **"Create Key"** 클릭
        5. 생성된 키 **복사** (sk-ant-api03-... 형식)
        6. 위 입력창에 **붙여넣기**
        """)
        st.link_button("🔗 API Keys 페이지 바로가기", "https://console.anthropic.com/settings/keys", use_container_width=True)

        st.markdown("---")
        st.warning("⚠️ API 키는 한 번만 보여줍니다. 복사해두세요!")
        st.success("💰 예상 비용: 전자책 1권당 약 200~500원")

    st.markdown("---")
    st.markdown("### 📊 진행 상황")
    progress = sum([bool(st.session_state['topic']), bool(st.session_state['target_persona']), bool(st.session_state['outline']), len(st.session_state['chapters']) > 0]) / 4
    st.progress(progress)

    st.markdown("---")
    st.markdown("### 🚀 빠른 이동")
    sidebar_pages = ["① 주제", "② 목차", "③ 본문", "④ 완성"]
    sidebar_mapping = [0, 4, 5, 7]
    for i, p in enumerate(sidebar_pages):
        if st.button(p, key=f"sidebar_nav_{i}", use_container_width=True):
            st.session_state['current_page'] = sidebar_mapping[i]
            st.rerun()

    # 사이드바 하단 제작자 정보
    st.markdown("---")
    st.markdown("""
    <div style="text-align:center; padding:10px 0; color:#C9A24B !important; font-size:12px;">
        <strong>CASHMAKER</strong><br>
        <span style="color:#ffffff !important;">제작: 남현우 작가</span>
    </div>
    """, unsafe_allow_html=True)

# ==========================================
# 헬퍼 함수
# ==========================================
def get_api_key():
    return st.session_state.get('api_key', '')

def clean_text(text):
    if not text:
        return ""
    text = re.sub(r'^#{1,6}\s*', '', text, flags=re.MULTILINE)
    text = re.sub(r'\*\*([^*]+)\*\*', r'「\1」', text)
    text = text.replace('**', '').replace('*', '').replace('###', '').replace('##', '').replace('#', '')
    return text.strip()

def clean_content(text, subtopic=None):
    if not text:
        return ""
    # HTML 테이블 및 모든 HTML 태그 제거
    text = re.sub(r'<table[^>]*>.*?</table>', '', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<[^>]+>', '', text)
    # 마크다운 제거
    text = re.sub(r'^#{1,6}\s*', '', text, flags=re.MULTILINE)
    # **굵은글씨** 패턴 완전 제거
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    text = text.replace('**', '').replace('*', '').replace('###', '').replace('##', '').replace('#', '')
    # 연속 줄바꿈 정리
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = text.strip()

    # 본문 첫 줄이 소제목과 동일하면 제거 (AI가 소제목을 본문 시작에 또 박는 경우)
    if subtopic:
        sub_clean = re.sub(r'[「」"\'\s\.\?!]+', '', subtopic).strip()
        lines = text.split('\n')
        if lines:
            first_line_clean = re.sub(r'[「」"\'\s\.\?!]+', '', lines[0]).strip()
            # 정확 일치 또는 거의 같은 경우 (90% 이상 매칭)
            if first_line_clean == sub_clean or (sub_clean and sub_clean in first_line_clean and len(first_line_clean) <= len(sub_clean) * 1.2):
                text = '\n'.join(lines[1:]).lstrip('\n').strip()
    return text

def parse_json(response):
    """JSON 파싱 - 개선된 에러 처리"""
    if not response:
        return None
    try:
        # 먼저 전체 응답에서 JSON 블록 찾기
        json_match = re.search(r'```json\s*([\s\S]*?)\s*```', response)
        if json_match:
            return json.loads(json_match.group(1))

        # JSON 블록이 없으면 중괄호로 시작하는 객체 찾기
        match = re.search(r'\{[\s\S]*\}', response)
        if match:
            json_str = match.group()
            # 불완전한 JSON 수정 시도
            json_str = re.sub(r',\s*}', '}', json_str)  # 마지막 쉼표 제거
            json_str = re.sub(r',\s*]', ']', json_str)  # 배열 마지막 쉼표 제거
            return json.loads(json_str)
    except json.JSONDecodeError as e:
        st.warning(f"JSON 파싱 경고: {str(e)[:50]}")
    except Exception as e:
        st.warning(f"파싱 오류: {str(e)[:50]}")
    return None

def ask_ai(prompt, temp=0.7, ensure_quality=False):
    """Claude API 호출

    ensure_quality=True 시 Haiku 선택해도 자동으로 Sonnet 4.5로 업그레이드.
    목차/본문/프롤로그/에필로그/컨셉 생성처럼 정교한 프롬프트를 따라야 하는 작업에 사용.
    """
    api_key = get_api_key()
    if not api_key:
        st.error("Claude API 키를 입력해주세요")
        return None

    if not CLAUDE_AVAILABLE:
        st.error("anthropic 패키지가 설치되지 않았습니다. pip install anthropic")
        return None

    # 선택된 모델 가져오기 (기본값: Sonnet 4.5)
    user_model = st.session_state.get('claude_model', 'claude-sonnet-4-5')

    # 핵심 생성 작업은 Haiku 자동 업그레이드 (품질 일관성 보장)
    if ensure_quality and 'haiku' in user_model.lower():
        model = 'claude-sonnet-4-5'
    else:
        model = user_model

    try:
        client = anthropic.Anthropic(api_key=api_key)
        message = client.messages.create(
            model=model,
            max_tokens=8000,
            temperature=temp,
            messages=[
                {"role": "user", "content": prompt}
            ]
        )
        return message.content[0].text
    except anthropic.AuthenticationError:
        st.error("API 키가 유효하지 않습니다. Claude API 키를 확인해주세요.")
        return None
    except anthropic.RateLimitError:
        st.error("API 할당량이 초과되었습니다. 잠시 후 다시 시도해주세요.")
        return None
    except anthropic.BadRequestError as e:
        try:
            err_msg = str(e).encode('utf-8', errors='ignore').decode('utf-8')[:100]
        except:
            err_msg = "요청 형식 오류"
        st.error(f"요청 오류: {err_msg}")
        return None
    except Exception as e:
        try:
            err_msg = str(e).encode('utf-8', errors='ignore').decode('utf-8')[:100]
        except:
            err_msg = "알 수 없는 오류"
        st.error(f"AI 오류: {err_msg}")
        return None

def generate_cover_image_gemini(title, subtitle, theme_keywords):
    """Google Gemini로 표지 배경 이미지 생성"""

    api_key = get_api_key()
    if not api_key:
        return None, "Gemini API 키가 필요합니다."

    if not IMAGEN_AVAILABLE:
        return None, "google-genai 패키지가 필요합니다: pip install google-genai"

    try:
        client = google_genai.Client(api_key=api_key)

        # 베스트셀러급 고급 표지 프롬프트 - 텍스트 절대 금지 강조
        prompt = f"""Create an ABSTRACT background image for a book cover.

Theme keywords: {theme_keywords}

STYLE: Dark, moody, cinematic atmosphere. Abstract shapes, gradients, smoke, light rays, or geometric patterns. Luxury aesthetic with gold/amber accent lighting on deep black background.

CRITICAL RULES:
- ONLY abstract visuals: smoke, light, shadows, gradients, textures
- NO objects, NO people, NO faces, NO hands
- NO text, NO letters, NO words, NO numbers, NO symbols, NO characters of ANY language
- NO Korean, NO English, NO Chinese, NO Japanese characters
- Pure abstract art only

OUTPUT: Dark dramatic background with subtle golden light accents, suitable for text overlay."""

        # Gemini 이미지 생성
        response = client.models.generate_content(
            model='gemini-2.0-flash-exp-image-generation',
            contents=prompt,
            config=genai_types.GenerateContentConfig(
                response_modalities=['IMAGE', 'TEXT']
            )
        )

        if response.candidates and response.candidates[0].content.parts:
            for part in response.candidates[0].content.parts:
                if hasattr(part, 'inline_data') and part.inline_data:
                    image_base64 = base64.b64encode(part.inline_data.data).decode('utf-8')
                    return image_base64, None

        return None, "이미지가 생성되지 않았습니다."

    except Exception as e:
        error_msg = str(e)
        if "quota" in error_msg.lower() or "limit" in error_msg.lower():
            return None, "API 할당량 초과. 잠시 후 다시 시도해주세요."
        elif "safety" in error_msg.lower():
            return None, "안전 필터에 의해 차단되었습니다. 다른 키워드로 시도해주세요."
        return None, f"이미지 생성 오류: {error_msg[:80]}"

def generate_cover_prompt_ai(title, subtitle, topic):
    """AI가 표지 디자인 컨셉과 이미지 프롬프트 생성"""
    prompt = f"""당신은 베스트셀러 책 표지 디자이너입니다.

책 제목: {title}
부제: {subtitle}
주제: {topic}

이 책의 표지 이미지를 위한 영문 프롬프트를 만들어주세요.

[요구사항]
1. 실제 베스트셀러 표지 스타일 분석 기반
2. 제목의 핵심 메시지를 시각적으로 표현
3. 고급스럽고 전문적인 느낌
4. 텍스트 오버레이를 위한 여백 고려
5. 추상적이거나 상징적인 이미지

[출력 형식]
IMAGE_PROMPT: (영문 이미지 생성 프롬프트, 50단어 이내)
COLOR_SCHEME: (추천 컬러 팔레트, 예: dark, gold, minimal)
STYLE: (디자인 스타일, 예: editorial, bold, elegant)

영문 프롬프트만 출력하세요. 한국어 설명 불필요."""

    result = ask_ai(prompt, temp=0.7)
    if result:
        # 파싱
        image_prompt = ""
        color_scheme = "dark"
        style = "editorial"

        for line in result.split('\n'):
            if 'IMAGE_PROMPT:' in line:
                image_prompt = line.split('IMAGE_PROMPT:')[-1].strip()
            elif 'COLOR_SCHEME:' in line:
                color_scheme = line.split('COLOR_SCHEME:')[-1].strip().lower()
            elif 'STYLE:' in line:
                style = line.split('STYLE:')[-1].strip().lower()

        return image_prompt, color_scheme, style
    return None, "dark", "editorial"


# ==========================================
# 고급 표지 렌더러 (외부 API 불필요, 벡터 SVG)
# ==========================================
COVER_TEMPLATES = {
    "noir_gold": "느와르 골드 — 블랙 + 금박",
    "editorial": "에디토리얼 — 아이보리 + 세리프",
    "deep_navy": "딥 네이비 — 남색 + 아이보리",
}


def _wrap_title_lines(title, max_chars=7, max_lines=3):
    """제목을 표지용으로 줄바꿈 (한글 글자수 기준)"""
    title = (title or "").strip()
    if not title:
        return ["제목"]
    words = title.split()
    lines, cur = [], ""
    for w in words:
        cand = (cur + " " + w).strip()
        if not cur or len(cand) <= max_chars:
            cur = cand
        else:
            lines.append(cur)
            cur = w
    if cur:
        lines.append(cur)
    # 띄어쓰기 없는 긴 단어는 강제 줄바꿈
    if len(lines) == 1 and len(lines[0]) > max_chars:
        s = lines[0]
        lines = [s[i:i + max_chars] for i in range(0, len(s), max_chars)]
    return lines[:max_lines]


import tempfile as _tempfile
import subprocess as _subprocess

# 표지용 한글 폰트 캐시 디렉터리
_KR_FONT_DIR = os.path.join(_tempfile.gettempdir(), "cashmaker_fonts")
_KR_FONT_CACHE = {}

# 한글 폰트 다운로드 후보(여러 미러). 환경에 설치된 폰트가 없을 때 사용.
_KR_FONT_SOURCES = {
    "NanumMyeongjo-Bold.ttf": [
        "https://github.com/google/fonts/raw/main/ofl/nanummyeongjo/NanumMyeongjo-Bold.ttf",
        "https://raw.githubusercontent.com/google/fonts/main/ofl/nanummyeongjo/NanumMyeongjo-Bold.ttf",
        "https://cdn.jsdelivr.net/gh/google/fonts@main/ofl/nanummyeongjo/NanumMyeongjo-Bold.ttf",
    ],
    "NanumMyeongjo-Regular.ttf": [
        "https://github.com/google/fonts/raw/main/ofl/nanummyeongjo/NanumMyeongjo-Regular.ttf",
        "https://raw.githubusercontent.com/google/fonts/main/ofl/nanummyeongjo/NanumMyeongjo-Regular.ttf",
        "https://cdn.jsdelivr.net/gh/google/fonts@main/ofl/nanummyeongjo/NanumMyeongjo-Regular.ttf",
    ],
    "NanumGothic-Bold.ttf": [
        "https://github.com/google/fonts/raw/main/ofl/nanumgothic/NanumGothic-Bold.ttf",
        "https://raw.githubusercontent.com/google/fonts/main/ofl/nanumgothic/NanumGothic-Bold.ttf",
        "https://cdn.jsdelivr.net/gh/google/fonts@main/ofl/nanumgothic/NanumGothic-Bold.ttf",
    ],
    "NanumGothic-Regular.ttf": [
        "https://github.com/google/fonts/raw/main/ofl/nanumgothic/NanumGothic-Regular.ttf",
        "https://raw.githubusercontent.com/google/fonts/main/ofl/nanumgothic/NanumGothic-Regular.ttf",
        "https://cdn.jsdelivr.net/gh/google/fonts@main/ofl/nanumgothic/NanumGothic-Regular.ttf",
    ],
}


def _fc_korean_fonts():
    """fontconfig(fc-list)로 시스템에 설치된 한글 폰트 경로를 찾는다."""
    try:
        out = _subprocess.run(
            ["fc-list", ":lang=ko", "file"],
            capture_output=True, text=True, timeout=5
        ).stdout
        paths = []
        for line in out.splitlines():
            p = line.split(":")[0].strip()
            if p and os.path.exists(p):
                paths.append(p)
        # 명조(serif) 계열을 앞쪽에 배치
        paths.sort(key=lambda p: (0 if ("Myeongjo" in p or "Serif" in p or "Batang" in p) else 1))
        return paths
    except Exception:
        return []


def _download_kr_font(filename):
    """한글 폰트를 캐시 디렉터리에 내려받고 경로를 반환. 실패 시 None."""
    try:
        os.makedirs(_KR_FONT_DIR, exist_ok=True)
    except Exception:
        return None
    dest = os.path.join(_KR_FONT_DIR, filename)
    if os.path.exists(dest) and os.path.getsize(dest) > 10000:
        return dest
    for url in _KR_FONT_SOURCES.get(filename, []):
        try:
            r = requests.get(url, timeout=15)
            if r.status_code == 200 and len(r.content) > 10000:
                with open(dest, "wb") as f:
                    f.write(r.content)
                return dest
        except Exception:
            continue
    return None


def _kr_font_path(serif=True):
    """사용 가능한 한글 폰트 경로 반환: 설치된 폰트 → 다운로드 순."""
    key = "serif" if serif else "sans"
    if key in _KR_FONT_CACHE:
        return _KR_FONT_CACHE[key]
    # 1) 시스템에 설치된 한글 폰트
    fc = _fc_korean_fonts()
    if fc:
        # serif 요청이면 명조 우선, 없으면 첫 번째
        pick = fc[0]
        _KR_FONT_CACHE[key] = pick
        return pick
    # 2) 다운로드
    fname = "NanumMyeongjo-Bold.ttf" if serif else "NanumGothic-Bold.ttf"
    p = _download_kr_font(fname)
    _KR_FONT_CACHE[key] = p
    return p


def _cover_font(candidates, size):
    """후보 경로들 중 존재하는 폰트를 로드. 없으면 한글 폰트를 자동 확보."""
    if not PILLOW_AVAILABLE:
        return None
    # 1) 명시된 시스템 경로 후보
    for path, idx in candidates:
        if path and os.path.exists(path):
            try:
                return ImageFont.truetype(path, size, index=idx)
            except Exception:
                try:
                    return ImageFont.truetype(path, size)
                except Exception:
                    continue
    # 2) 후보가 명조/serif 계열인지 판별
    is_serif = any(
        ("Serif" in p) or ("Myeongjo" in p) or ("batang" in p.lower())
        for p, _ in candidates
    )
    # 3) fontconfig 또는 다운로드로 한글 폰트 확보
    kr = _kr_font_path(serif=is_serif)
    if kr:
        try:
            return ImageFont.truetype(kr, size)
        except Exception:
            pass
    # 4) 최후의 수단(한글 미지원, 박스로 보일 수 있음)
    try:
        return ImageFont.load_default()
    except Exception:
        return None


def render_cover_png(path, title, subtitle, author, brand):
    """Pillow로 출판물급 표지 PNG 생성. 성공 시 path 반환, 실패 시 None."""
    if not PILLOW_AVAILABLE:
        return None
    try:
        W, H = 1240, 1748  # ~150dpi A5

        SERIF = [
            ("/usr/share/fonts/opentype/noto/NotoSerifCJK-Bold.ttc", 0),
            ("C:/Windows/Fonts/batang.ttc", 0),
            ("C:/Windows/Fonts/Batang.ttc", 0),
            ("/System/Library/Fonts/AppleSDGothicNeo.ttc", 0),
            ("/usr/share/fonts/truetype/nanum/NanumMyeongjo.ttf", 0),
        ]
        SERIF_MED = [
            ("/usr/share/fonts/opentype/noto/NotoSerifCJK-Medium.ttc", 0),
            ("C:/Windows/Fonts/batang.ttc", 0),
            ("/usr/share/fonts/truetype/nanum/NanumMyeongjo.ttf", 0),
        ]
        SANS = [
            ("/usr/share/fonts/opentype/noto/NotoSansCJK-Medium.ttc", 0),
            ("C:/Windows/Fonts/malgun.ttf", 0),
            ("/System/Library/Fonts/AppleSDGothicNeo.ttc", 0),
            ("/usr/share/fonts/truetype/nanum/NanumGothic.ttf", 0),
        ]

        def wrap(t, mx):
            words = t.split(); lines, cur = [], ""
            for w in words:
                cand = (cur + " " + w).strip()
                if not cur or len(cand) <= mx:
                    cur = cand
                else:
                    lines.append(cur); cur = w
            if cur:
                lines.append(cur)
            out = []
            for ln in lines:
                while len(ln) > mx + 2:
                    out.append(ln[:mx]); ln = ln[mx:]
                out.append(ln)
            return out

        img = Image.new("RGB", (W, H), (16, 14, 11))
        d = ImageDraw.Draw(img)
        top, bot = (26, 23, 16), (9, 8, 6)
        for y in range(H):
            t = y / H
            d.line([(0, y), (W, y)], fill=(
                int(top[0] + (bot[0] - top[0]) * t),
                int(top[1] + (bot[1] - top[1]) * t),
                int(top[2] + (bot[2] - top[2]) * t)))
        GOLD = (196, 154, 74); CREAM = (244, 239, 227)
        SUBC = (201, 192, 174); AUTHC = (237, 231, 216); DIM = (138, 130, 112)

        d.rectangle([56, 56, W - 56, H - 56], outline=GOLD, width=2)
        d.rectangle([68, 68, W - 68, H - 68], outline=GOLD, width=1)
        cx = W / 2

        def centered(y, text, font, fill, tracking=0):
            if not font or not text:
                return
            if tracking <= 0:
                b = d.textbbox((0, 0), text, font=font)
                d.text((cx - (b[2] - b[0]) / 2, y), text, font=font, fill=fill)
                return
            widths = [d.textbbox((0, 0), ch, font=font)[2] - d.textbbox((0, 0), ch, font=font)[0] for ch in text]
            total = sum(widths) + tracking * (len(text) - 1)
            x = cx - total / 2
            for ch, wch in zip(text, widths):
                d.text((x, y), ch, font=font, fill=fill); x += wch + tracking

        centered(160, brand or "", _cover_font(SANS, 32), GOLD, tracking=14)
        d.line([(cx - 46, 232), (cx + 46, 232)], fill=GOLD, width=2)

        t_lines = wrap(title, 9)
        longest = max((len(l) for l in t_lines), default=4)
        fs = 116 if longest <= 8 else 98
        ftitle = _cover_font(SERIF, fs)
        line_h = int(fs * 1.30)
        block_h = line_h * len(t_lines)
        ty = int(H * 0.44 - block_h / 2)
        for i, ln in enumerate(t_lines):
            centered(ty + i * line_h, ln, ftitle, CREAM, tracking=2)

        sep_y = ty + block_h + 46
        d.line([(cx - 70, sep_y), (cx + 70, sep_y)], fill=GOLD, width=2)
        if subtitle:
            centered(sep_y + 44, subtitle, _cover_font(SERIF_MED, 40), SUBC, tracking=2)

        d.line([(cx - 30, H - 290), (cx + 30, H - 290)], fill=GOLD, width=1)
        centered(H - 252, author or "저자", _cover_font(SANS, 42), AUTHC, tracking=8)
        centered(H - 190, "지음", _cover_font(SANS, 26), DIM, tracking=4)

        img.save(path, "PNG")
        return path
    except Exception:
        return None


def build_cover_svg(template, title, subtitle, author):
    """베스트셀러급 고급 표지 SVG 생성 (1600x2560, 무손실 벡터)"""
    t_lines = _wrap_title_lines(title, max_chars=7, max_lines=3)
    longest = max((len(l) for l in t_lines), default=4)
    if longest <= 4:
        fs = 230
    elif longest <= 6:
        fs = 185
    else:
        fs = 150

    sub = html.escape((subtitle or "").strip())
    auth = html.escape((author or "").strip())
    esc_lines = [html.escape(l) for l in t_lines]

    line_h = int(fs * 1.22)
    n = len(esc_lines)
    title_top = 1180 - int((n - 1) * line_h / 2)
    tspans = "".join(
        f'<tspan x="800" dy="{0 if i == 0 else line_h}">{ln}</tspan>'
        for i, ln in enumerate(esc_lines)
    )

    fonts = (
        '<style><![CDATA['
        "@import url('https://fonts.googleapis.com/css2?family=Noto+Serif+KR:wght@400;500;600;700;900&display=swap');"
        ']]></style>'
    )
    serif = "'Noto Serif KR','Nanum Myeongjo',serif"

    if template == "editorial":
        ink, accent, subcol = '#1A1A1A', '#8A7A55', '#5A5346'
        bg = (
            '<rect width="1600" height="2560" fill="#F4F0E8"/>'
            f'<rect x="78" y="78" width="1444" height="2404" fill="none" stroke="{ink}" stroke-width="2"/>'
        )
        deco = (
            f'<line x1="610" y1="{title_top - 130}" x2="990" y2="{title_top - 130}" stroke="{accent}" stroke-width="3"/>'
            f'<line x1="660" y1="1560" x2="940" y2="1560" stroke="{accent}" stroke-width="2"/>'
        )
        tw = 600
    elif template == "deep_navy":
        ink, accent, subcol = '#F3ECD9', '#C9A24B', '#C9C3B2'
        bg = (
            '<defs><radialGradient id="bgnavy" cx="50%" cy="38%" r="72%">'
            '<stop offset="0%" stop-color="#1B3457"/><stop offset="100%" stop-color="#0B1626"/>'
            '</radialGradient></defs>'
            '<rect width="1600" height="2560" fill="url(#bgnavy)"/>'
            f'<rect x="70" y="70" width="1460" height="2420" fill="none" stroke="{accent}" stroke-width="2" opacity="0.7"/>'
        )
        deco = (
            f'<line x1="660" y1="{title_top - 120}" x2="940" y2="{title_top - 120}" stroke="{accent}" stroke-width="3"/>'
            f'<rect x="780" y="1540" width="40" height="40" transform="rotate(45 800 1560)" fill="none" stroke="{accent}" stroke-width="2"/>'
        )
        tw = 600
    else:  # noir_gold
        ink, accent, subcol = '#E9D9A8', '#C9A24B', '#B8B3A8'
        bg = (
            '<defs><radialGradient id="bgnoir" cx="50%" cy="40%" r="75%">'
            '<stop offset="0%" stop-color="#1C1A14"/><stop offset="60%" stop-color="#0E0D0B"/>'
            '<stop offset="100%" stop-color="#060606"/></radialGradient></defs>'
            '<rect width="1600" height="2560" fill="url(#bgnoir)"/>'
            f'<rect x="64" y="64" width="1472" height="2432" fill="none" stroke="{accent}" stroke-width="1.5" opacity="0.55"/>'
            f'<rect x="84" y="84" width="1432" height="2392" fill="none" stroke="{accent}" stroke-width="0.75" opacity="0.35"/>'
        )
        deco = (
            f'<line x1="640" y1="{title_top - 120}" x2="960" y2="{title_top - 120}" stroke="{accent}" stroke-width="2"/>'
            f'<line x1="700" y1="1560" x2="900" y2="1560" stroke="{accent}" stroke-width="1.5" opacity="0.8"/>'
        )
        tw = 700

    sub_svg = (
        f'<text x="800" y="1640" text-anchor="middle" font-family="{serif}" '
        f'font-size="46" fill="{subcol}" letter-spacing="6">{sub}</text>'
    ) if sub else ""
    auth_svg = (
        f'<text x="800" y="2360" text-anchor="middle" font-family="{serif}" '
        f'font-size="40" fill="{accent}" letter-spacing="10">{auth}</text>'
    ) if auth else ""

    return (
        '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 1600 2560" '
        'width="100%" preserveAspectRatio="xMidYMid meet">'
        f'{fonts}{bg}{deco}'
        f'<text x="800" y="{title_top}" text-anchor="middle" font-family="{serif}" '
        f'font-size="{fs}" font-weight="{tw}" fill="{ink}" letter-spacing="3">{tspans}</text>'
        f'{sub_svg}{auth_svg}</svg>'
    )


def estimate_docx_pages(chapters, outline):
    """워드(A5) 실제 출력 구조를 그대로 모사한 페이지 수 추정.

    워드는 본문 10.5pt / 줄간격 1.85 / A5 + 소제목마다 새 페이지로 나뉘고,
    표지·판권·프롤로그·에필로그·저자소개·챕터 오프너가 더해진다.
    기존의 '글자수//500'은 이 구조를 반영하지 않아 워드 페이지 수와 크게 어긋났다.
    """
    import math
    CPP = 330  # A5 11.5pt/1.85 + 소제목별 페이지 나눔 실측(페이지당 약 320자) 반영
    FRONT = 2  # 표지 + 판권
    PROLOGUE = 2
    EPILOGUE = 2
    AUTHOR = 1

    pages = 0
    chapter_with_content = 0
    for ch in (outline or []):
        cd = (chapters or {}).get(ch)
        if not cd:
            continue
        sub_pages = 0
        has_content = False
        for s in cd.get('subtopics', []):
            content = cd.get('subtopic_data', {}).get(s, {}).get('content', '')
            if content:
                has_content = True
                chars = len(content.replace(' ', '').replace('\n', ''))
                sub_pages += max(1, math.ceil(chars / CPP))  # 소제목은 새 페이지에서 시작
        if has_content:
            chapter_with_content += 1
            pages += 1          # 챕터 오프너 페이지
            pages += sub_pages

    if chapter_with_content == 0:
        return 0
    return FRONT + PROLOGUE + pages + EPILOGUE + AUTHOR


def extract_video_id(url):
    """YouTube URL에서 video ID 추출"""
    patterns = [
        r'(?:youtube\.com\/watch\?v=|youtu\.be\/|youtube\.com\/embed\/)([^&\n?#]+)',
        r'youtube\.com\/watch\?.*v=([^&\n?#]+)',
    ]
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    return None

def get_youtube_transcript(video_id):
    """YouTube 자막 가져오기"""
    if not YOUTUBE_TRANSCRIPT_AVAILABLE:
        return None, "youtube-transcript-api가 설치되지 않았습니다. pip install youtube-transcript-api"

    try:
        # 한국어 자막 우선, 없으면 영어, 없으면 자동생성
        transcript_list = YouTubeTranscriptApi.list_transcripts(video_id)

        transcript = None
        # 수동 자막 먼저 시도
        for lang in ['ko', 'en']:
            try:
                transcript = transcript_list.find_transcript([lang])
                break
            except:
                continue

        # 수동 자막 없으면 자동 생성 자막
        if not transcript:
            try:
                transcript = transcript_list.find_generated_transcript(['ko', 'en'])
            except:
                # 아무 자막이나 가져오기
                for t in transcript_list:
                    transcript = t
                    break

        if transcript:
            fetched = transcript.fetch()
            full_text = ' '.join([item['text'] for item in fetched])
            return full_text, None
        else:
            return None, "자막을 찾을 수 없습니다"

    except Exception as e:
        return None, f"자막 추출 오류: {str(e)[:100]}"

def analyze_youtube_video_direct(url):
    """YouTube 영상 자막 기반 분석 (빠르고 정확)"""
    api_key = get_api_key()
    if not api_key:
        st.error("API 키를 입력해주세요")
        return None

    # 1. Video ID 추출
    video_id = extract_video_id(url)
    if not video_id:
        st.error("올바른 YouTube URL이 아닙니다")
        return None

    # 2. 자막 가져오기
    transcript, error = get_youtube_transcript(video_id)
    if error:
        st.warning(f"자막 추출 실패: {error}")
        st.info("자막이 없는 영상입니다. 텍스트 입력으로 직접 내용을 입력해주세요.")
        return None

    if not transcript or len(transcript) < 50:
        st.warning("자막 내용이 너무 짧습니다")
        return None

    # 3. 자막 기반 분석
    prompt = f"""다음은 YouTube 영상의 자막입니다. 이 내용을 분석해주세요.

[자막 내용]
{transcript[:15000]}

[분석 요청]
위 자막 내용을 바탕으로 분석해주세요. 자막에 없는 내용은 추측하지 마세요.

JSON 형식으로 응답:
{{
    "title": "영상의 핵심 주제 (자막 기반 추론)",
    "creator": "알 수 없음",
    "main_topic": "메인 주제 한 줄 요약",
    "key_points": ["핵심 포인트 1", "핵심 포인트 2", "핵심 포인트 3", "핵심 포인트 4", "핵심 포인트 5"],
    "detailed_notes": ["상세 내용 1", "상세 내용 2", "상세 내용 3"],
    "actionable_tips": ["실천 팁 1", "실천 팁 2", "실천 팁 3"],
    "quotes": ["인상적인 문장 1", "인상적인 문장 2"],
    "vocabulary": [{{"term": "용어", "definition": "설명"}}],
    "study_questions": ["학습 질문 1", "학습 질문 2"],
    "summary": "전체 내용 5-7문장 요약"
}}"""

    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-2.0-flash')
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        st.error(f"분석 오류: {str(e)[:150]}")
        return None

def get_full_content():
    full = ""
    for ch in st.session_state.get('outline', []):
        if ch in st.session_state.get('chapters', {}):
            ch_data = st.session_state['chapters'][ch]
            ch_content = ""
            for s in ch_data.get('subtopics', []):
                c = ch_data.get('subtopic_data', {}).get(s, {}).get('content', '')
                if c:
                    ch_content += f"\n\n【{s}】\n\n{clean_content(c)}"
            if ch_content:
                full += f"\n\n{'='*50}\n{ch}\n{'='*50}{ch_content}"
    return full.strip()

# ==========================================
# 전자책 워드 디자인 폰트 (출판물 톤)
# 본문 세리프 = 바탕(모든 Windows 기본), 라틴 = Georgia
# 디스플레이 산세리프 = 맑은 고딕, 라틴 = Arial
# ==========================================
EBOOK_SERIF_KR = '바탕'
EBOOK_SERIF_LATIN = 'Georgia'
EBOOK_SANS_KR = '맑은 고딕'
EBOOK_SANS_LATIN = 'Arial'

# 출판 톤 색상 팔레트
_INK = (28, 28, 30)      # 본문/제목 (거의 검정)
_SOFT = (92, 92, 98)     # 부제/캡션
_FAINT = (165, 165, 170) # 라벨/번호
_HAIR = (200, 200, 204)  # 가는 선
_GOLD = (150, 120, 60)   # 차분한 골드 포인트


def create_ebook_docx(title, subtitle, author, chapters_data, outline, interview_data=None):
    """베스트셀러 출판물 스타일 워드 문서 생성 (프리미엄 에디토리얼)."""
    if not DOCX_AVAILABLE:
        return None, "python-docx 패키지가 필요합니다: pip install python-docx"

    try:
        BRAND = "CASHMAKER"
        doc = Document()

        # 기본 문단 간격 0 (빈 문단이 멋대로 커지는 것 방지)
        normal = doc.styles['Normal']
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(0)
        normal.paragraph_format.line_spacing = 1.0

        def _apply_section(sec):
            sec.page_width = Cm(14.8)
            sec.page_height = Cm(21)
            sec.left_margin = Cm(2.0)
            sec.right_margin = Cm(2.0)
            sec.top_margin = Cm(2.3)
            sec.bottom_margin = Cm(2.2)

        _apply_section(doc.sections[0])

        # ── 폰트/런 ──
        def set_font(run, size, bold=False, color=None, italic=False,
                     serif=False, track=None):
            run.font.size = Pt(size)
            fam_latin = EBOOK_SERIF_LATIN if serif else EBOOK_SANS_LATIN
            fam_kr = EBOOK_SERIF_KR if serif else EBOOK_SANS_KR
            run.font.name = fam_latin
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.find(qn('w:rFonts'))
            if rfonts is None:
                rfonts = OxmlElement('w:rFonts')
                rpr.append(rfonts)
            rfonts.set(qn('w:ascii'), fam_latin)
            rfonts.set(qn('w:hAnsi'), fam_latin)
            rfonts.set(qn('w:eastAsia'), fam_kr)
            run.bold = bold
            run.italic = italic
            if color:
                run.font.color.rgb = RGBColor(*color)
            if track is not None:
                sp = OxmlElement('w:spacing')
                sp.set(qn('w:val'), str(int(track * 20)))
                rpr.append(sp)

        def vspace(pts):
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.space_before = Pt(0); pf.space_after = Pt(0)
            pf.line_spacing = Pt(pts)
            return p

        def track_text(text, n=1):
            return (" " * n).join(list(str(text)))

        def hairline(align=WD_ALIGN_PARAGRAPH.CENTER, width_cm=None, color=_HAIR,
                     space_before=0, space_after=12, size=6):
            p = doc.add_paragraph()
            p.alignment = align
            pf = p.paragraph_format
            pf.space_before = Pt(space_before); pf.space_after = Pt(space_after)
            if width_cm and align == WD_ALIGN_PARAGRAPH.CENTER:
                side = (10.8 - width_cm) / 2
                if side > 0:
                    pf.left_indent = Cm(side); pf.right_indent = Cm(side)
            pPr = p._p.get_or_add_pPr()
            pbdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), str(size))
            bottom.set(qn('w:space'), '1')
            _col = color if isinstance(color, str) else '{:02X}{:02X}{:02X}'.format(*color)
            bottom.set(qn('w:color'), _col)
            pbdr.append(bottom)
            pPr.append(pbdr)
            return p

        def ornament(glyph="✦", color=_FAINT, size=11, space=22):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(space)
            p.paragraph_format.space_after = Pt(space)
            set_font(p.add_run(glyph), size, serif=False, color=color)
            return p

        def _shade_para(p, hex_fill):
            """문단 배경색 칠하기."""
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_fill)
            pPr.append(shd)

        def _box_border(p, color_hex, sides=('top', 'bottom', 'left', 'right'), sz=6):
            """문단 사면(또는 일부) 테두리 — 박스 효과."""
            pPr = p._p.get_or_add_pPr()
            pbdr = OxmlElement('w:pBdr')
            for s in sides:
                el = OxmlElement(f'w:{s}')
                el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), str(sz))
                el.set(qn('w:space'), '6'); el.set(qn('w:color'), color_hex)
                pbdr.append(el)
            pPr.append(pbdr)

        def key_summary_box(items, chapter_no):
            """챕터 끝 '핵심 정리' 요약 박스 (베스트셀러식 정리 코너)."""
            # 제목 줄 (다크 배경 + 골드 글자)
            title_p = doc.add_paragraph()
            title_p.paragraph_format.space_before = Pt(20); title_p.paragraph_format.space_after = Pt(0)
            title_p.paragraph_format.left_indent = Pt(2); title_p.paragraph_format.right_indent = Pt(2)
            _shade_para(title_p, '2E2A24')
            r = title_p.add_run(f"  CHAPTER {chapter_no} · 핵심 정리  ")
            set_font(r, 11, bold=True, serif=False, color=(201, 162, 75), track=0.5)
            title_p.paragraph_format.space_after = Pt(0)
            # 항목들 (연한 배경 박스 안)
            for k, item in enumerate(items):
                ip = doc.add_paragraph()
                ip.paragraph_format.left_indent = Cm(0.2); ip.paragraph_format.right_indent = Cm(0.2)
                ip.paragraph_format.space_before = Pt(0)
                ip.paragraph_format.space_after = Pt(2 if k < len(items) - 1 else 0)
                ip.paragraph_format.line_spacing = 1.5
                _shade_para(ip, 'F7F4EE')
                num_r = ip.add_run(f"  {k+1}  ")
                set_font(num_r, 10.5, bold=True, serif=False, color=(150, 120, 60))
                txt_r = ip.add_run(item.strip() + "  ")
                set_font(txt_r, 10.5, serif=True, color=_INK)
            # 박스 하단 골드 라인
            hairline(align=WD_ALIGN_PARAGRAPH.LEFT, color=_GOLD, space_before=0, space_after=16, size=10)

        def body_paragraph(text, first=False, size=11.5, ls=1.85):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_font(p.add_run(text), size, serif=True, color=_INK)
            pf = p.paragraph_format
            pf.line_spacing = ls; pf.space_after = Pt(16)
            if not first:
                pf.first_line_indent = Cm(0.6)
            return p

        def drop_cap(letter, rest, lines=3, size=11.5, ls=1.85):
            # 드롭캡 제거: 첫 글자 크기 들쭉날쭉/과대 문제 → 모든 첫 문단을 동일하게 처리
            return body_paragraph((letter or "") + (rest or ""), first=True, size=size, ls=ls)

        def _cell_shade(cell, hex_fill):
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_fill)
            cell._tc.get_or_add_tcPr().append(shd)

        def _cell_para(cell, text, size, *, serif=True, bold=False, italic=False,
                       color=(255, 255, 255), align=WD_ALIGN_PARAGRAPH.CENTER,
                       track=None, space_before=0, space_after=0, ls=1.2, first=True):
            p = cell.add_paragraph() if not first else cell.paragraphs[0]
            p.alignment = align
            pf = p.paragraph_format
            pf.space_before = Pt(space_before); pf.space_after = Pt(space_after); pf.line_spacing = ls
            if text:
                set_font(p.add_run(text), size, serif=serif, bold=bold, italic=italic, color=color, track=track)
            return p

        def _no_cell_borders(cell):
            tcPr = cell._tc.get_or_add_tcPr()
            tcb = OxmlElement('w:tcBorders')
            for s in ('top', 'bottom', 'left', 'right'):
                e = OxmlElement(f'w:{s}'); e.set(qn('w:val'), 'nil'); tcb.append(e)
            tcPr.append(tcb)

        def page_number_footer(section):
            footer = section.footer
            footer.is_linked_to_previous = False
            p = footer.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            set_font(run, 9, serif=True, color=_SOFT)
            f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
            instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'PAGE'
            f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end')
            run._r.append(f1); run._r.append(instr); run._r.append(f2)

        def full_page_cover(title, subtitle, author, brand):
            """표지: Pillow PNG(출판물급)를 우선 시도, 실패 시 텍스트 표지로 폴백."""
            # 1) PNG 표지 시도 — 페이지에 꽉 차게 삽입
            try:
                import tempfile
                tmp_png = os.path.join(tempfile.gettempdir(), f"writey_cover_{uuid.uuid4().hex[:8]}.png")
                made = render_cover_png(tmp_png, title, subtitle, author, brand)
                if made and os.path.exists(made):
                    sec0 = doc.sections[0]
                    # 표지 섹션은 여백 0으로 → 이미지가 페이지 전체를 덮음
                    old = (sec0.left_margin, sec0.right_margin, sec0.top_margin, sec0.bottom_margin)
                    sec0.left_margin = Cm(0); sec0.right_margin = Cm(0)
                    sec0.top_margin = Cm(0); sec0.bottom_margin = Cm(0)
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
                    run = p.add_run()
                    run.add_picture(made, width=Cm(14.8), height=Cm(21.0))
                    # 다음 섹션(본문)부터 정상 여백 복원
                    new_sec = doc.add_section(WD_SECTION.NEW_PAGE)
                    _apply_section(new_sec)
                    return
            except Exception:
                pass

            # 2) 폴백: 텍스트 기반 표지 (PNG 불가 환경)
            tbl = doc.add_table(rows=1, cols=1)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = tbl.rows[0].cells[0]
            cell.width = Cm(10.8)
            tbl.rows[0].height = Cm(24.2)
            tbl.rows[0].height_rule = 2  # EXACT (한 페이지 고정)
            va = OxmlElement('w:vAlign'); va.set(qn('w:val'), 'center')
            cell._tc.get_or_add_tcPr().append(va)
            _cell_shade(cell, '141210')   # 짙은 잉크 배경
            _no_cell_borders(cell)
            tcMar = OxmlElement('w:tcMar')
            for mn, mv in [('top', '420'), ('left', '500'), ('bottom', '420'), ('right', '500')]:
                m = OxmlElement(f'w:{mn}'); m.set(qn('w:w'), mv); m.set(qn('w:type'), 'dxa'); tcMar.append(m)
            cell._tc.get_or_add_tcPr().append(tcMar)

            GOLD = (198, 158, 78)
            CREAM = (245, 240, 230)
            DIM = (150, 146, 138)

            # 상단 브랜드 라인
            _cell_para(cell, track_text(brand or "", 3), 9, serif=False, color=GOLD,
                       space_before=10, space_after=2, first=True)
            _cell_para(cell, "—— ◆ ——", 9, serif=False, color=GOLD, space_after=64, first=False)

            # 제목 (아주 큰 세리프, 크림색)
            tlines = _wrap_title_lines(title, max_chars=8, max_lines=3)
            for line in tlines:
                _cell_para(cell, line, 38, serif=True, bold=True, color=CREAM,
                           ls=1.16, space_after=0, first=False)

            # 골드 구분선
            gline = cell.add_paragraph(); gline.alignment = WD_ALIGN_PARAGRAPH.CENTER
            gline.paragraph_format.space_before = Pt(20); gline.paragraph_format.space_after = Pt(16)
            set_font(gline.add_run("⎯⎯⎯⎯⎯⎯"), 13, serif=False, color=GOLD)

            # 부제 (이탤릭, 골드빛 크림) — 단일 단락, 자동 줄바꿈
            if subtitle:
                sp = cell.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sp.paragraph_format.space_before = Pt(0); sp.paragraph_format.space_after = Pt(56)
                sp.paragraph_format.line_spacing = 1.45
                sp.paragraph_format.left_indent = Cm(0.3); sp.paragraph_format.right_indent = Cm(0.3)
                set_font(sp.add_run(subtitle), 11.5, serif=True, italic=True, color=(214, 200, 168))
            else:
                _cell_para(cell, "", 6, first=False, space_after=56)

            # 하단 저자
            _cell_para(cell, f"{author or '저자'}", 12.5, serif=False, color=CREAM,
                       space_after=2, first=False)
            _cell_para(cell, "지음", 9.5, serif=False, color=DIM, space_after=10, first=False)
            doc.add_page_break()



        def running_header(section, text):
            header = section.header
            header.is_linked_to_previous = False
            p = header.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(p.add_run(track_text(text, 1)), 8, serif=False, color=_FAINT)

        def add_bookmark(paragraph, name):
            clean = re.sub(r'[^\w가-힣]', '_', name)[:40]
            bs = OxmlElement('w:bookmarkStart')
            bs.set(qn('w:id'), str(abs(hash(clean)) % 100000))
            bs.set(qn('w:name'), clean)
            be = OxmlElement('w:bookmarkEnd')
            be.set(qn('w:id'), str(abs(hash(clean)) % 100000))
            paragraph._p.insert(0, bs)
            paragraph._p.append(be)

        def add_hyperlink(paragraph, text, name, size=10, bold=False, color=_SOFT, serif=True):
            clean = re.sub(r'[^\w가-힣]', '_', name)[:40]
            hl = OxmlElement('w:hyperlink'); hl.set(qn('w:anchor'), clean)
            r = OxmlElement('w:r'); rPr = OxmlElement('w:rPr')
            rf = OxmlElement('w:rFonts')
            fam_latin = EBOOK_SERIF_LATIN if serif else EBOOK_SANS_LATIN
            fam_kr = EBOOK_SERIF_KR if serif else EBOOK_SANS_KR
            rf.set(qn('w:ascii'), fam_latin); rf.set(qn('w:hAnsi'), fam_latin); rf.set(qn('w:eastAsia'), fam_kr)
            rPr.append(rf)
            sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(int(size * 2))); rPr.append(sz)
            if bold:
                rPr.append(OxmlElement('w:b'))
            c = OxmlElement('w:color'); c.set(qn('w:val'), '{:02X}{:02X}{:02X}'.format(*color)); rPr.append(c)
            r.append(rPr)
            t = OxmlElement('w:t'); t.text = text; r.append(t)
            hl.append(r); paragraph._p.append(hl)

        # ── 표 헬퍼 (기존 로직 보존) ──
        def parse_table_data(text):
            lines = text.strip().split('\n'); table_data = []

            def _is_separator_cells(cells):
                # 모든 셀이 ---, :---:, --- 같은 구분선 기호로만 이루어졌으면 구분선 행
                if not cells:
                    return False
                for c in cells:
                    cc = c.strip()
                    if cc == '' or re.fullmatch(r'[:\-\s─━]+', cc) is None:
                        return False
                return True

            for line in lines:
                line = line.strip()
                if not line:
                    continue
                if re.match(r'^\|[\s\-:]+\|$', line):
                    continue
                if re.match(r'^[─━┌┬┐├┼┤└┴┘│┃]+$', line):
                    continue
                if re.match(r'^[\s\-]+$', line) and len(line.replace(' ', '').replace('-', '')) == 0:
                    continue
                if line.startswith('|') and line.endswith('|'):
                    cells = [c.strip() for c in line.split('|')]; cells = [c for c in cells if c]
                    if cells and not _is_separator_cells(cells):
                        table_data.append(cells)
                elif '|' in line and not line.startswith('|'):
                    cells = [c.strip() for c in line.split('|')]; cells = [c for c in cells if c]
                    if cells and not _is_separator_cells(cells):
                        table_data.append(cells)
                elif '\t' in line:
                    cells = [c.strip() for c in line.split('\t')]; cells = [c for c in cells if c]
                    if len(cells) >= 2 and not _is_separator_cells(cells):
                        table_data.append(cells)
                elif ':' in line and not line.startswith('http'):
                    parts = line.split(':', 1)
                    if len(parts) == 2 and len(parts[0]) < 30:
                        table_data.append([parts[0].strip(), parts[1].strip()])
            return table_data

        def add_premium_table(table_data):
            if not table_data or len(table_data) < 1:
                return None
            rows = len(table_data); cols = max(len(r) for r in table_data)
            is_comp = cols == 2 and rows >= 2
            table = doc.add_table(rows=rows, cols=cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            for ci in range(cols):
                for row in table.rows:
                    if ci < len(row.cells):
                        width = (Cm(3) if ci == 0 else Cm(7.8)) if is_comp else Cm(10.8 / cols)
                        row.cells[ci].width = width
            for i, row_data in enumerate(table_data):
                row = table.rows[i]; row.height = Cm(0.95)
                for j, txt in enumerate(row_data):
                    if j < cols:
                        cell = row.cells[j]; cell.text = ''; para = cell.paragraphs[0]
                        if i == 0:
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            set_font(para.add_run(str(txt)), 9, bold=True, color=(255, 255, 255), serif=False)
                            shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), '2E2A24')
                            cell._tc.get_or_add_tcPr().append(shd)
                        elif is_comp and j == 0:
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            set_font(para.add_run(str(txt)), 9, bold=True, color=(60, 50, 35), serif=False)
                            shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), 'F5F1E8')
                            cell._tc.get_or_add_tcPr().append(shd)
                        else:
                            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            para.paragraph_format.left_indent = Pt(6)
                            set_font(para.add_run(str(txt)), 9, color=(60, 60, 60), serif=True)
                            if i % 2 == 0:
                                shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), 'FAF8F3')
                                cell._tc.get_or_add_tcPr().append(shd)
                        tcPr = cell._tc.get_or_add_tcPr()
                        tcMar = OxmlElement('w:tcMar')
                        for mn, mv in [('top', '70'), ('left', '110'), ('bottom', '70'), ('right', '110')]:
                            m = OxmlElement(f'w:{mn}'); m.set(qn('w:w'), mv); m.set(qn('w:type'), 'dxa'); tcMar.append(m)
                        tcPr.append(tcMar)
                        va = OxmlElement('w:vAlign'); va.set(qn('w:val'), 'center'); tcPr.append(va)
            tbl = table._tbl
            tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
            tblBorders = OxmlElement('w:tblBorders')
            for bn in ['top', 'bottom']:
                b = OxmlElement(f'w:{bn}'); b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '8'); b.set(qn('w:color'), '2E2A24'); tblBorders.append(b)
            for bn in ['left', 'right']:
                b = OxmlElement(f'w:{bn}'); b.set(qn('w:val'), 'nil'); tblBorders.append(b)
            for bn in ['insideH', 'insideV']:
                b = OxmlElement(f'w:{bn}'); b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '4'); b.set(qn('w:color'), 'E5E0D5'); tblBorders.append(b)
            tblPr.append(tblBorders)
            sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(14)
            return table

        def process_content_with_tables(text):
            blocks = []; current = []; tbuf = []
            lines = text.split('\n'); i = 0

            def is_start(line, nxt=None):
                s = line.strip()
                if s.startswith('|') and s.endswith('|') and s.count('|') >= 2:
                    return True
                if '|' in s and len(s.split('|')) >= 2:
                    if any(p.strip() and not re.match(r'^[\s\-:]+$', p) for p in s.split('|')):
                        return True
                if nxt and ':' in s and ':' in nxt:
                    if len(s.split(':')[0].strip()) < 30 and len(nxt.split(':')[0].strip()) < 30:
                        return True
                return False

            def is_cont(line):
                s = line.strip()
                if not s:
                    return False
                if s.startswith('|') and s.endswith('|'):
                    return True
                if re.match(r'^\|[\s\-:]+\|$', s):
                    return True
                if '|' in s:
                    return True
                if ':' in s and len(s.split(':')[0].strip()) < 30:
                    return True
                return False

            while i < len(lines):
                line = lines[i]
                nxt = lines[i + 1] if i + 1 < len(lines) else None
                if is_start(line, nxt):
                    if current:
                        blocks.append(('text', '\n'.join(current))); current = []
                    tbuf = [line]; i += 1
                    while i < len(lines) and is_cont(lines[i]):
                        tbuf.append(lines[i]); i += 1
                    blocks.append(('table', '\n'.join(tbuf)))
                else:
                    current.append(line); i += 1
            if current:
                blocks.append(('text', '\n'.join(current)))
            return blocks

        # ════════════════════════ 표지 (풀 페이지 짙은 배경) ════════════════════════
        full_page_cover(title, subtitle, author, BRAND)

        # ════════════════════════ 판권 ════════════════════════
        vspace(330)
        cr = [title, "", f"지은이  {author or '저자'}", f"펴낸이  {BRAND}", "",
              "이 책의 저작권은 저자에게 있습니다.", "무단 전재와 복제를 금합니다."]
        for i, line in enumerate(cr):
            cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if line:
                set_font(cp.add_run(line), 11 if i == 0 else 9, bold=(i == 0),
                         color=_INK if i == 0 else _SOFT, serif=True)
            cp.paragraph_format.space_after = Pt(3)
        doc.add_page_break()

        # ════════════════════════ 프롤로그 ════════════════════════
        vspace(36)
        pl = doc.add_paragraph(); pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(pl.add_run(track_text("PROLOGUE", 2)), 10, serif=False, color=_FAINT)
        pl.paragraph_format.space_after = Pt(8)
        hairline(width_cm=1.4, color=_GOLD, space_after=30, size=8)

        prologue_text = None
        if interview_data:
            prologue_prompt = f"""당신은 한국 자기계발 베스트셀러 작가입니다. 독자가 첫 문장에서 "이거 내 얘기야"라고 무릎 치고, 마지막 문장에서 "다음 페이지가 너무 궁금해"라며 책장을 넘기게 만드는 프롤로그를 작성하세요.

[저자 정보 - 참고용, 그대로 복사하지 말 것]
- 분야: {interview_data.get('field', '')}
- 경력: {interview_data.get('experience_years', '')}
- 책 주제: {interview_data.get('topic', '')}
- 타겟 독자: {interview_data.get('target_reader', '')}
- 독자의 고민: {interview_data.get('target_problem', '')}
- 집필 동기: {interview_data.get('why_write', '')}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🎯 프롤로그의 단 하나의 미션
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

독자가 첫 문장부터 마지막 문장까지 한숨에 읽고, 자기 이야기처럼 공감하면서, 본문이 미치도록 궁금해서 1장으로 넘어갈 수밖에 없게 만들 것.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
✍️ 톤: 공감 후킹 + 스토리텔링 + 호기심 갭
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[5막 구조 - 베스트셀러 프롤로그 공식]

1막) 누구나 겪는 장면 (3~4문장) - 공감 끌어올리기
   - 독자가 "이거 내 얘기야"라고 즉시 떠올리는 구체적 장면
   - 시간·장소·감정의 디테일 (예: "수요일 밤 11시. 또 인스타를 켰다. 친구 OO이 부자가 돼 있었다.")
   - 일반론 절대 금지. 손에 잡히는 장면 하나

2막) 저자의 고백 (3~4문장) - "저도 그랬습니다"
   - 구체적 실패담 + 그때의 감정 (수치심, 막막함, 분노 등)
   - 진심 어린 톤 (위에서 내려다보지 말 것)
   - 약점 노출이 신뢰를 만든다

3막) 결정적 전환 (2~3문장) - 그러던 어느 날
   - 발견 / 사건 / 만남의 순간
   - "그날 알게 된 한 가지가 모든 걸 바꿨다"
   - 그게 뭔지 다 말하지 말고 살짝 가리기

4막) 약속 + 변화 (2문장) - 이 책이 줄 것
   - 본문이 다룰 핵심 변화를 구체적으로 (숫자/기간 포함)
   - 추상 X, 구체 O (예: "이 책은 그 90초가 어떻게 작동하는지 단계별로 풀어냅니다.")

5막) 본문 미끼 (1~2문장) - 페이지 넘기게 하기
   - 본문 1장이 다룰 가장 강한 장면 또는 통찰의 일부만 흘리기
   - "그런데 그 출발점은 의외의 한 가지였습니다." 같은 호기심 갭
   - "다음 페이지부터 시작됩니다" 같은 직접 안내는 절대 X

[문체]
- 합쇼체 기본 + 구어체 자연스럽게 ("~거든요", "~더라고요", "~잖아요")
- 짧은 문장 위주, 가끔 긴 문장으로 호흡 변화
- 현재형/과거형 혼용으로 생생한 장면감
- 디테일이 살아 있는 묘사 (시간, 숫자, 표정, 사물, 장소)

[분량] 600~800자

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🚫 절대 금지
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
❌ 마크다운 문법 일체 (# ## ### **굵게** > 인용 - 글머리) — 제목/라벨 출력 금지, 본문 텍스트만
❌ "프롤로그", "Prologue", "Prologue.", "들어가며" 같은 제목/라벨 출력 금지 (이미 본문 위에 'Prologue' 표시됨)
❌ 위 저자 정보를 그대로 복사 붙여넣기
❌ 자청, 역행자, 자의식 해체, 유전자 역행, 원시인, 추월차선 (특정 작가 고유 표현)
❌ 교과서 표현: "여정", "발걸음", "함께 하시길 바랍니다", "진정한", "이 책의 여정"
❌ AI 어휘: "~의 중요성", "다양한", "효과적인", "~를 통해", "이 책을 통해"
❌ 과장: "놀라운", "혁신적인", "충격적인", "어마어마한", "기적의"
❌ 뻔한 말: "포기하지 마세요", "꾸준히 하세요", "당신도 할 수 있습니다", "함께 가요"
❌ 직접 호칭: 여러분, 당신, 독자님
❌ 위에서 내려다보는 어조 (당신은 이래서 안 됩니다 X)

본문 텍스트만 출력하세요. 어떤 마크다운 헤더(#)나 '프롤로그' 라벨도 출력하지 마세요. 첫 문장부터 바로 본문이 시작되어야 합니다."""
            gen = ask_ai(prologue_prompt, 0.7, ensure_quality=True)
            if gen:
                prologue_text = gen
        if not prologue_text:
            prologue_text = ("수요일 밤 열한 시였습니다. 또 휴대폰을 들었습니다. 피드 속 친구는 벌써 두 번째 "
                             "집을 샀다고 했습니다. 저는 통장을 열어 보지도 못했습니다.\n\n"
                             "그때는 몰랐습니다. 문제는 의지가 아니라 구조였다는 걸요.\n\n"
                             "어느 날 한 줄의 계산을 마주하고서야 알았습니다. 같은 돈을 벌어도 누구는 쌓고 "
                             "누구는 흘려보내는 이유가 따로 있다는 것을요.\n\n"
                             "이 책은 그 단순한 차이를 처음부터 끝까지 풀어냅니다.")
        prologue_text = re.sub(r'^\s*#+\s*(프롤로그|Prologue|들어가며|머리말)\s*\.?\s*$', '', prologue_text, flags=re.MULTILINE | re.IGNORECASE)
        prologue_text = prologue_text.lstrip('\n').strip()
        paras = [x for x in prologue_text.split('\n\n') if x.strip()]
        for i, t in enumerate(paras):
            if i == 0 and len(t.strip()) > 1:
                drop_cap(t.strip()[0], t.strip()[1:])
            else:
                body_paragraph(t.strip(), first=(i == 0))
        doc.add_page_break()

        # ════════════════════════ 목차 ════════════════════════
        vspace(44)
        ct = doc.add_paragraph(); ct.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(ct.add_run(track_text("CONTENTS", 4)), 17, serif=True, bold=True, color=_INK)
        ct.paragraph_format.space_after = Pt(5)
        cts = doc.add_paragraph(); cts.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(cts.add_run("차 례"), 10, serif=False, color=_FAINT, track=3)
        cts.paragraph_format.space_after = Pt(12)
        hairline(width_cm=1.6, color=_GOLD, space_after=34, size=10)

        # 목차 전용 색 (참고 디자인: 보라 박스 + 청록 점선)
        TOC_PURPLE = '5B2E8C'
        TOC_PURPLE_RGB = (91, 46, 140)
        TOC_TEAL = (138, 170, 160)

        # 분야 라벨(있으면) 추출용
        field_label = ''
        if interview_data:
            field_label = (interview_data.get('field', '') or '').strip()

        # 챕터별 시작 페이지 추정 (앞 구조 페이지 + 누적)
        import math as _math
        _CPP = 330
        page_cursor = 9  # 표지·판권·프롤로그·목차 등 앞부분 근사
        chapter_start_pages = {}
        sub_pages_map = {}
        for cidx, chap in enumerate(outline):
            chapter_start_pages[cidx] = page_cursor
            page_cursor += 1  # 챕터 오프너
            if chap in chapters_data:
                for sname in chapters_data[chap].get('subtopics', []):
                    sub_pages_map[(cidx, sname)] = page_cursor
                    c = chapters_data[chap].get('subtopic_data', {}).get(sname, {}).get('content', '')
                    chars = len((c or '').replace(' ', '').replace('\n', ''))
                    page_cursor += max(1, _math.ceil(chars / _CPP))

        def _toc_part_chip(idx, clean):
            """보라 PART 박스 + 제목을 한 줄 표로 (참고 디자인)."""
            t = doc.add_table(rows=1, cols=2)
            t.alignment = WD_TABLE_ALIGNMENT.LEFT
            t.autofit = False
            cbox, ctitle = t.rows[0].cells
            cbox.width = Cm(1.7); ctitle.width = Cm(9.1)
            # 보라 박스
            _cell_shade(cbox, TOC_PURPLE)
            _no_cell_borders(cbox)
            bp1 = cbox.paragraphs[0]; bp1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            bp1.paragraph_format.space_before = Pt(5); bp1.paragraph_format.space_after = Pt(0); bp1.paragraph_format.line_spacing = 1.0
            set_font(bp1.add_run(f"PART {idx+1}"), 8, serif=False, bold=True, color=(255, 255, 255), track=0.3)
            if field_label:
                bp2 = cbox.add_paragraph(); bp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                bp2.paragraph_format.space_before = Pt(1); bp2.paragraph_format.space_after = Pt(5); bp2.paragraph_format.line_spacing = 1.0
                set_font(bp2.add_run(field_label[:6]), 10, serif=False, bold=True, color=(255, 255, 255), track=1)
            else:
                bp1.paragraph_format.space_after = Pt(5)
            va = OxmlElement('w:vAlign'); va.set(qn('w:val'), 'center'); cbox._tc.get_or_add_tcPr().append(va)
            # 제목 셀
            _no_cell_borders(ctitle)
            tp = ctitle.paragraphs[0]
            tp.paragraph_format.left_indent = Cm(0.35)
            tp.paragraph_format.space_before = Pt(2); tp.paragraph_format.space_after = Pt(0); tp.paragraph_format.line_spacing = 1.2
            add_hyperlink(tp, clean, f"chapter_{idx+1}", size=15, bold=True, color=_INK, serif=True)
            va2 = OxmlElement('w:vAlign'); va2.set(qn('w:val'), 'center'); ctitle._tc.get_or_add_tcPr().append(va2)

        for idx, chapter in enumerate(outline):
            clean = chapter
            for pre in [f"PART {idx+1}.", f"PART{idx+1}.", f"PART {idx+1} ", f"{idx+1}.", f"{idx+1})"]:
                clean = clean.replace(pre, "").strip()
            if idx > 0:
                vspace(20)
            # 청록 점선 (상단)
            hairline(align=WD_ALIGN_PARAGRAPH.LEFT, color=TOC_TEAL, space_before=0, space_after=8, size=6)
            # 보라 PART 박스 + 제목
            _toc_part_chip(idx, clean)
            # 청록 점선 (하단)
            hairline(align=WD_ALIGN_PARAGRAPH.LEFT, color=TOC_TEAL, space_before=8, space_after=14, size=6)
            # 소제목 + 점선 리더 + 페이지 번호
            if chapter in chapters_data:
                for sub in chapters_data[chapter].get('subtopics', []):
                    si = chapters_data[chapter]['subtopics'].index(sub)
                    pg = sub_pages_map.get((idx, sub), '')
                    sr = doc.add_paragraph()
                    sr.paragraph_format.left_indent = Cm(0.3); sr.paragraph_format.space_after = Pt(9)
                    sr.paragraph_format.line_spacing = 1.3
                    sr.paragraph_format.tab_stops.add_tab_stop(
                        Cm(10.6), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
                    add_hyperlink(sr, sub, f"subtopic_{idx+1}_{si+1}", size=10.5, bold=False, color=(58, 58, 64), serif=True)
                    # 탭 + 페이지번호
                    tabrun = sr.add_run(f"\t{pg:03d}" if isinstance(pg, int) else "")
                    set_font(tabrun, 10, serif=False, color=TOC_PURPLE_RGB)
        doc.add_page_break()

        # ════════════════════════ 본문 (장마다 섹션) ════════════════════════
        for idx, chapter in enumerate(outline):
            clean = chapter
            for pre in [f"PART {idx+1}.", f"PART{idx+1}.", f"PART {idx+1} ", f"{idx+1}.", f"{idx+1})"]:
                clean = clean.replace(pre, "").strip()
            sec = doc.add_section(WD_SECTION.NEW_PAGE)
            _apply_section(sec)
            sec.different_first_page_header_footer = True
            running_header(sec, clean)
            page_number_footer(sec)

            vspace(90)
            lab = doc.add_paragraph(); lab.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(lab.add_run(track_text("PART", 3)), 10, serif=False, color=_GOLD)
            lab.paragraph_format.space_after = Pt(4)
            num = doc.add_paragraph(); num.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(num.add_run(f"{idx+1:02d}"), 72, serif=True, bold=False, color=(60, 56, 50))
            num.paragraph_format.space_after = Pt(8)
            hairline(width_cm=1.8, color=_GOLD, space_after=24, size=10)
            cn = doc.add_paragraph(); cn.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_bookmark(cn, f"chapter_{idx+1}")
            set_font(cn.add_run(clean), 20, bold=True, color=_INK, serif=True)
            cn.paragraph_format.line_spacing = 1.3; cn.paragraph_format.space_after = Pt(44)

            if chapter not in chapters_data:
                continue
            subs = chapters_data[chapter].get('subtopics', [])
            for si, sub in enumerate(subs):
                content = chapters_data[chapter].get('subtopic_data', {}).get(sub, {}).get('content', '')
                if not content:
                    continue
                if si > 0:
                    doc.add_page_break()
                # 소제목 번호 칩 (골드 배경 박스)
                chip = doc.add_paragraph()
                chip.paragraph_format.space_before = Pt(10); chip.paragraph_format.space_after = Pt(6)
                _shade_para(chip, '2E2A24')
                cr = chip.add_run(f"  {idx+1}.{si+1}  ")
                set_font(cr, 11, bold=True, serif=False, color=(201, 162, 75), track=1)
                # 소제목 제목 (크고 굵게)
                ht = doc.add_paragraph()
                add_bookmark(ht, f"subtopic_{idx+1}_{si+1}")
                set_font(ht.add_run(sub), 16, bold=True, color=_INK, serif=True)
                ht.paragraph_format.space_after = Pt(6); ht.paragraph_format.line_spacing = 1.25
                # 진한 더블 밑줄(챕터 오프너와 확실히 구분)
                hairline(align=WD_ALIGN_PARAGRAPH.LEFT, color='2E2A24', space_before=0, space_after=2, size=10)
                hairline(align=WD_ALIGN_PARAGRAPH.LEFT, color=_GOLD, space_before=0, space_after=20, size=4)

                cleaned = clean_content(content, subtopic=sub)
                first_para_done = False
                for btype, bcontent in process_content_with_tables(cleaned):
                    if btype == 'table':
                        td = parse_table_data(bcontent)
                        if td and len(td) >= 2:
                            sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(8)
                            add_premium_table(td)
                            first_para_done = True
                    else:
                        for pt in [p for p in bcontent.split('\n\n') if p.strip()]:
                            # 챕터 첫 소제목의 첫 문단만 드롭캡, 이후엔 일반
                            if (not first_para_done) and si == 0 and len(pt.strip()) > 1:
                                drop_cap(pt.strip()[0], pt.strip()[1:])
                            else:
                                body_paragraph(pt.strip(), first=(not first_para_done))
                            first_para_done = True
                if si < len(subs) - 1:
                    ornament()

            # 챕터 끝 핵심 정리 박스
            ch_summary = chapters_data[chapter].get('key_points', [])
            if ch_summary:
                key_summary_box(ch_summary, idx + 1)

        # ════════════════════════ 에필로그 ════════════════════════
        ep_sec = doc.add_section(WD_SECTION.NEW_PAGE)
        _apply_section(ep_sec)
        ep_sec.different_first_page_header_footer = True
        page_number_footer(ep_sec)
        vspace(36)
        el = doc.add_paragraph(); el.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(el.add_run(track_text("EPILOGUE", 2)), 10, serif=False, color=_FAINT)
        el.paragraph_format.space_after = Pt(8)
        hairline(width_cm=1.4, color=_GOLD, space_after=30, size=8)
        epilogue_text = None
        if interview_data:
            epilogue_prompt = f"""당신은 한국 자기계발 베스트셀러 작가입니다. 마지막 페이지를 덮은 독자가 한 번 더 처음으로 돌아가게 만드는 에필로그를 작성하세요.

[저자 정보 - 참고용, 그대로 복사하지 말 것]
- 분야: {interview_data.get('field', '')}
- 경력 기간: {interview_data.get('experience_years', '')}
- 책 주제: {interview_data.get('topic', '')}
- 타겟 독자: {interview_data.get('target_reader', '')}
- 독자에게 전하고 싶은 말: {interview_data.get('final_message', '')}
- 작가 경력/경험: {interview_data.get('author_career', '')}
- 어려움/실패 경험: {interview_data.get('struggle_story', '')}
- 극복 스토리: {interview_data.get('breakthrough', '')}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

## 에필로그 작성 원칙 (스토리텔링)

### 1. 나의 스토리로 시작 (3-4문장)
- 작가 경력/경험을 자연스럽게 녹여서
- "저는 ~했습니다" 형식으로 간결하게
- 구체적 숫자나 사실 포함

### 2. 왜 이 책을 썼는지 (2-3문장)
- 내가 겪은 어려움 + 극복 과정 힌트
- 독자를 위해 책을 쓴 진심

### 3. 독자에게 한마디 (2-3문장)
- 지금 당장 할 수 있는 구체적 행동 하나
- 진심 어린 마무리 (근데 뻔하지 않게)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[문체]
- 짧은 문장, 짧은 문단 (2-4문장)
- 구어체 + 합쇼체 ("~거든요", "~잖아요" OK)

[분량] 400-500자

[금지 - 절대 쓰지 말 것]
- 저자 정보를 그대로 복사 붙여넣기
- 자청, 역행자, 자의식 해체, 유전자 역행, 원시인, 추월차선 (특정 작가 고유 표현)
- 교과서 표현: "여정", "발걸음", "함께 하시길 바랍니다", "진정한"
- AI 표현: "~의 중요성", "다양한", "효과적인", "~를 통해"
- 과장: "놀라운", "혁신적인", "충격적인"
- 뻔한 말: "포기하지 마세요", "꾸준히 하세요", "화이팅"
- 직접 호칭: 여러분, 당신, 독자님
- 마크다운 문법

에필로그만 출력하세요."""
            gen = ask_ai(epilogue_prompt, 0.7, ensure_quality=True)
            if gen:
                epilogue_text = gen
        if not epilogue_text:
            epilogue_text = ("여기까지 읽어주셔서 고맙습니다.\n\n완벽하지 않아도 괜찮습니다. 오늘 할 수 "
                             "있는 한 가지만 시작해 보세요. 작은 시작이 가장 멀리 갑니다.")
        epilogue_text = re.sub(r'^\s*#+\s*(에필로그|Epilogue|마치며|맺음말)\s*\.?\s*$', '', epilogue_text, flags=re.MULTILINE | re.IGNORECASE)
        epilogue_text = epilogue_text.lstrip('\n').strip()
        for i, t in enumerate([x for x in epilogue_text.split('\n\n') if x.strip()]):
            body_paragraph(t.strip(), first=(i == 0))
        vspace(30)
        sg = doc.add_paragraph(); sg.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_font(sg.add_run(author or "저자"), 11, italic=True, color=_SOFT, serif=True)

        # ════════════════════════ 저자 소개 ════════════════════════
        doc.add_page_break()
        vspace(96)
        al = doc.add_paragraph(); al.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(al.add_run(track_text("ABOUT THE AUTHOR", 1)), 9, serif=False, color=_FAINT)
        al.paragraph_format.space_after = Pt(10)
        hairline(width_cm=1.4, color=_GOLD, space_after=22, size=8)
        an = doc.add_paragraph(); an.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(an.add_run(author or "저자"), 16, bold=True, color=_INK, serif=True)
        an.paragraph_format.space_after = Pt(18)
        if interview_data:
            field = interview_data.get('field', ''); exp = interview_data.get('experience_years', '')
            career = interview_data.get('author_career', ''); method = interview_data.get('core_method', '')
            if career:
                author_bio = f"{field} 분야에서 {exp} 활동해온 실전가.\n\n{career}\n\n{method[:100] if method else ''}"
            else:
                author_bio = f"{field} 분야에서 {exp} 활동해온 실전가.\n\n{method}"
        else:
            author_bio = "실전에서 직접 부딪히며 쌓은 노하우를 독자와 나누고자 이 책을 썼다."
        for t in [x for x in author_bio.split('\n\n') if x.strip()]:
            bp = doc.add_paragraph(); bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(bp.add_run(t.strip()), 10, color=_SOFT, serif=True)
            bp.paragraph_format.line_spacing = 1.7; bp.paragraph_format.space_after = Pt(10)
        vspace(40)
        em = doc.add_paragraph(); em.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(em.add_run("✦"), 11, color=_FAINT, serif=False)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue(), None

    except Exception as e:
        return None, f"문서 생성 오류: {str(e)}"
# ==========================================
# 기법(메서드) 이름 — 영문 통일 + 책마다 고유
# ==========================================
def load_used_method_names():
    """이전 전자책들에서 쓴 기법 이름 목록 (겹침 방지용)"""
    names = load_config().get('used_method_names', [])
    return names if isinstance(names, list) else []


def remember_method_name(name):
    if not name:
        return
    names = load_used_method_names()
    if name not in names:
        names.append(name)
        save_config({'used_method_names': names[-200:]})  # 최근 200개만 보관


def _fallback_method_name(topic):
    import random
    # 발음 가능한 2~3자 영문 약자 + 자연스러운 한글 접미사 (예: "ERP 공식")
    acronym = "".join(random.sample("ABCDEFGHJKLMNPRSTVWXYZ", k=random.choice([2, 3])))
    suffix = random.choice(["공식", "법칙", "시스템", "전략", "구조", "사이클"])
    return f"{acronym} {suffix}"


def get_or_create_method_name(topic, interview_data=None, force_new=False):
    """이 책의 기법 이름을 만들거나 가져온다.

    형식: '영문 약자(2~4자) + 한글 접미사'  예) "ERP 공식", "DPS 법칙"
    - force_new=True: 새 전자책 → 이전 책들과 겹치지 않는 새 이름
    - force_new=False: 같은 책 안의 재생성 → 이미 정해진 이름 재사용
    """
    if not force_new and st.session_state.get('method_name'):
        return st.session_state['method_name']

    used = load_used_method_names()
    used_str = ", ".join(used[-60:]) if used else "(아직 없음)"
    seed = uuid.uuid4().hex[:6]
    core = (interview_data or {}).get('core_method', '') if interview_data else ''

    prompt = f"""이 전자책의 '기법 이름' 딱 하나를 짓는다.

주제: {topic}
핵심 방법: {core}

형식 (반드시 지킬 것):
- "영문 대문자 약자(2~4자) + 한글 접미사" 한 덩어리.  예) "ERP 공식", "DPS 법칙", "ARC 전략", "PMR 시스템"
- 약자는 발음 가능하고, 각 글자가 주제와 관련된 영어 단어의 첫 글자여야 한다(의미 있는 약자).
  예) ERP = Earn-Reinvest-Profit, DPS = Discover-Plug-Scale
- 한글 접미사는 다음 중에서만: 공식 / 법칙 / 시스템 / 전략 / 구조 / 사이클
- 약자만 영어이고, 영어 '단어'를 길게 늘어놓지 말 것.
- 약자는 반드시 영문 '대문자'로만(예: ERP O, erp/Peak X). 일반 영어 단어(peak, loop, flow 등)를 약자 자리에 쓰지 말 것.

❌ 절대 금지 (어색한 예):
- "eBook Revenue Loop", "Cashflow Ladder" 같은 영어 다단어 구절 (부자연스러움)
- 한글 약자, 의미 없는 자음 나열, 발음 불가능한 약자
- 아래 이미 쓴 이름과 약자·접미사가 겹치는 것:
  {used_str}

다양성 시드 {seed} 참고해 매번 다르게.

JSON만 출력:
{{"acronym": "영문 대문자 2~4자", "expansion": "약자 풀이(영문, 하이픈 연결)", "suffix": "한글 접미사", "method_name": "약자 + 공백 + 접미사"}}"""

    name = None
    expansion = ""
    try:
        res = ask_ai(prompt, 0.9)
        parsed = parse_json(res) if res else None
        if parsed:
            ac = (parsed.get('acronym') or '').strip().upper()
            sf = (parsed.get('suffix') or '').strip()
            mn = (parsed.get('method_name') or '').strip()
            expansion = (parsed.get('expansion') or '').strip()
            # method_name 우선, 없으면 약자+접미사 조합
            name = mn if mn else (f"{ac} {sf}" if ac and sf else '')
    except Exception:
        name = None

    # 형식 검증: 영문 대문자 2~4자 + 공백 + 허용된 한글 접미사 ('루틴' 제외 — 어색)
    allowed_suffixes = ("공식", "법칙", "시스템", "전략", "구조", "사이클")
    suffix_re = "(" + "|".join(allowed_suffixes) + ")"
    pattern = r"^[A-Z]{2,4}\s" + suffix_re + r"$"

    def _normalize(n):
        """'PMR  공식' 같은 입력을 정규화하고, 약자를 대문자로 통일."""
        n = (n or "").strip()
        parts = n.split()
        if len(parts) == 2:
            ac, sf = parts[0], parts[1]
            # 약자가 알파벳으로만 이뤄졌으면 대문자화 (peak → PEAK 시도)
            if ac.isalpha() and ac.isascii():
                ac = ac.upper()
            return f"{ac} {sf}"
        return n

    def _valid(n):
        n = (n or "").strip()
        if not n or n in used:
            return False
        return re.fullmatch(pattern, n) is not None

    name = _normalize(name)
    if not _valid(name):
        name = _fallback_method_name(topic)
        expansion = ""  # 폴백 이름은 의미 있는 약자 풀이가 없음
        tries = 0
        while name in used and tries < 15:
            name = _fallback_method_name(topic)
            tries += 1

    st.session_state['method_name'] = name
    st.session_state['method_expansion'] = expansion
    remember_method_name(name)
    return name


def method_lock_rule(name):
    """모든 생성 단계에서 동일하게 끼워 넣는 '기법 이름 고정' 규칙 블록"""
    return f"""━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🔒 기법 이름 고정 (이 프롬프트의 다른 어떤 작명/금지 규칙보다 우선)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
이 책의 기법 이름은 정확히 "{name}" 으로 이미 확정되어 있다. (형식: 영문 약자 + 한글 접미사)
- 모든 곳에서 "{name}" 을 글자 그대로, 띄어쓰기까지 똑같이 사용한다.
- 약자를 영어 단어로 풀어쓰거나, 다른 약자/이름으로 바꾸거나, 접미사를 바꾸는 것을 절대 금지한다.
- ○○○ / [컨셉명] / [시스템명] 자리에는 전부 "{name}" 을 넣는다.
- "{name}" 에 또 다른 한글 접미사(시스템/공식 등)를 덧붙이지 않는다. 위 이름 그대로만 쓴다.
"""


def review_outline_grammar(outline_text, method_name=""):
    """생성된 목차를 '어법 전담'으로 한 번 더 검토·교정한다.
    자극/호기심은 유지하되, 어색한 표현만 자연스럽게 고친 목차를 반환.
    실패 시 원본을 그대로 반환(안전)."""
    if not outline_text or not outline_text.strip():
        return outline_text

    lock = f'기법 이름 "{method_name}" 은 글자 그대로 유지(절대 변경 금지).' if method_name else ""

    prompt = f"""너는 30년 경력의 한국어 교정 전문 편집자다. 아래는 자기계발서 목차다.
너의 임무는 '어색한 줄만 콕 집어 고치는 것'과 '같은 표현 반복을 없애는 것' 딱 두 가지다. 목차의 의미·자극·호기심·구성·강도는 절대 건드리지 마라.
⚠️ 가장 흔한 실수: 멀쩡한 줄까지 손대서 밋밋하게 만드는 것. 이미 자연스러운 줄은 단 한 글자도 바꾸지 마라(고칠 게 없으면 그대로 둔다). 강한 표현을 순화하거나 자극을 줄이면 교정 실패다.
{lock}

[고쳐야 할 어색함 — 반드시 잡아낼 것]
1. 비유 오용: 비유 명사를 '되다/안 되다'의 대상으로 쓴 것
   ❌ "배당주를 사도 월세가 안 되는 이유" → ✅ "배당주를 사도 월 100만원조차 못 받는 이유"
2. 사물에 안 맞는 동사: 추상/사물에 물리·생물 동사를 억지로 붙인 것
   ❌ "캘린더에 박은 습관" → ✅ "캘린더에 기재한 습관"
   ❌ "통장에 꽂히는 30만원" → ✅ "매달 통장에 들어오는 30만원"
   ❌ "계좌가 자란다 / 살아났다" → ✅ "잔고가 늘어난다 / 수익이 회복됐다"
   ❌ "재투자를 켠 사람" → ✅ "수익을 다시 넣은 사람"
3. 대비 짝 어긋남: "A는 ~, B는 ~"에서 양쪽 서술어 축이 다른 것
   ❌ "누구는 월세처럼 받고 누구는 한 푼도 못 쓴" → ✅ "누구는 월세처럼 받고 누구는 한 푼도 못 버는"
4. 주어-서술어 비호응, 억지 은유, 뜻이 통하지 않는 단어 조합
5. 원어민이 읽다가 "이게 무슨 말이지?" 하고 멈칫하는 모든 줄

[반복 제거 — 어법 교정만큼 중요]
- 목차 전체(25줄)를 훑어 '같은 핵심 단어·표현·문장 끝 패턴'이 두 번 이상 나오면, 의미·자극을 유지한 채 한쪽만 다른 표현으로 바꿔 중복을 없앤다.
  · 예: '결정적'이 두 줄에 있으면 한쪽을 '단 하나의 / 의외의' 등으로. '~하는 이유'가 두 줄이면 한쪽을 '~한 순간 / ~의 차이' 등으로. '무너지다'가 두 줄이면 한쪽을 '주저앉다 / 흔들리다'로.
  · 조사·접속사 같은 기능어와 컨셉명("{method_name}")·'○○○'는 반복으로 치지 않는다.
- 같은 숫자(예: 3개 + 3개월 + 3년)가 여러 번이면 일부를 다른 숫자·표현(열에 아홉, 대부분 등)으로 바꿔 겹침을 푼다.

[교정 원칙]
- 어법이 깨진 줄(비유 오용·사물에 안 맞는 동사·주어서술어 비호응·뜻 모를 조합)만 같은 의도·같은 자극을 유지한 채 자연스럽게 고친다.
- 자연스럽고 강한 줄은 손대지 마라. 의미를 바꾸거나 밋밋하게 만들면 실패다.
- 5개 PART 제목이 이어 읽혔을 때의 이야기 흐름(바닥→도약)과 순서를 절대 깨지 마라.
- 형식(PART 줄, 소제목 줄, 줄 수, 순서)은 입력과 똑같이 유지한다.
- 설명·주석·코드블록 없이, 교정된 목차 텍스트만 그대로 출력한다.

[교정할 목차]
{outline_text}"""

    try:
        fixed = ask_ai(prompt, 0.3)
        if fixed and fixed.strip() and len(fixed.strip()) > len(outline_text.strip()) * 0.5:
            return fixed.strip()
    except Exception:
        pass
    return outline_text


def generate_outline_only(interview_data, progress_placeholder):
    """인터뷰 데이터를 기반으로 목차까지만 생성 (본문 제외)"""
    try:
        topic = interview_data.get('topic', '')
        if not topic:
            return False

        # 1. 타겟 자동 설정
        progress_placeholder.info("🎯 1/4 타겟 독자 분석 중...")
        target = f"{interview_data.get('target_reader', '')} - {interview_data.get('target_problem', '')}"
        st.session_state['target_persona'] = target

        # 새 전자책 → 이전 책들과 겹치지 않는 영문 기법 이름을 새로 확정
        method_name = get_or_create_method_name(topic, interview_data, force_new=True)
        method_expansion = st.session_state.get('method_expansion', '')

        # 2. 책 고유 컨셉 생성 (가장 중요!)
        progress_placeholder.info("💡 2/4 책 고유 컨셉 설계 중...")
        concept_prompt = f"""당신은 크몽/클래스101 베스트셀러 전자책 기획자입니다.
이 책의 기법 이름은 이미 "{method_name}" (영문 약자 + 한글 접미사 형식)으로 확정되어 있습니다.
당신의 일은 새 이름을 짓는 게 아니라, 이 기법을 중심으로 책의 관점과 메시지를 설계하는 것입니다.

{method_lock_rule(method_name)}

이 기법의 약자 풀이(참고): {method_expansion if method_expansion else "(풀이 없음)"}
약자가 의미 있는 머리글자라면, 핵심 관점·메시지에 그 단계적 흐름이 자연스럽게 묻어나게 하라. 단 "Y:~" 식 사전 나열은 금지.

[저자 정보]
주제: {topic}
핵심 방법: {interview_data.get('core_method', '')}
저자만의 차별점: {interview_data.get('unique_point', '')}
타겟의 고민: {interview_data.get('target_problem', '')}
저자의 구체적 노하우·실전 디테일(★ 책의 핵심 재료 — 컨셉에 반드시 반영): {interview_data.get('detailed_knowhow', '') or '(미입력)'}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
✅ 만들어야 할 것
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

1. 핵심 관점 — 이 주제를 "{method_name}"(이)라는 기법으로 어떻게 새롭게 보는가? 남들과 다른 접근법.
2. 핵심 메시지 — "{method_name}(으)로 ~할 수 있다" 형식의 한 문장.

❌ 절대 금지:
- 기법 이름을 새로 짓거나 한글로 바꾸는 것 (이미 "{method_name}"으로 고정)
- "제국을 건설/왕좌/전설의/세계 최초" 같은 과장
- 황금·보물·비밀·마법·연금술 같은 유치한 단어

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📝 출력 형식
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[고유 시스템/공식 이름]
{method_name}

[핵심 관점]
(2~3문장, 자연스럽게)

[핵심 메시지]
(한 문장, "{method_name}(으)로 ~하는 방법")

[목차에서 반복할 키워드]
{method_name}"""

        book_concept = ask_ai(concept_prompt, 0.8, ensure_quality=True)
        st.session_state['book_concept'] = book_concept

        # 3. 제목 생성
        progress_placeholder.info("📝 3/4 제목 생성 중...")
        title_prompt = f"""당신은 크몽/클래스101 베스트셀러 전자책 제목을 만드는 전문가입니다.
결제 버튼을 누르게 만드는 제목을 써주세요.

[이 책의 컨셉]
{book_concept}

이 책의 기법 이름은 "{method_name}"으로 확정됨. 제목에 기법 이름을 넣을 경우 반드시 이 이름 그대로 사용(약자 풀어쓰기·재작명 금지).

[주제]
{topic}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🔥 실제 잘 팔리는 전자책 제목 분석
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[컨셉 중심형 - 짧은 신조어 + 부제]
• 단단한 돈 - 잃지 않는 사람의 7가지 원칙
• 돈의 속성 - 최소한 이것만은 알아야 할
• 1억 모으는 통장 - 30대 직장인의 5단계 공식

[신사임당/클래스101 스타일 - 결과 중심]
• 퇴사 후 월 1000만원 버는 글쓰기
• 블로그로 월 300 만드는 현실적인 방법
• 투잡러의 시간관리 비법

[크몽 베스트셀러 - 구체적 약속]
• 30일 만에 첫 수익 내는 스마트스토어
• 3개월 안에 월 100 만드는 전자책 공식
• 회사 다니면서 월 200 추가 수입 만들기

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
✅ 좋은 제목의 공식
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[공식 1] 컨셉형 (2~4단어)
저자만의 프레임워크/용어가 들어간 제목
예: 짧고 강렬한 신조어/약어 (3~5자), 동사의 명사화, 은유적 압축어

[공식 2] 결과형 (구체적 숫자 포함)
기간 + 결과가 명확한 제목
예: "3개월 만에 월 300", "100일 글쓰기"

[공식 3] 타겟형 (누구를 위한)
특정 대상의 고민을 건드리는 제목
예: "퇴사 준비생의 월급 독립기", "직장인의 두 번째 월급"

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🚫 절대 금지
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

❌ 유치한 비유:
나침반, 지도, 열쇠, 보물, 황금, 마법, 연금술

❌ AI스러운 제목:
"~의 이해", "~가이드", "~완벽 정복"
"효과적인 ~", "성공적인 ~"

❌ 너무 추상적:
의미를 알 수 없는 신조어
무슨 내용인지 전혀 감이 안 오는 제목

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📝 출력
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

제목: 컨셉이 드러나면서도 무슨 책인지 알 수 있게
부제: 구체적인 결과/약속을 담아 15~25자

JSON만 출력:
{{
    "title": "제목 (컨셉+내용이 드러나게, 3~8단어)",
    "subtitle": "부제 (구체적 결과/약속, 15~25자)"
}}"""

        title_result = ask_ai(title_prompt, 0.4)
        title_data = parse_json(title_result)
        if title_data:
            st.session_state['book_title'] = title_data.get('title', topic)
            st.session_state['subtitle'] = title_data.get('subtitle', '')

        # 4. 목차 생성 (책 컨셉 기반)
        progress_placeholder.info("📋 4/4 목차 설계 중...")
        outline_prompt = f"""당신은 한국 자기계발 분야 톱 0.1% 기획자입니다. 서점에서 단 5초간 목차만 본 사람이 책을 손에서 못 놓게 만드는 5장짜리 목차를 씁니다.

목차의 단 하나의 목적: 독자가 "이 책을 안 읽으면 평생 손해"라고 느끼게 만드는 것.
정보 전달은 본문이 한다. 목차는 100% 구매심리만 다룬다.

[이 책의 시그니처 컨셉/시스템]
{book_concept}

{method_lock_rule(method_name)}

[주제]: {topic}

[저자의 구체적 노하우·실전 디테일 (★ 목차에 자연스럽게 녹여라)]
{interview_data.get('detailed_knowhow', '') or '(미입력 — 일반 원리로 구성)'}
→ 위 노하우에 담긴 구체적 수치·단계·실수·기준·도구를 소제목에 녹여, 저자만의 디테일이 느껴지게 하라. 단 노하우를 그대로 베끼지 말고 호기심을 유발하는 제목으로 재가공한다.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🛒 구매 결정 5초 룰 (모든 규칙 중 1순위)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

독자는 목차를 5초만 본다. 그 5초 안에 다음 3가지가 동시에 작동해야 결제한다.

[1] 정체성 변화 발견 — "이걸 읽으면 나는 OO한 사람이 된다"
   → 5개 챕터 제목을 이어 읽으면 한 사람의 결정적 변화가 보여야 한다.
   ✅ "단단해진 멘탈은 인생을 통째로 바꾼다" (변화 서사 O)
   ❌ "멘탈 관리의 다양한 기법" (정체성 변화 X — 즉시 폐기)

[2] 손실회피 작동 — "이걸 모르면 평생 OO한다"
   → 통념 박살(인지부조화) 챕터/소제목이 최소 3개 들어가야 한다.
   ✅ "의지로 버틴 사람일수록 더 크게 무너진다"
   ❌ "멘탈 관리의 중요성" (잃을 게 안 보임)

[3] 구체성 — 추상 명사 1개당 구체적 숫자/장면 1개
   → "많은 사람" 금지, "월급 280만원짜리 7년차 회사원" 가능
   → 시간(90초, 47일), 금액(34만원, 1억), 비율(99%, 8할) 적극 사용

25줄(챕터 5 + 소제목 20) 중 한 줄이라도 "그냥 정보"가 섞이면 그 목차는 평이해진다.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
💎 기법 이름 사용 (위 고정 규칙을 그대로 따른다)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

기법 이름은 위에서 "{method_name}"(영문 약자+한글 접미사)으로 이미 확정되었다.
- 새로 작명하지 말고, 약자를 풀어쓰지 말고, 글자 그대로 쓴다.
- 아래 예시의 ○○○ 자리에는 전부 "{method_name}" 을 넣는다.
- "{method_name}" 에 또 다른 접미사를 덧붙이지 않는다.

[약자의 의미를 목차에 '간접적으로' 녹여라]
이 기법의 약자 풀이: {method_expansion if method_expansion else "(풀이 없음 — 약자가 무엇의 머리글자인지 책 흐름으로 자연스럽게 암시)"}
- 약자가 의미 있는 머리글자(예: ERP = Earn-Reinvest-Profit)라면, 그 각 단계가 PART 1~5의 흐름이나 소제목에 '자연스럽게' 드러나도록 목차를 설계하라. 독자가 목차만 읽어도 "아, 이 책이 이런 단계를 다루는구나"를 감 잡게.
- 단, "Y: ~, M: ~" 처럼 글자별로 사전식 나열은 절대 금지. 풀이를 직접 받아쓰지 말 것.
- 어디까지나 자연스러운 호기심 자극 제목 안에 의미가 스며들게 한다. 예) 약자 첫 단계가 'Earn(벌기)'이면 PART 1을 버는 단계의 통념 박살로 구성하는 식.
- 풀이가 없으면 억지로 짜맞추지 말고, 평소처럼 호기심 중심으로 쓰되 책 전체가 하나의 방법론을 단계적으로 다룬다는 느낌만 유지하라.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🏆 골드 스탠다드 — 이 '결'을 그대로 복제하라 (규칙 100개보다 이게 먼저다)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

아래 둘은 '완성된 목차'의 본보기다. 한국 정보성 전자책(자청·프드프 류) 베스트셀러의 목차 톤 — 설명하지 않고, 한 줄마다 '장면 + 의문'을 던져 손가락이 결제 버튼으로 가게 만든다. 규칙을 다 지키는 것보다, 아래 목차를 처음 본 독자가 느끼는 '안 사면 손해 같은 긴장감'을 똑같이 재현하는 게 1순위다. 네가 쓴 목차를 아래와 나란히 놓았을 때 결이 조금이라도 약하거나 '설명문' 같으면 전부 다시 써라.

[본보기 A — 30대 직장인 부업 · 월 500]
PART 1. 직장인 99%가 부업 첫 달을 못 넘기고 포기하는 이유
- 퇴근 후 3시간을 갈아 넣고도 통장이 그대로인 진짜 이유
- 부지런한 사람일수록 더 빨리 나가떨어지는 역설
- "일단 시작하라"는 그 말이 당신을 가난하게 붙잡아 둔다
- 결국 모든 답은 '○○○' 안에 있었다
PART 2. 노력보다 '자리'가 먼저다
- 같은 시간을 써도 누구는 월 30, 누구는 월 500인 단 하나의 차이
- 잘 파는 사람들이 절대 입 밖에 안 내는 자리 잡는 순서
- 카톡 프로필 한 줄만 바꿨더니 문의가 들어오기 시작했다
- 남들 다 뛰어든 시장에서 6개월 만에 따라잡히는 사람들의 공통점
PART 3. 첫 30만원이 통장에 찍힌 그날
- 정확히 47일째, 잔고가 처음 늘어난 그 순간
- 27만원짜리 첫 정산서가 알려준 단 한 가지
- ○○○가 통장에 처음 흔적을 남기는 네 가지 신호
- 첫 수익을 본 사람 열에 아홉이 거기서 멈추는 함정
PART 4. 자는 동안에도 매출이 들어오기 시작한다
- 일하지 않는 시간에도 돈이 들어오게 만드는 구조의 정체
- 월 100을 넘긴 사람들이 가장 먼저 끊은 습관 한 가지
- 일을 덜 했더니 오히려 매출이 두 배가 된 이상한 달
- 한 번 자리잡은 사람은 다시 월급만 기다리지 않는다
PART 5. 월 500 다음, 억대로 넘어가는 길
- 월 500에서 멈추는 사람과 5천을 넘기는 사람의 단 한 줄 차이
- 돈과 시간이 동시에 따라오기 시작하는 전환점
- 1년 만에 사직서를 낸 사람들의 똑같았던 마지막 한 달
- '○○○'로 자리잡은 사람들의 5년 뒤가 완전히 다른 이유

[본보기 B — 멘탈 · 자기관리]
PART 1. 1년 안에 99%가 다시 무너지는 결정적 이유
- 마음 다잡기를 백 번 해도 사흘이면 원점으로 돌아가는 이유
- 의지가 강한 사람일수록 더 크게 무너지는 역설
- "긍정적으로 생각하라"는 조언이 당신을 더 망가뜨린다
- 여기서부터 진짜 이야기, '○○○'가 시작된다
PART 2. 의지로 잡으려는 순간 뇌는 반대로 움직인다
- 감정이 폭발할 때 뇌가 정확히 어디부터 멈추는가
- 참을수록 더 크게 터지는 사람들의 공통된 패턴
- 명상도 운동도 당신에게만 효과가 없었던 진짜 이유
- 무너진 다음 날 가장 먼저 해야 할 단 한 가지
PART 3. 회복의 8할이 결판나는 폭발 직후 90초
- 이 90초를 놓치면 다시 일어서는 데 며칠이 걸린다
- 단번에 회복하는 사람들이 무의식적으로 하는 행동
- ○○○를 처음 적용한 사람들이 가장 먼저 느낀 변화
- 감정에 휘둘리던 사람이 처음으로 멀쩡했던 그날
PART 4. 한 번 회복한 뇌는 다시는 무너지지 않는다
- 같은 일을 겪어도 더는 흔들리지 않게 되는 메커니즘
- 무너지는 횟수가 한 달 만에 절반으로 줄어드는 신호
- 단단해진 사람들이 더 이상 하지 않는 다섯 가지 행동
- 주변에서 "사람이 달라졌다"는 말을 듣기 시작하는 시점
PART 5. 단단해진 멘탈은 인생을 통째로 바꾼다
- 멘탈이 바뀌자 통장과 인간관계가 함께 달라진 이유
- 다시는 예전으로 돌아가지 않는 사람들의 결정적 차이
- 6개월 뒤 가장 크게 달라지는 단 한 가지
- '○○○'로 회복한 사람들의 5년 뒤가 완전히 다른 이유

[본보기 C — 주식·배당 투자]
PART 1. 10년을 모아도 배당으로 월 50만원 못 받는 사람들의 공통점
- 배당률 숫자만 보고 골랐더니 3년째 원금이 그대로인 이유
- 매달 꼬박꼬박 넣었는데도 통장이 안 불어난 8할의 패턴
- "분산투자"가 오히려 당신을 제자리에 묶어두는 방식
- 결국 모든 답은 '○○○' 안에 있었다
PART 2. 받은 배당을 쓰지 않은 사람만 3년 뒤 게임이 달랐다
- 같은 원금을 넣었는데 누구는 매달 80만원, 누구는 한 푼도 못 받는 이유
- 배당을 재투자한 계좌와 생활비로 쓴 계좌의 5년 뒤 잔고
- 증권사 앱을 하루 세 번 이상 열어본 해에 수익률이 가장 낮았다
- 남들이 팔 때 오히려 더 사 모은 계좌만 폭락장에서 살아남은 이유
PART 3. 처음으로 배당이 생활비를 넘어선 그달
- 재투자를 3년 유지한 계좌에서 배당이 두 배로 늘어난 순간
- ○○○가 처음으로 통장에 흔적을 남기는 네 가지 장면
- 원금보다 배당이 먼저 눈에 들어오기 시작하는 전환의 신호
- 처음 이 구조를 이해한 사람 열에 여덟이 바로 그때 저지르는 실수
PART 4. 쓰지 않아도 매달 들어오는 돈이 생긴 사람의 뇌는 다르다
- 근로소득만 바라보던 사람이 더 이상 월급날을 기다리지 않게 되는 때
- 배당 수령액이 월 100만원을 넘긴 사람들이 가장 먼저 바꾼 한 가지
- 계좌를 덜 들여다볼수록 잔고가 더 빠르게 불어난 이상한 메커니즘
- 직장을 그만두지 않아도 퇴직 준비가 끝나는 정확한 시점
PART 5. 배당이 월급을 넘어선 날, 인생의 선택지가 달라졌다
- 노동 없이 들어오는 돈이 생기자 직장에서 더 당당해진 이유
- 5년 뒤 가장 크게 갈리는 건 투자 금액이 아니라 이것 하나
- 배당으로 퇴직한 사람들이 다시 월급쟁이로 돌아가지 않는 구조
- '○○○'를 3년 이상 유지한 사람들의 10년 뒤가 완전히 다른 이유

[본보기 D — 직장인 퇴사·독립]
PART 1. 퇴사를 꿈꾸다 40대가 된 직장인의 정확한 실수
- 퇴사 준비를 3년째 하고 있는데 통장이 그대로인 진짜 이유
- "조건이 갖춰지면 나간다"는 그 말이 당신을 평생 붙잡는다
- 먼저 퇴사한 동기가 6개월 만에 다시 돌아온 날 배운 것
- 여기서부터 진짜 이야기, '○○○'가 시작된다
PART 2. 월급이 없어도 버티는 사람과 석 달 만에 무너지는 사람
- 퇴사 첫 달보다 여섯 번째 달이 더 무서운 진짜 이유
- 돈이 아니라 이것 하나가 먼저 무너진 사람이 결국 돌아간다
- "나만의 일"을 찾느라 1년을 날린 사람들의 공통된 함정
- 퇴사 후 6개월 살아남은 사람들이 처음 한 달에 한 일
PART 3. 월 200이 통장에 처음 찍힌 날, 두려움이 사라졌다
- 정확히 73일째, 첫 입금 알림이 울린 그 새벽
- ○○○가 처음으로 작동한다는 걸 알아챈 단 하나의 신호
- 첫 수익을 본 사람이 바로 그 순간 가장 많이 저지르는 실수
- 이 시점을 넘긴 사람과 못 넘긴 사람의 1년 뒤 통장
PART 4. 다시는 월급날을 기다리지 않는 사람이 된다
- 수입이 불규칙해도 흔들리지 않는 뇌가 만들어지는 조건
- 사업자 등록증 한 장이 삶의 선택지를 통째로 바꾼 이유
- 퇴사 1년 차에 가장 먼저 끊은 것과 가장 먼저 시작한 것
- 주변에서 "달라 보인다"는 말을 듣기 시작하는 정확한 시점
PART 5. 월 500을 넘긴 뒤, 그다음 게임은 완전히 다르다
- 돈을 좇던 사람이 돈이 따라오게 만드는 사람이 된 차이
- 1년 만에 사직서를 쓴 사람들의 마지막 3개월이 닮은 이유
- 지금도 퇴사를 못 한 사람과 해낸 사람 사이의 단 하나
- '○○○'로 자립한 사람들이 5년 뒤 가장 잘했다고 꼽은 결정

[본보기에서 훔쳐야 할 6가지]
1. 첫 줄(PART 1)부터 독자의 뜨끔한 현실을 정면으로 때린다 — "직장인 99%가 못 넘기는", "1년 안에 99%가 무너지는".
2. 통념을 뒤집어 "내 탓이 아니었다"는 안도와 충격을 동시에 준다 — "부지런한 사람일수록 더 빨리 나가떨어진다".
3. 결과·장면은 보여주고 방법·이유는 숨긴다 — "47일째 잔고가 늘어난 그 순간"(왜·어떻게는 본문).
4. 5개 PART를 세로로 읽으면 '바닥→깨달음→첫 사건→전환→도약'이 한 편의 이야기로 흐른다.
5. 한 줄도 "~의 방법", "~의 중요성" 같은 교과서 표현이 없다. 전부 사람·장면·숫자·의문이다.
6. ★ "너 지금 이러고 있지?" 직접 저격 — 자청 목차의 가장 강한 무기. 독자가 읽다가 "어? 나 얘기인데?"라고 멈춰야 결제한다. 소제목 20개 중 최소 5개는 독자의 현재 행동을 정확하게 묘사해서 들키게 만들어라. "퇴사 준비를 3년째 하고 있는데 통장이 그대로인 진짜 이유", "공부만 하다 1년을 통째로 날리는 사람들의 5가지 변명", "증권사 앱을 하루 세 번 이상 열어본 해에 수익률이 낮았다" — 독자가 어제 한 행동을 정확히 집어내는 줄이 목차에서 가장 세게 작동한다.

⚠️ 위 A·B·C·D는 '톤·구조의 본보기'다. 문장을 그대로 베끼지 말고, 주어진 [주제]·[저자 노하우]에 맞춰 같은 강도로 완전히 새로 써라. A·B의 문장 패턴을 그대로 옮기면 즉시 탈락이다.

[실전 실패 사례 — 실제로 이렇게 평이하게 나왔다. 같은 실수를 반복하지 마라]
✗ 두 PART가 사실상 같은 말이라 흐름이 안 보인 경우:
   PART 3. 첫 개선 후 매출이 두 배로 뛴 그날
   PART 4. 개선할수록 매출이 따라 올라온다   ← 둘 다 '개선하면 매출 오른다' = 같은 내용. 5막이 무너짐.
   ✓ 역할을 완전히 갈라라 — PART 3은 '첫 성과가 터진 장면', PART 4는 '자동으로 팔리는 사람이 됨':
   PART 3. 사십칠 번째 주문이 찍힌 새벽의 정산 알림
   PART 4. 자는 사이에 휴대폰이 울리는 판매자가 된다

✗ "리뷰 3개가 알려준 결정적 신호" → '결정적 신호'는 금지된 AI 클리셰 + 무슨 신호인지 안 보여 밍밍.
   ✓ "후기 한 줄이 본문 백 페이지를 이기는 이유" (장면 + 의문)
✗ "SRP 사이클이 본격 작동하기 시작하는 4가지 신호" → 시스템 의인화('사이클이 작동') + '신호' 반복.
   ✓ "○○○가 통장에 첫 흔적을 남기는 네 가지 장면"
✗ "매출과 시간이 동시에 따라오기 시작하는 전환점" → 추상적이라 장면이 없음.
   ✓ "일 년 만에 사표를 던진 사람들의 닮은 마지막 한 주"

★ 위 실패작이 평이했던 결정적 이유: 같은 말이 계속 반복됐다 — '전자책' 다섯 번, '개선' 네 번, '매출' 세 번, '가장 먼저'·'결정적'·'신호'·'한 번' 각 두 번, 숫자 3이 네 번. 이 반복이 목차를 가장 평이하게 만든다. 25줄에서 핵심 단어와 문장 끝은 전부 한 번씩만 써라.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🧠 마케팅 뇌과학 8대 트리거 (목차 전체에 골고루 박을 것)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

뇌가 읽기를 멈출 수 없게 만드는 메커니즘. 뇌는 (1) 예측이 깨질 때(예측오류) (2) 빈칸을 보면 채우려 할 때(정보격차) (3) 잃을 것이 보일 때(손실회피) 도파민이 분비돼 "더 알고 싶다"는 충동을 만든다. 아래 트리거는 전부 이 셋을 노린다.

1. **손해회피 + 충격 통계** — "1년 안에 99%가 다시 무너진다", "1억치 강의 들어도 망한다"
2. **인지부조화/통념 박살(예측오류)** — "의지로 버틴 사람일수록 더 크게 무너진다", "노력할수록 가난해진다"
3. **권위 어휘 (과학/임상)** — 뇌, 신경회로, N주 후, 임상, 데이터, 알고리즘 (절대 비유로 남용 금지, 사실 진술로만)
4. **임박감 + 절대성** — "이 90초를 놓치면 며칠 걸린다", "다시는 ~하지 않는다"
5. **정체성 전환 약속** — "회복한 뇌는 다시 무너지지 않는다", "단단해진 사람은 ~한다"
6. **인그룹 사회증명** — "○○를 익힌 사람들의 5년 뒤", "상위 1%만 도달하는"
7. **이중/삼중 보상** — "통장과 인간관계가 함께 변한다", "돈도 사람도 따라온다"
8. **호기심 갭(정보격차)** — "두 달 안에 가장 먼저 끊은 한 가지", "정확히 어디부터 멈추는가"

※ 가장 강한 줄은 위 트리거를 '겹쳐' 쓴다. 예) "10명 중 8명이(사회증명) 배당주를 사도 월 100만원조차 못 받는(손실회피+구체금액) 이유(정보격차)" — 한 줄에 트리거 2~3개가 겹칠수록 결제 충동이 폭발한다.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🧲 호기심 갭(Information Gap) 강화 — 결제를 부르는 가장 강력한 무기
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

조지 로웬스타인의 정보격차 이론: 사람은 자기가 모르는 것이 '뭔지'는 알지만 '내용'은 모를 때 가장 강하게 끌린다. 25줄 중 최소 10줄에 이 갭을 넣어야 목차만 보고 결제한다.

[호기심 갭 6대 공식 - 결과만 보이고 방법/이유는 본문으로 미루기]

1. **숫자 + 미공개 결과** — '정확히/딱'으로 시작해 결과만 보여주고 내용은 숨기기
   ✅ "정확히 47일째, 통장 잔고가 처음으로 불어나기 시작한 순간"
   ✅ "27만원짜리 첫 정산서가 알려준 단 한 가지"
   ❌ "47일 만에 돈을 버는 방법" (방법을 다 보여줘버림 → 결제 안 함)

2. **이미 벌어진 사건 + 원인 숨김** — '왜?'를 유발하는 결과만
   ✅ "3년 차 베테랑이 신입에게 6개월 만에 따라잡힌 단 하나의 이유"
   ✅ "월 1,000을 찍은 사람들이 가장 먼저 끊은 습관 한 가지"

3. **묘하게 구체적인 행동/대상 + 이유 숨김** — 디테일이 호기심을 폭발시킨다
   ✅ "성공한 부업러가 매일 밤 11시에 반드시 끄는 것"
   ✅ "1년 만에 1억 모은 사람들이 절대 안 쓰는 5단어"
   ✅ "월 500 넘긴 사람들 카톡 프로필에서 사라진 한 단어"

4. **반대 결과 미스터리** — 통념과 정반대 결과만 던지고 메커니즘은 본문
   ✅ "더 열심히 할수록 더 가난해진 7년의 비밀"
   ✅ "잠을 늘렸더니 매출이 2배가 된 이상한 메커니즘"

5. **'딱 하나' 절대성** — 수많은 변수 중 단 하나만 보여주기
   ✅ "월 100 / 월 500을 가르는 단 한 줄의 차이"
   ✅ "결국 모든 게 무너지는 사람들의 공통점 단 하나"

6. **시간 압축 미스터리** — 짧은 시간에 큰 일이 일어났는데 그 사이를 숨기기
   ✅ "퇴근 후 90분이 1년 뒤 인생을 갈라놓는다"
   ✅ "주말 4시간이 5년치 월급을 바꾼 그 과정"

[호기심 갭 만들 때 절대 어기지 말 것]
• 답을 같은 줄에 다 보여주지 마라. "왜 ~한가" "어떻게 ~하는가"로 끝나면 본문을 사야 알 수 있게.
• "방법" "비법" "노하우" 같은 말로 끝내면 갭이 닫힘 → 결제 안 함.
• 실행 순서와 디테일(종목·비율·타이밍)은 숨기되, 작동 원리의 '큰 줄기'는 한 조각 드러내라 — 원리까지 다 숨기면 '근거 없는 주장'이 되어 안 팔린다(아래 ⚖️ 섹션 참고).
• 한 줄 안에 "장면 + 의문"이 같이 있어야 호기심이 작동한다.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
⚖️ 호기심 갭의 함정 — '주장만 있고 근거가 없으면' 안 팔린다 (★ 매우 중요)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

방금까지 '방법을 숨기라'고 했다. 하지만 모든 줄을 '결과 + 미스터리'로만 채우면 목차가 '근거 없는 주장 모음'이 되어 매력이 죽는다. 독자는 "오, 이렇게 하니까 되는 거구나"라는 믿음(=작동 원리 감)이 있어야 결제한다. 지금 가장 흔한 실패가 바로 이것 — "달라진다 / 완성된다 / 무너지지 않는다 / 쌓이기 시작한다" 같은 결과 선언만 스무 줄 반복해서 공허해지는 것.

[숨길 것 vs 보여줄 것 — 이 구분이 전부다]
• 숨겨라(본문에서) = 실행 디테일: 정확한 종목·비율·금액·타이밍·순서·도구 설정. "이대로 따라 하면 된다" 수준은 절대 노출 금지.
• 보여줘라(목차에서) = '지렛대 한 조각 + 그게 왜 되는지 원리'. 독자가 큰 흐름은 감 잡되, 혼자 실행하진 못하는 수준.

[나쁜 예 — 주장만, 근거 0 (실제로 이렇게 나왔다)]
✗ "배당금이 통장에 찍히는 순간부터 게임이 완전히 달라진다" → 뭐가 어떤 원리로 달라지는지 지렛대가 안 보임. 게다가 '게임이 달라진다'는 AI 클리셰.
✗ "단단해진 포트폴리오는 폭락장에서도 무너지지 않는다" → 왜 안 무너지는지 근거 0. 공허한 선언.
✗ "현금흐름과 자산 가치가 동시에 따라오기 시작하는 전환점" → 추상적 주장, 작동 원리 없음.
✗ "한 번 흐름을 만든 사람은 다시 원금만 바라보지 않는다" → 무슨 '흐름'인지 레버가 안 보임.

[좋은 예 — 지렛대·원리는 보여주고 디테일은 숨김]
✓ "받은 배당을 쓰지 않고 그대로 다시 사 모은 사람만, 같은 원금으로 더 많은 배당을 받게 된 이유" (지렛대=재투자, 원리=복리가 보임 / 얼마를·어느 종목은 숨김)
✓ "폭락장에 팔지 않고 오히려 더 사 모은 계좌만 이듬해 가장 빨리 회복한 이유" (행동=급락 매수, 인과의 큰 흐름 보임 / 비중·시점은 숨김)
✓ "월급은 일해야 들어오고 배당은 가만히 있어도 들어온다 — 이 둘을 갈라놓은 사람만 자산이 따로 쌓인다" (구조의 논리가 보임)
✓ "수익을 빼 쓰지 않고 3년 내내 재투자만 한 사람이 원금을 두 배로 불린 까닭" (지렛대=재투자 누적, 결과의 근거 보임)

[규칙 — 이 강도로 써라]
- 소제목 20개 중 최소 12개는 '구체적 지렛대(행동·원리) 한 조각'을 드러내라. 막연한 '달라진다 / 완성된다 / 무너지지 않는다 / 따라온다'로 끝나는 순수 주장 줄은 PART당 1개 이내로 제한.
- 20개를 다 읽으면 이 방법의 '큰 골격(무엇을 → 왜 → 대략 어떻게)'이 손에 잡혀야 한다. 단, 그 골격을 '혼자 실행할 수 있을 만큼' 구체적으로 주지는 마라(그건 본문의 몫).
- 컨셉명 약자(예: DIME)의 각 단계가 PART 1~5 흐름에 자연스럽게 드러나, 독자가 "이런 순서·논리로 가는 방법이구나"를 목차만 보고 감 잡게 하라.
- 즉, 호기심(디테일 숨김)과 신뢰(원리 노출)를 한 줄 안에 같이 넣는다 — "[구체적 지렛대]로 [결과]가 일어난 [까닭/순간]"이 가장 강한 공식이다.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🎯 단 하나의 미션 — 목차가 곧 '예고편 스토리'다
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

5개 PART 제목은 독립된 5개의 광고 문구가 아니다. 이어 읽으면 한 사람(독자 자신)의 추락→반전→도약을 그리는 '한 편의 예고편'이어야 한다.
목차를 5초 훑은 사람의 머릿속에 "어? 이거 딱 내 얘기인데, 그래서 마지막에 어떻게 됐는데?"라는 한 문장이 떠올라야 결제한다.

[★ 스토리 흐름 강제 규칙 — 가장 자주 어기는 부분, 반드시 지켜라]
PART 1~5를 '제목만' 세로로 이어 읽었을 때, 끊긴 5개 슬로건이 아니라 시간 순서대로 이어지는 한 사람의 이야기로 읽혀야 한다.

(1) 같은 주인공 — 5개 PART 내내 '같은 한 사람'의 시점이 유지돼야 한다. PART마다 화자가 바뀌면 흐름이 깨진다.
(2) 시간의 화살 — PART가 넘어갈수록 시점이 '과거→현재→미래'로 전진해야 한다. 뒤로 돌아가거나 제자리걸음 금지.
    PART 1 (지금, 바닥) → PART 2 (왜 그랬는지 깨달음) → PART 3 (처음 바뀐 그 사건) → PART 4 (이제 다른 사람이 됨) → PART 5 (1년 뒤, 완전히 다른 인생)
(3) 감정 에스컬레이션 — PART 1은 가장 아프게(좌절·손실), PART 5는 가장 벅차게(도약·보상). PART가 올라갈수록 잃을 공포 또는 얻을 흥분이 한 단계씩 더 커져야 한다. 5개 제목의 '온도'가 평평하면 실패다.
(4) 갈고리 연결 — 각 PART 제목은 바로 다음 PART를 안 보고는 못 배기게 끝나야 한다. PART 1을 읽으면 "그럼 어떻게?"가, PART 3을 읽으면 "그 다음엔?"이 자동으로 떠오르게.

[이어 읽기 예시 — 이렇게 '한 편'으로 읽혀야 한다]
PART 1. 1년 안에 99%가 다시 무너지는 결정적 이유      ← (바닥: 너도 곧 무너진다)
PART 2. 의지로 잡으려는 순간 뇌는 반대로 움직인다       ← (깨달음: 그게 네 탓이 아니었다)
PART 3. 회복의 8할이 결판나는 폭발 직후 90초          ← (사건: 바로 이 순간이 갈림길)
PART 4. 한 번 회복한 뇌는 다시는 무너지지 않는다        ← (전환: 이제 너는 다른 사람)
PART 5. 단단해진 멘탈은 인생을 통째로 바꾼다           ← (도약: 1년 뒤 완전히 다른 인생)
→ 5줄을 이어 읽으면 "무너지던 내가 → 원인을 알고 → 결정적 순간을 잡아 → 다른 사람이 되고 → 인생이 바뀐다"는 한 편의 이야기가 또렷이 보인다. 이 '세로 흐름'이 안 보이면 제목이 아무리 자극적이어도 다시 써라.

한 줄 한 줄은 다음 챕터를 못 참게, 5줄 전체는 사지 않으면 결말이 궁금해 미치게 만들어야 한다.

[컨셉명 등장 규칙 - 3번, 자연스럽게 분산]
컨셉명은 챕터 제목에는 등장 X. 소제목에만 정확히 3번 자연스럽게 넣는다.

⚠️ 컨셉명 글자 수/단어 수를 단정하는 표현 절대 금지 (가장 자주 어기는 실수!)
   ❌ "○○○ 세 글자에서 시작한다" — 컨셉명이 3글자 아니면 거짓말이 됨
   ❌ "○○○ 네 글자만 기억하라" — 글자 수 단정 금지
   ❌ "단 한 단어가 모든 걸 바꾼다" — 단어 수 단정 금지
   ✅ 컨셉명이 'DPS'(3자)든 '단단한 매출 구조'(8자)든 '복리 자산 공식'(7자)이든 모두 자연스럽게 작동하는 문장만 사용

[PART 1 마지막 소제목 - 도입] (아래 풀에서 1개 선택, 책에 가장 어울리는 것)
   • "결국 모든 답은 '○○○' 안에 있었다"
   • "이 책의 모든 페이지는 '○○○' 하나를 향해 간다"
   • "지금부터 '○○○' 단 하나만 기억하면 된다"
   • "여기서부터 진짜 이야기, '○○○'가 시작된다"
   • "마지막에 도달하는 곳은 결국 '○○○'다"
   • "'○○○'를 만나기 전과 후는 완전히 다른 게임이다"
   • "이 모든 혼란을 한 줄로 정리하는 '○○○'"

[PART 3 또는 PART 4 안 1곳 - 작동·전환점]
   • "○○○가 본격 작동하기 시작하는 4가지 신호"
   • "○○○를 처음 적용한 사람들이 가장 먼저 느낀 변화"
   • "○○○가 통장에 처음 흔적을 남기는 순간"
   • "○○○ 한 달 차에 가장 먼저 무너지는 한 가지"

[PART 5 마지막 소제목 - 확장·사회증명]
   • "○○○를 익힌 사람들의 5년 뒤가 완전히 다른 이유"
   • "○○○로 자리잡은 사람들이 다시는 돌아가지 않는 이유"
   • "○○○ 이후, 1년 만에 가장 크게 달라지는 단 한 가지"

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📐 챕터 제목 형식 (가장 중요)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[형식]
- 9~18자의 자연스러운 한국어 한 문장. 길게 풀어 설명하지 말고 짧고 단단하게 쳐라("노력보다 자리가 먼저다"처럼 짧고 센 게 더 강하다). 단, 짧다고 라벨·단어 나열이 되면 안 됨 — 끝까지 '말이 되는 한 문장'
- 단어 한두 개짜리 라벨 금지, 라벨 분리(— ㅣ :) 금지
- 평서문 또는 단언형. "~이유"로 끝나는 설명체는 한 PART에만 사용
- 명사 엔딩과 동사 엔딩을 챕터별로 섞어라 (5개 모두 명사 엔딩 금지)
- 5개를 이으면 [좌절 → 통념 박살 → 첫 사건 → 안정화 → 도약]의 5막이 '시간 순서'로 또렷이 보여야 한다 (위 🎯 스토리 흐름 강제 규칙 준수)
  · PART 1 = 가장 아픈 바닥(잃을 것), PART 5 = 가장 벅찬 도약(얻을 것). 1→5로 갈수록 감정 온도가 한 단계씩 올라가야 한다. 5개 제목의 세기가 평평하면 실패.
  · 각 PART 제목 끝은 다음 PART를 안 보면 못 배기게(갈고리). 5개가 따로 노는 독립 슬로건이면 즉시 다시 써라.

[좋은 예 - 성공적인 멘탈 관리 비결]
PART 1. 1년 안에 99%가 다시 무너지는 결정적 이유
PART 2. 의지로 잡으려는 순간 뇌는 반대로 움직인다
PART 3. 회복의 8할이 결판나는 폭발 직후 90초
PART 4. 한 번 회복한 뇌는 다시는 무너지지 않는다
PART 5. 단단해진 멘탈은 인생을 통째로 바꾼다

[좋은 예 - 30대 직장인 N잡 월 500]
PART 1. 직장인 99%가 부업 30일을 못 버티는 이유
PART 2. 노력보다 자리가 먼저다
PART 3. 첫 30만원이 통장에 찍힌 그날
PART 4. 새벽 3시에도 매출이 들어온다
PART 5. 월 500 다음, 억대 수익으로 가는 길

[나쁜 예 - 즉시 폐기]
- "발굴", "폭로", "전환" 같은 한두 단어짜리
- "발굴 — 부업의 90%는 첫 단추에서 망한다" (라벨 + 대시)
- "DPS의 첫 관문, 노력 없이도 돈이 따라오는 자리를 찾는 법" (시스템명 라벨화 + 너무 김)
- "이제 무너지는 게 더 이상 사건이 아니다" (사건이 아니다 ← 말이 안 됨)
- "한 번 흔들려도 다음 날엔 흔적도 없다" (AI식 과장)
- "월수도 시스템의 첫 설계" (의미 불명 + 설계라는 설명체 어휘)
- "MDS 파이프라인" (파이프라인이라는 영어 외래어를 시스템 접미사로)
- "주가 -12% 떨어져도 매도 안 하는 뇌 회로가 박혔다" (뇌 회로가 박히다 = 어법 어색, 비유 남용)
- "신경회로가 새로 깔리고 있다는 신호" 류 (한 번까진 OK, 같은 비유 두 번 X)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
✍️ 소제목 톤: 한국 자기계발 베스트셀러 + 마케팅 뇌과학
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

각 PART당 소제목 4개. 모두 다른 패턴 + 명사·동사 엔딩 섞기 (4개 모두 같은 엔딩 금지).

[좋은 패턴 풀 - 매번 다르게]
1. 통계 충격형 — "1년 안에 99%가 다시 무너진다", "한 달 차에 90%가 다시 무너지는 정확한 이유"
2. 인지부조화형 — "의지로 버틴 사람일수록 더 크게 무너진다"
3. 뇌과학 권위형 — "감정이 폭발할 때 뇌는 정확히 어디부터 멈추는가", "신경회로가 새로 깔리고 있다는 4가지 신호"
4. 임박감/손해회피형 — "이 90초를 놓치면 다시 일어서는 데 며칠이 걸린다"
5. 정체성 전환형 — "한 번 회복한 뇌는 다시는 무너지지 않는다"
6. 호기심 갭형 — "단번에 회복한 사람들이 모두 무의식적으로 하는 행동"
7. 이중/삼중 보상형 — "회복 후 6개월 만에 통장과 인간관계가 함께 변한다"
8. 인그룹 사회증명형 — "○○를 익힌 사람들의 5년 뒤가 완전히 다른 이유"

[★ 가장 중요 — 목차만 보고 결제하게 만들어라 (설명문 금지)]
이 목차의 단 하나의 목적: 서점에서 5초 훑은 사람이 "이건 사야 해"라고 결제하게 만드는 것.
정보 전달은 본문이 한다. 목차는 100% 구매 욕구만 자극한다.

지금 가장 흔한 실패 = "설명문처럼 평이함". 아래 셋 중 하나라도 걸리면 그 줄은 죽은 줄이니 다시 써라:
  (1) 정보 전달형("~하는 법", "~의 중요성", "~란 무엇인가") — 교과서 목차
  (2) 실행 디테일(종목·비율·타이밍·순서)까지 다 알려줌 — 본문 살 이유 사라짐 (단, 작동 '원리·지렛대'의 큰 흐름은 오히려 보여줘야 믿고 산다)
  (3) 자극이 없음 — 심장이 안 뛰면 결제 안 한다

[자청식 후킹 8가지 무기 — 소제목 20개 중 최소 14개에 적용 (밍밍한 정보형은 거의 없어야 함)]
1. 통념 정면 박살: "열심히 할수록 가난해지는 이유", "절약이 당신을 평생 가난하게 만든다"
2. 자기기만 적발(뜨끔하게): "공부만 하다 1년을 통째로 날리는 사람들의 5가지 변명"
3. 단정 선언(반박 불가 톤): "결국 전부 이거 하나에서 갈렸다", "답은 처음부터 정해져 있었다"
4. 날 선 대비: "버는 사람과 버는 척하는 사람", "3년 버틴 사람과 3개월에 접은 사람의 통장"
5. 충격 숫자: "10명 중 8명이 1년 안에 무너지는 진짜 이유", "상위 3%만 아는 한 가지"
6. 공포·손실 회피: "지금 이걸 모르면 5년 뒤 똑같은 자리에 있다", "당신이 놓치는 사이 벌어지는 일"
7. 미스터리 갭(정체를 가림): "퇴사 3개월 만에 다시 돌아온 회사원이 깨달은 단 하나"
8. ★ 독자 행동 직접 저격 (가장 강한 무기) — "너 지금 이러고 있지?" 톤으로 독자의 현재 행동·습관·변명을 정확하게 집어내라. 독자가 어제 한 행동을 보여서 "어? 내 얘기잖아"라고 멈춰야 결제한다.
   예: "퇴사 준비를 3년째 하고 있는데 통장은 그대로인 사람의 패턴"
   예: "증권사 앱을 하루 세 번 이상 열어본 해에 수익률이 가장 낮았다"
   예: "조건이 갖춰지면 시작하겠다고 한 사람이 5년째 같은 말을 반복하는 이유"
   예: "유튜브 강의 50개를 들었는데 통장은 그대로인 사람들의 공통점"
   ※ 소제목 20개 중 최소 5개는 이 무기를 써라. 독자의 현재 행동을 실제로 알고 있어야 쓸 수 있으므로, [타겟 독자]와 [독자가 지금 느끼는 통증]을 다시 읽고 그 사람이 오늘 실제로 한 행동·생각·변명 중 가장 뜨끔한 것을 골라라."

[좋은 변환 — 평이 → 자극, 어법은 자연스럽게]
- "배당주 고르는 법" → "10년 배당을 받고도 한 푼도 못 쓴 사람들의 공통점"
- "복리의 중요성" → "같은 돈을 넣었는데 7년 뒤 잔고가 두 배로 갈린 이유"
- "감정 관리가 필요하다" → "계좌를 자주 열어본 해일수록 오히려 수익률이 낮았던 이유"
- "분산 투자를 하자" → "한 종목에 몰빵한 사람이 그해 오히려 덜 잃은 까닭"
- "꾸준함이 답이다" → "3년을 버틴 사람과 3개월에 그만둔 사람의 통장이 비슷했던 이유"

[규칙]
- 소제목 절반 이상에 구체 숫자(금액·기간·비율·인원)를 넣는다.
- 통념을 뒤집거나 자기기만을 찌르는 줄을 PART마다 최소 1개.
- 답을 같은 줄에서 다 말하지 마라. "왜/어떻게/무엇이"의 정체는 본문에 숨긴다.
- 사람을 모욕하지 마라(욕설·인신공격 금지). 통념과 '행동'을 때리되 독자를 적으로 만들지 않는다.
- 자청·역행자 등 특정 작가의 '단어'는 쓰지 말 것. 톤만 가져온다.
- 자극을 위해 억지 비유나 말 안 되는 조합을 만들지 마라. 아래 [자연스러움 원칙]이 항상 우선.

[자연스러움 원칙 - 절대 어기지 말 것]
⚠️ 어법이 1순위다. 호기심보다 "말이 되는가"가 먼저다. 어법이 어색하면 호기심을 줄여서라도 자연스러운 문장으로 바꾼다.
- 모든 소제목/챕터 제목은 한국어 원어민 편집자가 손대지 않고 통과시킬, 문법적으로 완결된 자연스러운 문장이어야 한다.
- 단어를 억지로 조합한 "그럴듯해 보이지만 뜻이 안 통하는" 표현은 즉시 폐기. 특히 은유·비유를 무리하게 끼워 맞추지 마라.
  ❌ "배당 사이 파이프 굵기 차이는 정확히 얼마인가" (파이프 굵기? 배당 사이? — 무슨 말인지 알 수 없음)
  ❌ "수익이 자라는 토양의 산도를 맞추는 법" (억지 비유, 어법 어색)
  ✅ "같은 배당주를 사도 누구는 매달 80만원을 받고 누구는 한 푼도 못 버는 이유" (구체 금액으로 자연스럽고 강력)
- 어법 검사: 주어와 서술어가 자연스럽게 연결되는가? ("회복이 굴러간다" X — 회복은 굴러가지 않음)
- 대비 구조의 짝 맞추기(매우 중요): "A는 ~, B는 ~" 형태로 두 사람/상황을 대비할 때, 양쪽 서술어가 같은 축에서 정반대가 되어야 한다. 받는 것의 반대는 "못 받는 것", 버는 것의 반대는 "못 버는 것"이다. 축이 어긋나면 즉시 폐기.
  ❌ "같은 ETF를 샀는데 누구는 월세처럼 받고 누구는 한 푼도 못 쓴 이유" (받다 ↔ 못 쓰다 — 축이 어긋남)
  ✅ "같은 ETF를 샀는데 누구는 월세처럼 받고 누구는 한 푼도 못 버는 이유" (받다 ↔ 못 벌다 — 같은 축)
- 비유는 '명사'로만 쓰고 '되다/안 되다'의 대상으로 삼지 마라(비유 오용 금지). "월세"는 받는 금액의 비유지, "월세가 된다/안 된다"처럼 술어의 주어가 될 수 없다. 어색하면 비유를 버리고 '구체 금액'으로 직진하라 — 이게 더 강력하다.
  ❌ "10명 중 8명이 배당주를 사도 월세가 안 되는 이유" (월세가 안 된다? — 비유 오용, 말 안 됨)
  ✅ "10명 중 8명이 배당주를 사도 월 100만원조차 못 받는 이유" (구체 금액 — 자연스럽고 더 셈)
  ✅ "배당주 3년을 모아도 매달 통장에 30만원도 안 들어오는 사람들의 공통점" (구체 금액·장면)
- 추상 개념 + 물리·생물 동사 금지: 계좌·통장·재투자·수익·복리·멘탈·습관·시스템 같은 추상/사물에 "켜다/끄다/돌리다/감다/조이다/풀다/꽂다/박다/심다/굴린다/얹는다/멈춘다/자란다/살아난다/숨쉰다/깨어난다" 같은 동사를 붙이지 마라. 동사가 그 대상에 물리적으로 말이 되는지 검사하라.
  ❌ "캘린더에 박은 습관" (일정을 '박는다'? 어색)  ✅ "캘린더에 기재한 습관" / "매일 같은 시간에 반복한 습관" (자연스러움)
  ❌ "통장에 꽂히는 30만원" (꽂히다 = 비유 오용)  ✅ "매달 통장에 들어오는 30만원"
  → 의심되면 그 사물에 자연스럽게 쓰는 동사로 바꾼다(캘린더↔기재하다/적다, 돈↔들어오다/받다, 습관↔반복하다/들이다 등).
  ❌ "재투자를 켠 사람과 끈 사람" / "수익률 화면을 끄니 계좌가 살아났다" / "내버려 두면 계좌가 자란다"
  ✅ "수익을 다시 넣은 사람과 빼서 쓴 사람의 3년 뒤" / "계좌를 덜 들여다본 해에 수익률이 더 높았던 이유" (대비·호기심은 살리되 주어-서술어가 말이 되게)
- 과장 형용사 금지: "흔적도 없다", "통째로", "완전히" (꼭 필요할 때만)
- 추상 X, 구체 O: "많은 사람" → "월급 280만원짜리 7년차 회사원"
- 도구/플랫폼명 적극: 네이버, 카카오, 노션, 카톡, 캘린더, 구글 시트
- 출력 직전, 모든 줄을 소리내어 읽어라. 원어민이 한 번에 이해 못 하거나 "이게 무슨 말이지?" 싶은 줄은 전부 다시 쓴다.

[한 줄 = 한 호흡 — 정보 욱여넣기가 비문과 '노잼'의 가장 큰 원인]
한 소제목에 비교·숫자·기간·원리를 두세 개씩 밀어 넣으면 반드시 문장이 꼬여 비문이 되고, 호기심도 흩어져 끌리지 않는다. 가장 센 핵심 '하나'만 남기고 나머지는 과감히 쳐내라. 짧고 선명한 줄이 더 자연스럽고 훨씬 더 읽고 싶어진다.
✗ "배당률 5%를 찾아 헤맨 3년과 3%로 시작한 1년의 통장이 뒤집힌 순간"
   → 비교(5%·3%)+기간(3년·1년)+결과를 한 줄에 다 욱여넣어 비문. '통장이 뒤집히다'도 어색.
   ✓ "고배당만 좇던 3년보다, 낮은 배당으로 묵힌 1년이 더 많이 남은 이유" (비교를 한 축으로, 의외성으로 호기심)
✗ "계좌를 하루 세 번 확인한 사람이 수익률을 스스로 갉아먹은 증거"
   → '수익률을 갉아먹다'는 어색한 은유, '증거'로 끝나니 뻣뻣하고 김이 샌다.
   ✓ "하루에도 몇 번씩 계좌를 열어본 해에 오히려 수익률이 가장 낮았다" (행동→의외의 결과, 자연스럽고 궁금함)

[은유는 버리고 직진하라 — 어색한 은유는 어법도 자극도 둘 다 죽인다]
'갉아먹다 / 뒤집히다 / 몸집을 불리다 / 굴러가다 / 살아나다 / 깨어나다' 같은 은유를 돈·계좌·수익·복리·멘탈 같은 추상·사물에 붙이면 어색해지고 호기심도 식는다. 은유를 빼고 '구체적 장면 + 숫자 + 의외의 사실'로 직진하면 더 자연스럽고 더 센 호기심이 나온다. (예: "복리가 굴러간다" X → "재투자만 했는데 5년 뒤 받는 배당이 두 배가 됐다" O)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🚫 즉시 폐기 표현
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

특정 작가 고유어 (절대 금지): 자청 / 역행자 / 자의식 해체 / 유전자 역행 / 원시인 / 추월차선 / 아토믹 해빗 / 언카피어블
AI 클리셰: 졸업 / 정체 / 마지막 한 수 / 다른 차원 / 결정적 시그널 / 진짜 게임 / 흔적도 없다 / 사건이 아니다 / 회로가 박혔다
시스템 의인화: "○○가 멈춘 날", "○○ 위에 얹다", "○○를 졸업한", "회복이 굴러가다"
뇌과학 비유 남용: "뇌 회로가 박혔다", "뇌 회로가 새로 깔린다" (전체 목차에 뇌·신경회로는 사실 진술로 1~2회만, 비유 남용 X)
밍밍: 효과적인 / 성공적인 / ~의 모든 것 / ~하는 방법 / 알아야 할 / 의 중요성
유치 비유: 나침반 / 열쇠 / 보물 / 황금 / 마법 / 파이프라인 / 엔진 / 톱니바퀴 / 사이클 / 눈덩이
참고서: 첫걸음 / 완벽가이드 / 핵심정리 / 기초/중급/고급 / 첫 설계 / 첫 셋업
챕터 제목 라벨: "발굴 —", "1단계:", "STEP 1." 같은 분리 형식
의문문 문어체: "왜 ~는 ~하지 못하는가" 식의 한 PART에 1개까지만
콜론(:) — 단 한 번도 쓰지 마라
숫자 중복 금지: 전체 목차에서 같은 숫자(예: 3시간 + 3개월) 두 번 등장 금지

★★ 표현·단어 반복 절대 금지 — 목차 전체(25줄)에서 같은 말은 '딱 한 번'만 ★★
한 번 쓴 단어·표현·문장 끝 패턴은 그걸로 끝이다. 두 번째부터는 무조건 다른 말로 바꾼다. (목차가 평이해 보이는 가장 큰 원인이 '같은 말 반복'이다.)
- 핵심 명사·동사·형용사 반복 금지: '무너지다'를 한 줄에 썼으면 다른 줄에선 '주저앉다 / 흔들리다 / 제자리로 돌아가다' 등 다른 표현으로. ('이유·차이·순간·진짜·결정적·공통점·단 하나' 등 자주 겹치는 말도 전부 한 번씩만)
- 문장 끝 패턴 반복 금지: '~하는 이유'로 끝나는 줄은 전체에서 1개만. 나머지는 '~한 순간 / ~의 차이 / ~하는 사람들 / ~한 그날 / ~인 신호' 등 25줄이 전부 다른 엔딩이 되도록.
- 같은 통계 표현 반복 금지: '99%'를 썼으면 다른 줄은 '열에 아홉 / 대부분 / 상위 3%만' 처럼 다르게.
- 예외: 조사·접속사 같은 기능어, 컨셉명(정확히 3번)과 '○○○' 자리표시는 반복으로 치지 않는다.
- 출력 직전 점검: 25줄의 명사·동사·문장 끝을 쭉 훑어, 두 번 이상 나온 단어·표현이 있으면 한쪽을 반드시 다른 말로 교체한 뒤 출력한다.
숫자 표기 - 부호 금지: "-12%" "+30%" 같은 부호 사용 X. "12% 폭락에도", "30% 상승하면" 식으로

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📝 출력 형식 (이 형식 외 어떤 텍스트도 출력 금지)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

PART 1. [13~18자, 좌절 + 통계 충격]
- [통계 충격 또는 결정적 함정형]
- [인지부조화/통념 박살형]
- [패턴 - 위 8개 중 다른 것]
- [컨셉명 첫 등장: 자연스러운 도입]

PART 2. [13~18자, 통념 박살 + 뇌과학 권위]
- [패턴]
- [패턴 - 다른 것]
- [패턴 - 또 다른 것]
- [패턴 - 또 다른 것]

PART 3. [13~18자, 첫 사건/결정적 순간]
- [패턴]
- [임박감/손해회피형]
- [패턴]
- [컨셉명 등장 가능: "○○○를 처음 적용한 사람들이 가장 먼저 느낀 변화" — PART 4에 넣을 거면 여기는 일반 패턴]

PART 4. [13~18자, 정체성 전환 선언]
- [컨셉명 등장 가능: "○○○가 본격 작동하는 4가지 신호" — PART 3에 안 넣었다면 여기에]
- [뇌과학 권위형 또는 통계형]
- [패턴]
- [패턴]

PART 5. [13~18자, 도약/이중 보상]
- [호기심 갭형]
- [이중/삼중 보상형]
- [패턴 - 또 다른 것]
- [컨셉명 세 번째 등장: 인그룹 사회증명]

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🔍 출력 전 자가 점검 (반드시 통과)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

출력 직전에 5가지 모두 통과해야 한다. 하나라도 No면 다시 써라.

체크 1. (스토리 흐름) 5개 PART 제목만 세로로 이어 읽었을 때, 같은 한 사람의 이야기가 '시간 순서(바닥→깨달음→사건→전환→도약)'로 이어지는가? 화자가 바뀌거나, 순서가 뒤섞이거나, 5줄이 따로 노는 독립 슬로건이면 실패 — 다시 써라.
체크 1-2. (감정 온도) PART 1이 가장 아프고(잃을 공포) PART 5가 가장 벅찬가(얻을 흥분)? 1→5로 세기가 한 단계씩 커지지 않고 평평하면 실패. 그리고 각 PART가 다음 PART를 안 보면 못 배기게 끝나는가?
체크 2. 통념 박살(인지부조화) 패턴이 5개 챕터+20개 소제목 안에 3개 이상 박혀있는가?
체크 3. 구체적 숫자(시간/금액/비율)가 8개 이상 등장하는가?
체크 4. 평이한 표현("~의 방법", "~의 모든 것", "~의 중요성", "효과적인", "성공적인")이 0개인가?
체크 5. 5초간 훑어본 가상 독자가 "이건 안 사면 손해"라고 느낄 만한 손실회피 트리거가 챕터 제목 5개 중 2개 이상에 있는가?
체크 6. (호기심+신뢰 균형) 소제목 20개 중 최소 12개가 '구체적 지렛대(행동·원리) 한 조각'을 드러내 "이렇게 하니까 되는구나"라는 믿음을 주는가? 막연히 '달라진다/완성된다/무너지지 않는다/따라온다'로 끝나는 순수 주장 줄이 PART당 2개 이상이면 실패 — 그 줄에 작동 원리나 지렛대를 한 조각 넣어 다시 써라. (디테일=종목·비율·타이밍은 여전히 숨긴다.)
체크 7. 컨셉명 글자 수를 단정하는 표현("세 글자", "네 글자", "단 한 단어")이 단 하나도 없는가? — 하나라도 있으면 즉시 전체 다시 쓰기.
체크 8. (어법 - 가장 중요) 모든 챕터 제목·소제목을 소리내어 읽었을 때, 원어민이 한 번에 이해되고 어법이 자연스러운가? 아래 셋 중 하나라도 걸리면 그 줄을 자연스럽게 다시 써라.
   (a) 한 줄에 비교·숫자·기간·원리를 두세 개씩 욱여넣어 길고 꼬인 비문("5%를 찾아 헤맨 3년과 3%로 시작한 1년의 통장이 뒤집힌 순간" 류) → 핵심 하나만 남기고 짧게.
   (b) 추상·사물에 어색한 은유 동사를 붙인 줄("수익률을 갉아먹다", "통장이 뒤집히다", "복리가 굴러가다" 류) → 은유 빼고 구체 장면·숫자로 직진.
   (c) "배당 사이 파이프 굵기 차이는…" 같은 억지 조합·뜻 모를 비유.
체크 9. (어법) 추상 개념에 기계·물리 동사를 붙인 줄("재투자를 켠/끈", "수익을 돌린다" 류)이 하나도 없는가? 있으면 즉시 자연스러운 표현으로 교체.
체크 10. (자극) 통념을 정면으로 뒤집거나 자기기만을 찌르는 줄이 PART마다 최소 1개 있는가? 전부 정보 전달형이면 평이한 것 — 다시 써라.
체크 11. (어법-대비) "A는 ~, B는 ~" 대비 줄에서 양쪽 서술어가 같은 축의 정반대인가? "받고 ↔ 못 쓴"처럼 축이 어긋난 줄이 하나라도 있으면 "받고 ↔ 못 받는"처럼 짝을 맞춰 다시 써라.
체크 12. (자극-결정타) 20개 소제목을 다시 훑어, 심장이 안 뛰는(밍밍한) 줄을 전부 골라내라. 그 줄은 [자청식 후킹 7가지 무기] 중 하나를 적용해 더 세게 다시 쓴다. "정보를 알려주는 줄"이 아니라 "안 읽으면 불안한 줄"이 되어야 한다.
체크 13. (비유 오용) "월세가 안 된다", "통장이 못 된다"처럼 비유 명사를 '되다/안 되다'의 대상으로 쓴 줄이 하나라도 있으면, 비유를 버리고 구체 금액("월 100만원조차 못 받는")으로 다시 써라.
체크 14. (뇌과학 겹치기) 20개 소제목 중 최소 6개는 뇌과학 트리거를 2개 이상 겹쳐 썼는가?(예: 사회증명+손실회피+구체금액) 한 줄에 트리거가 하나뿐인 평범한 줄이 절반을 넘으면 평이한 목차다 — 더 겹쳐 써라.
체크 15. (최종 결제 테스트 — 강화) 완성된 목차를 처음부터 끝까지 5초간 훑었다고 상상하라. "이 목차 정도는 나도 알겠다" 싶으면 실패다. 25줄을 한 줄씩 보며 '심장이 뛰는 정도'를 채점하라. 평범한 경쟁 도서 목차에도 들어갈 법한 줄(=교체 가능한 줄)은 전부 죽은 줄이다. 가장 약한 8줄을 골라 [자청식 후킹 7가지 무기] 중 하나로 한 단계 더 세게 다시 써라. "정보를 주는 줄"이 아니라 "안 읽으면 손해라 불안한 줄"이 25줄 중 다수가 될 때까지 반복한다. 단, 세기를 올리되 [자연스러움 원칙]과 어법 검사는 절대 희생하지 않는다 — 자극과 자연스러움이 부딪치면 어법이 이긴다.

체크 16. (반복 금지) 25줄의 핵심 단어·표현·문장 끝 패턴을 훑었을 때, 두 번 이상 나온 게 단 하나도 없는가?(기능어·컨셉명 제외) '이유/차이/결정적/신호/처음/가장 먼저/한 번' 같은 말이나 같은 숫자가 두 줄에 있으면 한쪽을 다른 표현으로 바꿔 다시 써라.
체크 17. (흐름-PART 중복) 5개 PART 제목 중 '사실상 같은 말'을 하는 PART가 둘 있는가?(예: 둘 다 "개선하면 매출 오른다") 있으면 흐름이 무너진 것 — 각 PART에 바닥/깨달음/첫 사건/전환/도약 중 '서로 다른 단 하나의 역할'만 주어 다시 써라.

목차만 출력. 콜론 금지. 매 소제목 다른 패턴. 명사·동사 엔딩 섞기. 같은 단어·표현·문장 끝은 전체에서 한 번만. 5개 PART는 이어 읽으면 서로 다른 5단계의 한 이야기. 각 PART는 정확히 소제목 4개. 컨셉명 글자수 단정 금지. 자가점검 결과는 출력하지 말 것. ★ 어법이 전부다 — 출력 직전 모든 줄을 소리내어 읽어라. 원어민이 한 번에 이해 못 하거나 '이게 무슨 말이지?' 싶은 줄은 자극을 줄이더라도 자연스러운 문장으로 바꿔 출력한다. 어색한 채로 내보내지 마라."""

        chapters = []
        subtopics = {}
        for _outline_attempt in range(2):
            outline_result = ask_ai(outline_prompt, 0.92, ensure_quality=True)
            if not outline_result:
                continue

            chapters = []
            subtopics = {}
            current_ch = None

            lines = outline_result.split('\n')
            for i, orig_line in enumerate(lines):
                line = orig_line.strip()
                if not line:
                    continue

                # 챕터 감지 (PART, 파트, Chapter, 1., 2. 등 다양한 형식)
                is_chapter = False
                ch_name = None

                # 마크다운 강조/헤더 기호를 벗겨낸 감지용 라인
                # (신형 Claude 모델이 **PART 1. ...**, ## PART 1, > 등으로 헤더를 감싸 출력해도 인식)
                detect_line = re.sub(r'^[\s>#\*_`~]+', '', line)
                detect_line = re.sub(r'[\*_`~]+$', '', detect_line).strip()

                # PART 1. 제목 형식
                if re.match(r'^(PART|파트|Part)\s*\d+[\.\s]', detect_line, re.IGNORECASE):
                    is_chapter = True
                    ch_name = detect_line
                # Chapter 1. 제목 형식
                elif re.match(r'^(Chapter|챕터)\s*\d+[\.\s]', detect_line, re.IGNORECASE):
                    is_chapter = True
                    ch_name = detect_line
                # 마크다운 헤더 형식
                elif re.match(r'^#+\s*(PART|파트|Chapter|챕터|\d+)', line, re.IGNORECASE):
                    is_chapter = True
                    ch_name = re.sub(r'^#+\s*', '', line)
                # 1. 제목 형식 (숫자로 시작, 들여쓰기 없음)
                elif re.match(r'^\d+[\.\)]\s', detect_line) and not orig_line.lstrip('*_`~ ').startswith(' '):
                    is_chapter = True
                    ch_name = detect_line
                # 【1부】 형식
                elif re.match(r'^[【\[]?\s*\d+\s*(부|장|편)[】\]]?', detect_line):
                    is_chapter = True
                    ch_name = detect_line

                if is_chapter and ch_name:
                    ch_name = re.sub(r'^[#\*\-\s]+', '', ch_name)
                    ch_name = ch_name.replace('**', '').replace('*', '').replace('#', '').strip()
                    if ch_name and len(ch_name) > 3:
                        current_ch = ch_name
                        if current_ch not in chapters:
                            chapters.append(current_ch)
                            subtopics[current_ch] = []
                elif current_ch:
                    # 소제목 감지
                    is_subtopic = False
                    st_name = None

                    # - 소제목 형식
                    if re.match(r'^[\-\•\·\*\→\▶]\s*', line):
                        is_subtopic = True
                        st_name = re.sub(r'^[\-\•\·\*\→\▶]\s*', '', line)
                    # 1) 소제목, a) 소제목 형식
                    elif re.match(r'^[a-z\d][\)\.\:]\s', line, re.IGNORECASE):
                        is_subtopic = True
                        st_name = re.sub(r'^[a-z\d][\)\.\:]\s*', '', line, flags=re.IGNORECASE)
                    # 들여쓰기된 라인
                    elif orig_line.startswith('  ') or orig_line.startswith('\t'):
                        is_subtopic = True
                        st_name = line.lstrip('- •·*→▶0123456789.):\t ')
                    # 챕터가 아닌 일반 텍스트 (이전이 챕터였고, 현재가 짧은 문장이면 소제목으로 간주)
                    elif len(chapters) > 0 and not re.match(r'^(PART|파트|Part|Chapter|챕터|\d+[\.\)])', line, re.IGNORECASE):
                        if len(line) > 5 and len(line) < 100:
                            is_subtopic = True
                            st_name = line.lstrip('- •·*→▶0123456789.):\t ')

                    if is_subtopic and st_name:
                        st_name = st_name.replace('**', '').replace('*', '').replace('#', '').strip()
                        st_name = re.sub(r'^\d+[\.\)\:]\s*', '', st_name)  # 앞 숫자 제거
                        if st_name and len(st_name) > 3 and len(subtopics[current_ch]) < 4:
                            # 챕터 이름과 동일하면 스킵
                            if st_name.lower() != current_ch.lower() and st_name not in subtopics[current_ch]:
                                subtopics[current_ch].append(st_name)

            # 챕터가 충분히 파싱됐으면 재시도 중단, 부족하면 한 번 더 생성 시도
            if len(chapters) >= 5:
                break

        if chapters:
            st.session_state['outline'] = chapters
            st.session_state['chapters'] = {}
            for ch in chapters:
                st.session_state['chapters'][ch] = {
                    'subtopics': subtopics.get(ch, []),
                    'subtopic_data': {s: {'questions': [], 'answers': [], 'content': ''} for s in subtopics.get(ch, [])}
                }

        # 목차가 생성되지 않았으면 기본 목차 생성
        if not st.session_state.get('outline'):
            progress_placeholder.warning("⚠️ AI 목차 생성/파싱에 실패해 기본 목차를 사용합니다. (API 키·모델 설정을 확인하거나 다시 시도해 주세요)")
            default_chapters = [
                "PART 1. 왜 지금인가",
                "PART 2. 진짜 비밀",
                "PART 3. 실전 공식",
                "PART 4. 수익화",
                "PART 5. 다음 단계"
            ]
            default_subtopics = {
                default_chapters[0]: [f"90%가 {topic}에 실패하는 이유", "아무도 말해주지 않는 진실", "지금 시작해야 하는 3가지 이유"],
                default_chapters[1]: ["전문가들이 숨기는 핵심 원칙", f"{topic}의 본질을 꿰뚫는 법", "이것만 알면 절반은 성공"],
                default_chapters[2]: ["바로 써먹는 5단계 공식", "실패 없이 시작하는 체크리스트", "첫 달에 결과 내는 비법"],
                default_chapters[3]: ["월 100만원 만드는 구조", "자동화로 시간 벌기", "확장 전략 A to Z"],
                default_chapters[4]: ["1년 후 당신의 모습", "다음 레벨로 가는 로드맵", "지금 바로 해야 할 첫 번째 행동"]
            }
            st.session_state['outline'] = default_chapters
            st.session_state['chapters'] = {}
            for ch in default_chapters:
                st.session_state['chapters'][ch] = {
                    'subtopics': default_subtopics.get(ch, []),
                    'subtopic_data': {s: {'questions': [], 'answers': [], 'content': ''} for s in default_subtopics.get(ch, [])}
                }

        # 저자명 및 인터뷰 데이터 저장
        st.session_state['author_name'] = interview_data.get('author_name', '')
        st.session_state['interview_data'] = interview_data
        st.session_state['topic'] = topic

        progress_placeholder.success("✅ 목차 생성 완료! 목차를 확인하고 수정할 수 있습니다.")
        return True

    except Exception as e:
        progress_placeholder.error(f"오류 발생: {str(e)}")
        return False

def regenerate_single_subtopic(chapter_name, subtopic_index, existing_subtopics):
    """개별 소제목 AI 재생성 - 자기계발 베스트셀러 톤"""
    topic = st.session_state.get('topic', '')
    book_concept = st.session_state.get('book_concept', '')

    # 기존 소제목들 (중복 방지용)
    other_subtopics = [s for i, s in enumerate(existing_subtopics) if i != subtopic_index]

    prompt = f"""당신은 한국 자기계발 베스트셀러 편집자입니다. 목차만 보고 결제하게 만드는 소제목 하나를 써주세요.

[책 컨셉]
{book_concept}

[챕터]: {chapter_name}
[주제]: {topic}

[기존 소제목들 - 이것들과 완전히 다르게]
{chr(10).join(f'- {s}' for s in other_subtopics)}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🔥 한국 자기계발 베스트셀러 톤 예시 (참고만, 그대로 베끼지 말 것)
- "5분만에 돈 되는 사업 아이템 찾는 비법"
- "월급 280만원이 월 1,000만원으로 바뀌기까지 47일"
- "회사 몰래 부업하다 잘리는 직장인의 진짜 이유"
- "노력이 결과로 안 바뀌는 결정적 함정"
- "첫 정산 받고 인생이 달라진 그날의 기록"

✅ 형식 (하나 선택, 매번 다른 패턴):
- 비법/법형: "5분만에 ~하는 비법", "~하는 정확한 방법"
- 충격형: "~는 거짓말이다", "~하면 오히려 망한다"
- 간증형: "~받고 인생이 달라졌다", "~를 깨닫기까지"
- 도발형: "~은 필요 없다", "~만 있으면 된다"
- 질문형: "왜 ~은 실패하는가"
- 숫자형: "정확히 47일 만에 일어난 일"
- 호기심 갭형(가장 강력): "월 500을 찍은 사람들이 가장 먼저 끊은 습관 한 가지" (결과만 보이고 방법은 숨김)

🧲 호기심 갭 우선 — 결과/사건/디테일은 보이고, 원리/방법/순서는 숨기면 결제 전환률 폭증

⚠️ 컨셉명을 쓰는 경우, 글자 수 단정 표현 절대 금지 ("세 글자", "네 글자", "단 한 단어" 등 — 컨셉명 길이와 안 맞으면 즉시 폐기)

❌ 금지:
- 자청, 역행자, 유전자, 원시인, 추월차선 등 특정 작가 고유 표현
- 설명서 표현: "~의 이해", "~하는 방법"
- 유치한 비유: 나침반, 열쇠, 마법, 톱니바퀴, 파이프라인
- AI 어휘: 효과적인, 다양한, ~를 통해
- 기존 소제목과 비슷한 패턴
- 컨셉명 글자 수 단정 ("세 글자에서 시작한다" 류)

소제목 하나만 (15~30자, 기호 없이):"""

    result = ask_ai(prompt, 0.9)
    if result:
        return result.strip().strip('"').strip("'").strip('-').strip()
    return None

def regenerate_chapter_subtopics(chapter_name, chapter_index):
    """챕터의 모든 소제목 AI 재생성 - 자기계발 베스트셀러 톤"""
    topic = st.session_state.get('topic', '')
    book_concept = st.session_state.get('book_concept', '')
    outline = st.session_state.get('outline', [])

    # 다른 챕터들의 소제목들 (중복 방지)
    other_chapter_subtopics = []
    for ch in outline:
        if ch != chapter_name:
            ch_data = st.session_state['chapters'].get(ch, {})
            other_chapter_subtopics.extend(ch_data.get('subtopics', []))

    # 챕터별 역할 정의
    chapter_roles = {
        0: "착각/각성 - 독자가 몰랐던 불편한 진실을 폭로",
        1: "해체 - 기존 상식과 믿음을 완전히 무너뜨림",
        2: "구조/재구축 - 저자만의 새로운 방법론 제시",
        3: "실전 - 구체적이고 따라할 수 있는 방법",
        4: "도약 - 변화된 미래와 행동 촉구"
    }
    current_role = chapter_roles.get(chapter_index, "핵심 내용 전달")

    prompt = f"""당신은 한국 자기계발 베스트셀러 편집자입니다. 이 챕터의 소제목 4개를 결제하고 싶게 써주세요.

[책 컨셉]
{book_concept}

[주제]: {topic}
[챕터]: {chapter_name}
[이 챕터의 역할]: {current_role}

[다른 챕터 소제목들 - 완전히 다르게 써야 함]
{chr(10).join(f'- {s}' for s in other_chapter_subtopics[:8])}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🔥 한국 자기계발 베스트셀러 톤 예시 (참고만, 그대로 베끼지 말 것)
- "5분만에 돈 되는 사업 아이템 찾는 비법"
- "월급 280만원이 월 1,000으로 바뀌기까지 47일"
- "회사 몰래 부업하다 잘리는 직장인의 진짜 이유"
- "첫 정산 받고 인생이 달라진 그날의 기록"
- "노력이 결과로 안 바뀌는 결정적 함정"
- "통장 잔고 23만원에서 시작한 한 가지 시도"

✅ 4개 소제목 모두 다른 형식으로 (4개 다 같은 패턴 금지):
1번: 통계 충격/숫자형 (예: "정확히 47일째에 달라진 한 가지", "1년 안에 99%가 다시 무너지는 결정적 이유")
2번: 간증/사건형 (예: "27만원짜리 첫 정산서가 알려준 한 가지", "~받고 인생이 달라진 그날")
3번: 충격/통념 박살형 (예: "노력할수록 가난해진다", "~만 있으면 된다")
4번: 호기심 갭형 (결과만 보이고 방법은 숨김 — 예: "월 500을 찍은 사람들이 가장 먼저 끊은 습관 한 가지", "성공한 부업러가 매일 밤 11시에 반드시 끄는 것")

🧲 호기심 갭 규칙 (반드시 4개 중 1개 이상 포함)
   - 결과/사건/디테일은 보이고, 원리/방법/순서는 숨겨라
   - "방법", "비법", "노하우"로 끝내면 갭이 닫혀 결제 안 함
   - 한 줄에 "묘하게 구체적인 디테일 + 의문"이 같이 있어야 작동

⚠️ 컨셉명을 쓰는 경우, 글자 수 단정 표현 절대 금지 ("세 글자", "네 글자", "단 한 단어" 등)

❌ 금지:
- 자청, 역행자, 유전자, 원시인, 추월차선 등 특정 작가 고유 표현
- 설명서 표현: "~의 이해", "~하는 방법", "효과적인", "다양한"
- 유치한 비유: 나침반, 열쇠, 마법, 톱니바퀴, 파이프라인
- 같은 패턴 반복

소제목 정확히 4개만 출력 (줄바꿈으로 구분, 기호/번호 없이, 각 줄 15~30자):"""

    result = ask_ai(prompt, 0.8)
    if result:
        lines = [line.strip().strip('"').strip("'").strip('-').strip() for line in result.strip().split('\n') if line.strip() and len(line.strip()) > 5]
        return lines[:4] if lines else None
    return None

def generate_body_from_outline(interview_data, progress_placeholder):
    """생성된 목차를 기반으로 본문만 생성"""
    try:
        topic = interview_data.get('topic', '')
        book_concept = st.session_state.get('book_concept', '')

        if not st.session_state.get('outline') or not st.session_state.get('chapters'):
            progress_placeholder.error("먼저 목차를 생성해주세요.")
            return False

        # 본문 생성
        total_subtopics = sum(len(st.session_state['chapters'][ch]['subtopics']) for ch in st.session_state['outline'])
        done = 0

        for ch in st.session_state['outline']:
            ch_data = st.session_state['chapters'][ch]
            for sub in ch_data['subtopics']:
                done += 1
                progress_placeholder.info(f"✍️ 본문 작성 중... ({done}/{total_subtopics}) - {sub[:20]}...")

                # 이전 소제목들의 내용 요약 (중복 방지용)
                prev_contents = []
                for prev_ch in st.session_state['outline']:
                    if prev_ch == ch:
                        break
                    prev_ch_data = st.session_state['chapters'].get(prev_ch, {})
                    for prev_sub in prev_ch_data.get('subtopics', []):
                        prev_content = prev_ch_data.get('subtopic_data', {}).get(prev_sub, {}).get('content', '')
                        if prev_content:
                            prev_contents.append(f"- {prev_sub}: {prev_content[:100]}...")

                # 현재 챕터의 이전 소제목들
                current_ch_prev = []
                for prev_sub in ch_data['subtopics']:
                    if prev_sub == sub:
                        break
                    prev_content = ch_data.get('subtopic_data', {}).get(prev_sub, {}).get('content', '')
                    if prev_content:
                        current_ch_prev.append(f"- {prev_sub}: {prev_content[:100]}...")

                prev_summary = "\n".join(prev_contents[-5:] + current_ch_prev) if (prev_contents or current_ch_prev) else "없음"

                # 소제목 인덱스에 따라 다른 시작 스타일 선택
                # 더 다양한 시작 스타일 (15가지)
                hook_styles = [
                    "도발적 질문 (예: '왜 99%는 이걸 모를까요?')",
                    "충격적 고백 (예: '저도 3년간 완전히 잘못하고 있었습니다.')",
                    "반전 사실 (예: '사실 정반대였습니다.')",
                    "구체적 숫자 (예: '정확히 47일 만에 달라졌습니다.')",
                    "생생한 에피소드 (예: '그날 카페에서 노트북을 열었을 때였습니다.')",
                    "단호한 선언 (예: '결론부터 말씀드리겠습니다.')",
                    "대화체 시작 (예: '\"이게 진짜 돼요?\" 처음 들었을 때 저도 그랬습니다.')",
                    "before/after (예: '6개월 전만 해도 저는 완전히 다른 사람이었습니다.')",
                    "상식 뒤집기 (예: '노력하면 된다? 완전히 틀렸습니다.')",
                    "비유로 시작 (예: '이건 마치 고장난 네비게이션을 따라가는 것과 같습니다.')",
                    "독자 공감 (예: '혹시 이런 경험 있으신가요?')",
                    "미래 제시 (예: '3개월 후, 완전히 다른 결과를 보게 될 겁니다.')",
                    "실패담 (예: '처음엔 완전히 망했습니다.')",
                    "발견의 순간 (예: '그때 깨달았습니다. 방법이 틀렸던 거였죠.')",
                    "핵심 한 줄 (예: '핵심은 딱 하나입니다.')",
                ]
                current_hook_style = hook_styles[done % len(hook_styles)]

                # 저자 이름 가져오기
                author_name = interview_data.get('author_name', '') or st.session_state.get('author_name', '') or '저자'

                content_prompt = f"""당신은 한국 자기계발 베스트셀러 작가입니다. 첫 문장으로 독자를 멈춰 세우고, 권석천 기자 칼럼처럼 정확한 디테일과 통찰로 끝까지 끌고 가는 본문을 씁니다.

🚨🚨🚨 최우선 규칙 🚨🚨🚨
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
1. 첫 문장이 가장 중요! 반드시 이 스타일로 시작:
   👉 {current_hook_style}

2. 이전 내용과 절대 중복 금지!
   아래 내용은 이미 썼으니 완전히 다른 이야기를 해라:
   {prev_summary}

3. 독자 직접 호칭 금지
   ❌ "여러분", "당신", "독자님", "~하시는 분들"
   ✅ "저는", "우리는", "제가"

4. 특정 작가 고유 표현 절대 금지 (저작권/표절 위험)
   ❌ "자청", "자청아", "자청씨"
   ❌ "역행자", "유전자 역행", "유전자 오작동", "자의식 해체"
   ❌ "원시인", "추월차선", "아토믹 해빗", "언카피어블"
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[집필 정보]
주제: {topic}
챕터: {ch}
소제목: {sub}
핵심 방법론: {interview_data.get('core_method', '')}
저자의 구체적 노하우·실전 디테일(★ 본문에 적극 활용 — 수치·사례·단계·실수·기준·도구를 이 소제목에 맞게 녹여라): {interview_data.get('detailed_knowhow', '') or '(미입력)'}

[책 컨셉]
{book_concept}

이 책의 기법 이름은 "{st.session_state.get('method_name', '')}"으로 확정됨. 본문에서 기법을 언급할 때 반드시 이 이름 그대로 사용(약자 풀어쓰기·재작명 금지).

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
✍️ 본문 톤: 자기계발 후킹 + 권석천 칼럼 깊이
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[문체]
- 합쇼체 기본 ("~입니다", "~합니다") + 가끔 구어체 ("~거든요", "~더라고요")
- 현재 시제로 장면을 그리듯
- 짧은 문장과 긴 문장을 교차해 리듬을 만듦
- 추상보다 구체. "많은 사람" → "월급 280만원짜리 7년차 회사원"
- 결론부터, 분석은 그 다음

[권석천식 깊이 - 칼럼처럼 인과를 추적]
- 사실 → 분석 → 통찰 순서로 전개
- 가설 검증식 흐름: "왜 그럴까. 이유를 되짚어봤습니다"
- 사회적 맥락이나 통계, 책/논문 인용 자연스럽게 섞기
- 결론을 강요하지 말고, 독자가 스스로 깨닫게 단서를 깔기

[자기계발식 후킹 - 뇌를 멈춰 세우는 첫 문장]
- 사건/숫자/대사/의문 중 하나로 시작 (일반론 절대 금지)
- 본문 중간에 작은 반전 1회 ("그런데 진짜 흥미로운 건 그 다음이었습니다")
- 마지막 문장은 발견의 결과로서의 통찰 한 줄

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🎯 구체성을 끝까지 밀어붙여라
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

✅ 숫자: "많이" X → "월 340만원, 정확히 47일" O
✅ 도구/플랫폼: 네이버, 카카오, 노션, 카톡, 구글 시트 등 실제 이름
✅ 실행 순서: "무엇을 → 어디서 → 어떻게" 명시

✅ 사례:
- 내 경험: "제가 처음 시작했을 때", "그날 회의실에서"
- 타인 사례: 수강생, 지인 (이름은 가끔만, 매번 X)
- 사례 없어도 OK. 일반적 원리·논리 전개로 충분

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🚫 절대 금지
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
❌ 자청, 역행자, 자의식 해체, 유전자 역행, 원시인, 추월차선, 아토믹 해빗 (특정 작가/책 고유 표현)
❌ 같은 이름 반복 (민준, 지수가 계속 나오면 안 됨)
❌ "김씨", "이씨" 같은 성씨 호칭
❌ 유치한 표현: 후다닥, 짜잔, 대박, ㅋㅋ, 어마어마한
❌ 유치한 비유: 마법, 황금열쇠, 나침반, 로켓, 눈덩이
❌ 억지 메타포: 순환법, 엔진, 고리, 파이프라인, 톱니바퀴
❌ AI스러운: 중요합니다, 따라서, 결론적으로, ~를 통해, 다양한, 효과적인, 진정한
❌ 형식: 1. 2. 첫째, 둘째, 글머리 기호, 이모지

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📊 표는 '필요할 때만' (강제 아님)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
- 비교(전/후, A안/B안), 단계별 정리, 수치 묶음처럼 표로 보여주면 확 명확해지는 내용이 있을 때만 표 1개를 넣어라.
- 그런 내용이 없으면 표를 넣지 마라. 모든 소제목에 표가 있을 필요 없다(대략 3개 중 1개 정도가 자연스럽다).
- 표 형식은 반드시 아래 마크다운 파이프 형식. 첫 줄이 헤더다. (HTML 태그 금지)
  | 구분 | 기존 방식 | 새로운 방식 |
  | 비용 | 월 30만원 | 월 0원 |
  | 시간 | 하루 3시간 | 하루 30분 |
- 표 앞뒤로는 반드시 설명 문단을 둔다. 표만 툭 던지지 마라.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🔥 몰입·충격 어조 (자청식 몰입감 — 특정 작가 고유어는 제외)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

- 도입부터 통념을 정면으로 깬다. 독자가 "어, 내가 알던 거랑 반대네" 하고 멈칫하게.
- 글 중간마다 '작은 충격'을 심어라: 예상과 반대되는 사실, 뒤집히는 데이터, 의외의 결론.
- 한 호흡에 읽히는 리듬 — 짧은 단정문으로 치고, 긴 문장으로 풀고, 다시 끊는다.
- "왜?"를 계속 만들어 다음 문단을 안 읽고는 못 배기게 한다(궁금증 갭).
- 충격 뒤엔 반드시 '그래서 무엇을 어떻게'의 실질이 온다. 겁만 주고 끝내지 않는다.
- 자청·역행자 등 특정 작가의 '단어'는 절대 쓰지 않는다. 몰입감과 태도만 가져온다.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🎓 전문성 (아마추어 글과 가르는 지점)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

- 추상 주장 1개당 근거 1개(데이터·연구·구체 사례·숫자)를 반드시 붙인다.
- 메커니즘을 설명하라. "왜 그렇게 되는가"의 원리와 과정을 단계로 풀어라.
- 바로 따라 할 수 있게: 무엇을·어디서·어떤 순서로·얼마나, 구체적으로.
- 흔한 조언의 한계를 짚고, 더 정확한 기준·예외를 제시한다(깊이의 증거).

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📏 분량: 2400~3000자 (최소 2400자 이상, 전문성 있는 밀도로)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

⚠️ 길이는 사례·데이터·단계별 설명·메커니즘으로 채운다. 같은 말 반복·미사여구로 늘리면 즉시 실패.

'{sub}' 본문 작성.
- 시작: {current_hook_style}
- 이전 내용과 완전히 다른 새로운 이야기
- 권석천 칼럼처럼 사실 → 분석 → 통찰 순서로 인과를 추적
- 마지막 한 줄에 발견된 통찰 하나
- 표가 정말 도움이 되는 내용이면 위 마크다운 파이프 형식으로 1개까지 넣어도 됨(아니면 넣지 말 것)

⛔ 절대 금지: 본문 첫 줄에 소제목('{sub}')을 다시 쓰지 마라.
   소제목은 위에 이미 표시되므로, 본문은 곧장 첫 후킹 문장으로 시작한다.
   ❌ 잘못: "{sub}\\n\\n그날 새벽 두 시였습니다..." (소제목 반복)
   ✅ 올바름: "그날 새벽 두 시였습니다..." (바로 본문 시작)"""

                content = ask_ai(content_prompt, 0.7, ensure_quality=True)
                if content:
                    content = clean_content(content, subtopic=sub)  # 이모티콘/마크다운 제거 + 소제목 중복 제거
                    ch_data['subtopic_data'][sub]['content'] = content

            # ── 챕터 끝: 핵심 정리(key_points) 생성 ──
            try:
                chapter_text = ""
                for s in ch_data['subtopics']:
                    c = ch_data.get('subtopic_data', {}).get(s, {}).get('content', '')
                    if c:
                        chapter_text += f"\n[{s}]\n{c[:600]}\n"
                if chapter_text.strip():
                    summary_prompt = f"""다음은 책 한 챕터의 본문이다. 이 챕터에서 독자가 반드시 기억해야 할 핵심을 3~4개로 정리하라.

[챕터 제목]: {ch}
[본문 발췌]
{chapter_text[:3000]}

규칙:
- 각 항목은 한 문장(공백 포함 45자 이내), 자연스러운 한국어 평서문.
- 추상적 구호 금지. 본문에서 실제로 다룬 구체적 통찰·행동을 요약.
- 번호·기호·이모지 없이 문장만.

JSON만 출력: {{"points": ["...", "...", "..."]}}"""
                    sres = ask_ai(summary_prompt, 0.5)
                    sparsed = parse_json(sres) if sres else None
                    pts = (sparsed or {}).get('points', []) if sparsed else []
                    pts = [p.strip() for p in pts if isinstance(p, str) and p.strip()][:4]
                    if pts:
                        ch_data['key_points'] = pts
            except Exception:
                pass

        # 완료 처리
        st.session_state['interview_completed'] = True
        progress_placeholder.success("✅ 본문 생성 완료!")
        return True

    except Exception as e:
        progress_placeholder.error(f"오류 발생: {str(e)}")
        return False

def go_next():
    if st.session_state['current_page'] < 6:
        st.session_state['current_page'] += 1

def go_prev():
    if st.session_state['current_page'] > 0:
        st.session_state['current_page'] -= 1

def auto_generate_all(topic, progress_placeholder):
    """주제만 입력하면 목차+본문까지 자동 생성"""
    try:
        # 1. 타겟 자동 생성
        progress_placeholder.info("🎯 1/4 타겟 분석 중...")
        target_result = suggest_targets(topic)
        targets = parse_json(target_result)
        if targets and targets.get('targets'):
            first_target = targets['targets'][0]
            persona = f"{first_target.get('name', '')} - {first_target.get('description', '')}"
            st.session_state['target_persona'] = persona

            # 페인포인트 분석
            pain_result = analyze_pains_deep(topic, persona)
            pain_data = parse_json(pain_result)
            if pain_data:
                st.session_state['pains'] = pain_data.get('pains', [])

        # 2. 목차 자동 생성
        progress_placeholder.info("📋 2/4 목차 생성 중...")
        outline_result = generate_outline(
            topic,
            st.session_state.get('target_persona', ''),
            st.session_state.get('pains', [])
        )

        # 목차 텍스트 파싱 (PAGE 4와 동일한 방식)
        if outline_result:
            chapters = []
            subtopics = {}
            current_ch = None

            for line in outline_result.split('\n'):
                orig_line = line
                line = line.strip()
                if not line:
                    continue

                # 챕터 감지 (PART, 1., 2. 등)
                is_chapter = False
                ch_name = None

                if re.match(r'^(PART|파트)\s*\d+', line, re.IGNORECASE):
                    is_chapter = True
                    ch_name = line
                elif re.match(r'^\d+[\.\)]\s', line):
                    is_chapter = True
                    ch_name = line
                elif re.match(r'^#+\s*(PART|파트|\d+)', line, re.IGNORECASE):
                    is_chapter = True
                    ch_name = re.sub(r'^#+\s*', '', line)

                if is_chapter and ch_name:
                    ch_name = ch_name.replace('**', '').replace('*', '').replace('#', '').strip()
                    if ch_name and len(ch_name) > 3:
                        current_ch = ch_name
                        if current_ch not in chapters:
                            chapters.append(current_ch)
                            subtopics[current_ch] = []
                elif current_ch:
                    # 소제목 감지
                    is_subtopic = False
                    st_name = None

                    if line.startswith('-') or line.startswith('•') or line.startswith('·'):
                        is_subtopic = True
                        st_name = line.strip().lstrip('-•· ')
                    elif re.match(r'^\s+[\da-z][\)\.]', orig_line):
                        is_subtopic = True
                        st_name = re.sub(r'^[\s\da-z\)\.\-]+', '', line).strip()

                    if is_subtopic and st_name:
                        st_name = st_name.replace('**', '').replace('*', '').replace('#', '').strip()
                        if st_name and len(st_name) > 3 and not re.match(r'^(PART|파트|Chapter|챕터)', st_name, re.IGNORECASE):
                            subtopics[current_ch].append(st_name)

            if chapters:
                st.session_state['outline'] = chapters
                st.session_state['chapters'] = {}
                for ch in chapters:
                    st.session_state['chapters'][ch] = {
                        'subtopics': subtopics.get(ch, []),
                        'subtopic_data': {s: {'questions': [], 'answers': [], 'content': ''} for s in subtopics.get(ch, [])}
                    }

        # 3. 본문 자동 생성
        progress_placeholder.info("✍️ 3/4 본문 작성 중...")
        if st.session_state.get('outline') and st.session_state.get('chapters'):
            total_subtopics = sum(len(st.session_state['chapters'][ch]['subtopics']) for ch in st.session_state['outline'])
            done = 0

            for ch in st.session_state['outline']:
                ch_data = st.session_state['chapters'][ch]
                for sub in ch_data['subtopics']:
                    done += 1
                    progress_placeholder.info(f"✍️ 본문 작성 중... ({done}/{total_subtopics})")

                    content = generate_content_premium(sub, ch, [], [], topic, st.session_state.get('target_persona', ''))
                    if content:
                        ch_data['subtopic_data'][sub]['content'] = content
                        ch_data['subtopic_data'][sub]['formatted'] = format_content_html(content)

        # 4. 완료
        progress_placeholder.success("✅ 완료! 본문 페이지로 이동합니다...")
        return True

    except Exception as e:
        progress_placeholder.error(f"오류 발생: {str(e)}")
        return False


# ==========================================
# AI 함수들
# ==========================================
def analyze_market_deep(topic):
    prompt = f"""주제: {topic}

이 주제로 전자책 시장을 분석해주세요.

[중요] 모든 답변은 반드시 한국어로만 작성하세요.

JSON:
{{
    "verdict": "강력 추천/추천/보류/비추천 중 하나",
    "verdict_reason": "판정 이유 한국어로",
    "total_score": 85,
    "search_data": {{
        "naver_monthly": "네이버 월간 검색량 예시: 12,000회",
        "google_monthly": "구글 월간 검색량 예시: 8,500회",
        "naver_blog_posts": "블로그 게시물 수",
        "youtube_videos": "유튜브 영상 수",
        "search_trend": "상승 또는 유지 또는 하락"
    }},
    "market_size": {{
        "score": 85,
        "level": "매우 큼/큼/보통/작음 중 하나",
        "analysis": "분석 2문장 한국어로"
    }},
    "competition": {{
        "score": 70,
        "level": "치열함/보통/낮음 중 하나",
        "your_opportunity": "차별화 기회 한국어로"
    }},
    "profit": {{
        "score": 80,
        "price_range": "권장 가격대",
        "monthly_revenue": "예상 월 수익"
    }},
    "popular_ebooks": [
        {{
            "title": "이 주제 관련 인기 전자책 제목",
            "platform": "크몽/탈잉/클래스101/리디북스/yes24 중 하나",
            "url": "해당 전자책 실제 URL (예: https://kmong.com/xxx)",
            "price": "가격"
        }},
        {{
            "title": "두번째 인기 전자책",
            "platform": "플랫폼명",
            "url": "URL",
            "price": "가격"
        }},
        {{
            "title": "세번째 인기 전자책",
            "platform": "플랫폼명",
            "url": "URL",
            "price": "가격"
        }}
    ],
    "recommendation": "최종 권장 2문장 한국어로"
}}"""
    return ask_ai(prompt, 0.5)


def suggest_targets(topic):
    prompt = f"""주제: {topic}

이 주제의 전자책을 가장 많이 구매할 것 같은 핵심 타겟 3개만 추천해주세요.
가장 적합하고 구매 가능성이 높은 타겟만 엄선해서 3개만 알려주세요.

[중요] 모든 답변은 반드시 한국어로만 작성하세요.

JSON:
{{
    "personas": [
        {{
            "name": "타겟 이름 (구체적으로)",
            "demographics": "연령대, 직업",
            "needs": "이 타겟이 이 책을 사는 이유",
            "pain_points": ["핵심 고민1", "고민2", "고민3", "고민4", "고민5"]
        }}
    ]
}}"""
    return ask_ai(prompt, 0.7)


def analyze_pains_deep(topic, persona):
    prompt = f"""주제: {topic}
타겟: {persona}

이 타겟의 고민을 아주 깊이 분석해주세요.

[중요] 모든 답변은 반드시 한국어로만 작성하세요. 외국어 사용 금지.

JSON:
{{
    "surface_pains": {{
        "pains": ["표면적 고민1", "고민2", "고민3", "고민4", "고민5"],
        "description": "표면적 고민 설명 3문장"
    }},
    "hidden_pains": {{
        "pains": ["숨겨진 진짜 고민1", "고민2", "고민3", "고민4"],
        "description": "숨겨진 고민 설명 3문장"
    }},
    "emotional_pains": {{
        "pains": ["감정적 고통1", "고통2", "고통3"],
        "description": "감정적 고통 설명 2문장"
    }},
    "failed_attempts": {{
        "attempts": ["시도했지만 실패한 것1", "것2", "것3"],
        "why_failed": "실패 이유 2문장"
    }},
    "dream_outcome": {{
        "ideal_result": "이상적인 결과",
        "timeline": "원하는 기간",
        "what_changes": "달라지는 것 2문장"
    }},
    "buying_triggers": {{
        "triggers": ["구매 요인1", "요인2", "요인3"],
        "objections": ["망설임 이유1", "이유2"]
    }},
    "marketing_hook": "마케팅 훅 한 문장"
}}"""
    return ask_ai(prompt, 0.6)


def analyze_competitor_reviews(topic):
    prompt = f"""주제: {topic}

이 주제 관련 전자책/도서의 부정적 리뷰를 분석해주세요.

[매우 중요]
- 모든 답변은 반드시 한국어로만 작성하세요.
- 영어, 러시아어 등 외국어 절대 사용 금지
- 한글과 숫자만 사용하세요.

JSON:
{{
    "analysis_scope": {{
        "books_analyzed": "287권",
        "reviews_analyzed": "3,842개",
        "negative_reviews": "892개 (23%)",
        "platforms": ["크몽", "예스24", "알라딘", "교보문고"]
    }},
    "negative_patterns": [
        {{
            "pattern": "불만 패턴 한국어로",
            "frequency": "67%",
            "example_reviews": ["실제 리뷰 예시 한국어로", "리뷰2"],
            "reader_emotion": "독자 감정 한국어로",
            "hidden_need": "숨겨진 니즈 한국어로",
            "solution": "해결책 한국어로"
        }},
        {{
            "pattern": "두 번째 불만",
            "frequency": "54%",
            "example_reviews": ["리뷰1", "리뷰2"],
            "reader_emotion": "감정",
            "hidden_need": "니즈",
            "solution": "해결책"
        }},
        {{
            "pattern": "세 번째 불만",
            "frequency": "41%",
            "example_reviews": ["리뷰1", "리뷰2"],
            "reader_emotion": "감정",
            "hidden_need": "니즈",
            "solution": "해결책"
        }}
    ],
    "hidden_needs_summary": {{
        "needs": ["숨겨진 니즈1", "니즈2", "니즈3"],
        "insight": "핵심 인사이트 2문장"
    }},
    "concept_suggestions": [
        {{
            "concept": "차별화 컨셉1 한국어로",
            "why_works": "이유 한국어로",
            "unique_point": "차별점 한국어로"
        }},
        {{
            "concept": "컨셉2",
            "why_works": "이유",
            "unique_point": "차별점"
        }}
    ],
    "success_formula": {{
        "must_have": ["필수1", "필수2", "필수3"],
        "must_avoid": ["금지1", "금지2"],
        "differentiation": "차별화 전략 한국어로 2문장"
    }}
}}"""
    return ask_ai(prompt, 0.6)


def generate_titles_bestseller(topic, persona, pains):
    prompt = f"""당신은 교보문고 종이책 베스트셀러와 크몽·클래스101 전자책 베스트셀러를 동시에 분석하는 제목 카피라이터입니다.

주제: {topic}
독자: {persona}
독자의 고민: {pains}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
가장 중요한 규칙 (이걸 어기면 전부 실패)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
제목은 반드시 "실제로 말이 되는 자연스러운 한국어 한 구절"이어야 합니다.
주제에서 뽑은 단어들을 기계적으로 이어 붙이지 마세요.

❌ 단어 짜깁기 실패 (절대 이렇게 만들지 말 것):
  "수면 매출 설계도"  → 수면+매출+설계도, 서로 관계없는 명사를 그냥 붙인 말
  "다이어트 부자 공식"  → 의미가 안 통하는 조합
  "관계 성장 엔진"      → 추상명사만 나열
  이런 제목은 소리내어 읽으면 "이게 무슨 말이지?" 싶고, 어떤 서점에도 존재하지 않습니다.

✅ 자연스러운 제목 = 사람이 실제로 쓰는 어순과 의미가 살아있는 구절
  (수면 주제 예) "잠든 사이 일어나는 일" / "초저녁의 기술" / "다시 잠드는 법"

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
두 시장의 제목 문법
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
[종이책 베스트셀러 — 품격·여운]
  돈의 속성 / 불변의 법칙 / 마흔에 읽는 쇼펜하우어 / 역행자 / 세이노의 가르침
  → 명사+의+명사, 한 단어 임팩트, 약간의 문학성. 광고 냄새가 없다.

[전자책 베스트셀러 — 구체적 약속·호기심]
  크몽·클래스101 상위권은 "독자가 무엇을 얻는지"가 살짝 드러난다.
  단, 유치하지 않게. "~하는 법/방법/노하우/비법/공식" 같은 흔한 꼬리표는 피하고,
  의외의 단어나 시선 전환을 한 번 넣어 호기심을 만든다.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
피해야 할 것
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
- 의미 없는 명사 나열 (가장 큰 실패 원인)
- 비밀/비법/공식/바이블/마법/머니/시스템/파이프라인/연금술
- 완벽한·궁극의·최고의·기적의 같은 과장 형용사
- "월 1000만원" 류 숫자 과시, "직장인을 위한" 류 타겟 명시

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
출력
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
'{topic}' 주제로 서로 다른 결의 제목을 정확히 5개.
(3개는 종이책 품격 톤, 2개는 전자책 호기심 톤)
길이는 자유 — 보통 2~7어절. 짧게 만드는 것보다 자연스러움이 우선이다.

각 제목마다 self_check를 채워, 소리내어 읽었을 때 말이 되는지 스스로 검증할 것.
self_check가 "어색하다/말이 안 된다"면 그 제목은 버리고 다시 만들 것.

JSON만 출력:
{{
    "titles": [
        {{"title": "제목", "subtitle": "호기심을 더하는 부제 한 줄 (20자 이내)", "concept": "이 제목이 매력적인 이유 한 줄", "self_check": "소리내어 읽으면 자연스러운가에 대한 한 문장 자기검증"}}
    ]
}}"""
    return ask_ai(prompt, 0.6, ensure_quality=True)


def analyze_text_content(text, source=""):
    prompt = f"""출처: {source}
내용: {text[:5000]}

분석:

JSON:
{{
    "title": "주제",
    "key_points": ["핵심1", "핵심2", "핵심3", "핵심4", "핵심5"],
    "insights": ["인사이트1", "인사이트2", "인사이트3"],
    "action_items": ["실행1", "실행2", "실행3"],
    "ebook_ideas": ["아이디어1", "아이디어2"],
    "summary": "요약 3문장"
}}"""
    return ask_ai(prompt, 0.5)


def summarize_all_knowledge(items, topic):
    """전체 학습 내용 통합 요약"""
    all_points = []
    all_tips = []
    all_ideas = []

    for item in items:
        if isinstance(item, dict):
            all_points.extend(item.get('key_points', []))
            all_tips.extend(item.get('actionable_tips', item.get('action_items', [])))
            all_ideas.extend(item.get('ebook_applications', item.get('ebook_ideas', [])))

    prompt = f"""전자책 주제: {topic}

학습한 모든 정보를 통합 분석해주세요.

수집된 핵심 포인트들:
{chr(10).join([f"• {p}" for p in all_points[:25]])}

실행 팁들:
{chr(10).join([f"• {t}" for t in all_tips[:15]])}

전자책 활용 아이디어:
{chr(10).join([f"• {i}" for i in all_ideas[:10]])}

JSON:
{{
    "integrated_summary": "전체 학습 내용 통합 요약 5문장",
    "core_insights": [
        "핵심 인사이트 1",
        "인사이트 2",
        "인사이트 3",
        "인사이트 4",
        "인사이트 5"
    ],
    "action_plan": [
        "즉시 실행할 것 1",
        "실행 2",
        "실행 3"
    ],
    "ebook_structure": [
        "추천 목차 1장",
        "2장",
        "3장",
        "4장"
    ],
    "unique_angle": "이 전자책만의 차별화된 관점",
    "study_plan": {{
        "week1": "1주차: 무엇을 할지",
        "week2": "2주차: 무엇을 할지",
        "week3": "3주차: 무엇을 할지",
        "week4": "4주차: 무엇을 할지"
    }},
    "expert_tips": [
        "전문가 팁 1",
        "팁 2",
        "팁 3"
    ]
}}"""
    return ask_ai(prompt, 0.6)


def generate_outline(topic, persona, pains, gaps=None):
    """한국 자기계발 베스트셀러 톤: 결제 버튼을 누르게 하는 자극형 목차"""

    # 페르소나/고통/시장 빈틈 정리 (있으면 활용, 없어도 작동)
    persona_block = f"[타겟 독자]\n{persona}\n" if persona else ""
    if isinstance(pains, list):
        pains_text = "\n".join(f"- {p}" for p in pains if p)
    else:
        pains_text = str(pains) if pains else ""
    pains_block = f"[독자가 지금 느끼는 통증]\n{pains_text}\n" if pains_text else ""
    if gaps:
        gaps_text = "\n".join(f"- {g}" for g in gaps) if isinstance(gaps, list) else str(gaps)
        gaps_block = f"[시장의 빈틈 - 경쟁자가 안 다루는 것]\n{gaps_text}\n"
    else:
        gaps_block = ""

    # 같은 책이면 기존 영문 기법 이름 재사용, 없으면 새로(고유) 생성
    method_name = get_or_create_method_name(topic, None, force_new=False)
    method_expansion = st.session_state.get('method_expansion', '')

    # 저자가 입력한 구체적 노하우 (세션에서 가져오기)
    _iv = st.session_state.get('interview_data') or st.session_state.get('temp_interview') or {}
    _knowhow = (_iv.get('detailed_knowhow', '') or '').strip()
    knowhow_block = (
        f"\n[저자의 구체적 노하우·실전 디테일 (★ 목차에 자연스럽게 녹여라)]\n{_knowhow}\n"
        "→ 위 노하우의 구체적 수치·단계·실수·기준·도구를 소제목에 녹여 저자만의 디테일이 느껴지게 하되, 그대로 베끼지 말고 호기심 유발 제목으로 재가공한다.\n"
    ) if _knowhow else ""

    prompt = f"""당신은 한국 자기계발 분야 톱 0.1% 기획자입니다. 서점에서 단 5초간 목차만 본 사람이 책을 손에서 못 놓게 만드는 5장짜리 목차를 씁니다.

목차의 단 하나의 목적: 독자가 "이 책을 안 읽으면 평생 손해"라고 느끼게 만드는 것.
정보 전달은 본문이 한다. 목차는 100% 구매심리만 다룬다.

{method_lock_rule(method_name)}

[주제]: {topic}
{persona_block}{pains_block}{gaps_block}{knowhow_block}
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🛒 구매 결정 5초 룰 (모든 규칙 중 1순위)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

독자는 목차를 5초만 본다. 그 5초 안에 다음 3가지가 동시에 작동해야 결제한다.

[1] 정체성 변화 발견 — "이걸 읽으면 나는 OO한 사람이 된다"
   → 5개 챕터 제목을 이어 읽으면 한 사람의 결정적 변화가 보여야 한다.
   ✅ "단단해진 멘탈은 인생을 통째로 바꾼다" (변화 서사 O)
   ❌ "멘탈 관리의 다양한 기법" (정체성 변화 X — 즉시 폐기)

[2] 손실회피 작동 — "이걸 모르면 평생 OO한다"
   → 통념 박살(인지부조화) 챕터/소제목이 최소 3개 들어가야 한다.
   ✅ "의지로 버틴 사람일수록 더 크게 무너진다"
   ❌ "멘탈 관리의 중요성" (잃을 게 안 보임)

[3] 구체성 — 추상 명사 1개당 구체적 숫자/장면 1개
   → "많은 사람" 금지, "월급 280만원짜리 7년차 회사원" 가능
   → 시간(90초, 47일), 금액(34만원, 1억), 비율(99%, 8할) 적극 사용

25줄(챕터 5 + 소제목 20) 중 한 줄이라도 "그냥 정보"가 섞이면 그 목차는 평이해진다.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🏆 골드 스탠다드 — 이 '결'을 그대로 복제하라 (규칙 100개보다 이게 먼저다)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

아래 둘은 '완성된 목차'의 본보기다. 한국 정보성 전자책(자청·프드프 류) 베스트셀러의 목차 톤 — 설명하지 않고, 한 줄마다 '장면 + 의문'을 던져 손가락이 결제 버튼으로 가게 만든다. 규칙을 다 지키는 것보다, 아래 목차를 처음 본 독자가 느끼는 '안 사면 손해 같은 긴장감'을 똑같이 재현하는 게 1순위다. 네가 쓴 목차를 아래와 나란히 놓았을 때 결이 조금이라도 약하거나 '설명문' 같으면 전부 다시 써라.

[본보기 A — 30대 직장인 부업 · 월 500]
PART 1. 직장인 99%가 부업 첫 달을 못 넘기고 포기하는 이유
- 퇴근 후 3시간을 갈아 넣고도 통장이 그대로인 진짜 이유
- 부지런한 사람일수록 더 빨리 나가떨어지는 역설
- "일단 시작하라"는 그 말이 당신을 가난하게 붙잡아 둔다
- 결국 모든 답은 '○○○' 안에 있었다
PART 2. 노력보다 '자리'가 먼저다
- 같은 시간을 써도 누구는 월 30, 누구는 월 500인 단 하나의 차이
- 잘 파는 사람들이 절대 입 밖에 안 내는 자리 잡는 순서
- 카톡 프로필 한 줄만 바꿨더니 문의가 들어오기 시작했다
- 남들 다 뛰어든 시장에서 6개월 만에 따라잡히는 사람들의 공통점
PART 3. 첫 30만원이 통장에 찍힌 그날
- 정확히 47일째, 잔고가 처음 늘어난 그 순간
- 27만원짜리 첫 정산서가 알려준 단 한 가지
- ○○○가 통장에 처음 흔적을 남기는 네 가지 신호
- 첫 수익을 본 사람 열에 아홉이 거기서 멈추는 함정
PART 4. 자는 동안에도 매출이 들어오기 시작한다
- 일하지 않는 시간에도 돈이 들어오게 만드는 구조의 정체
- 월 100을 넘긴 사람들이 가장 먼저 끊은 습관 한 가지
- 일을 덜 했더니 오히려 매출이 두 배가 된 이상한 달
- 한 번 자리잡은 사람은 다시 월급만 기다리지 않는다
PART 5. 월 500 다음, 억대로 넘어가는 길
- 월 500에서 멈추는 사람과 5천을 넘기는 사람의 단 한 줄 차이
- 돈과 시간이 동시에 따라오기 시작하는 전환점
- 1년 만에 사직서를 낸 사람들의 똑같았던 마지막 한 달
- '○○○'로 자리잡은 사람들의 5년 뒤가 완전히 다른 이유

[본보기 B — 멘탈 · 자기관리]
PART 1. 1년 안에 99%가 다시 무너지는 결정적 이유
- 마음 다잡기를 백 번 해도 사흘이면 원점으로 돌아가는 이유
- 의지가 강한 사람일수록 더 크게 무너지는 역설
- "긍정적으로 생각하라"는 조언이 당신을 더 망가뜨린다
- 여기서부터 진짜 이야기, '○○○'가 시작된다
PART 2. 의지로 잡으려는 순간 뇌는 반대로 움직인다
- 감정이 폭발할 때 뇌가 정확히 어디부터 멈추는가
- 참을수록 더 크게 터지는 사람들의 공통된 패턴
- 명상도 운동도 당신에게만 효과가 없었던 진짜 이유
- 무너진 다음 날 가장 먼저 해야 할 단 한 가지
PART 3. 회복의 8할이 결판나는 폭발 직후 90초
- 이 90초를 놓치면 다시 일어서는 데 며칠이 걸린다
- 단번에 회복하는 사람들이 무의식적으로 하는 행동
- ○○○를 처음 적용한 사람들이 가장 먼저 느낀 변화
- 감정에 휘둘리던 사람이 처음으로 멀쩡했던 그날
PART 4. 한 번 회복한 뇌는 다시는 무너지지 않는다
- 같은 일을 겪어도 더는 흔들리지 않게 되는 메커니즘
- 무너지는 횟수가 한 달 만에 절반으로 줄어드는 신호
- 단단해진 사람들이 더 이상 하지 않는 다섯 가지 행동
- 주변에서 "사람이 달라졌다"는 말을 듣기 시작하는 시점
PART 5. 단단해진 멘탈은 인생을 통째로 바꾼다
- 멘탈이 바뀌자 통장과 인간관계가 함께 달라진 이유
- 다시는 예전으로 돌아가지 않는 사람들의 결정적 차이
- 6개월 뒤 가장 크게 달라지는 단 한 가지
- '○○○'로 회복한 사람들의 5년 뒤가 완전히 다른 이유

[본보기 C — 주식·배당 투자]
PART 1. 10년을 모아도 배당으로 월 50만원 못 받는 사람들의 공통점
- 배당률 숫자만 보고 골랐더니 3년째 원금이 그대로인 이유
- 매달 꼬박꼬박 넣었는데도 통장이 안 불어난 8할의 패턴
- "분산투자"가 오히려 당신을 제자리에 묶어두는 방식
- 결국 모든 답은 '○○○' 안에 있었다
PART 2. 받은 배당을 쓰지 않은 사람만 3년 뒤 게임이 달랐다
- 같은 원금을 넣었는데 누구는 매달 80만원, 누구는 한 푼도 못 받는 이유
- 배당을 재투자한 계좌와 생활비로 쓴 계좌의 5년 뒤 잔고
- 증권사 앱을 하루 세 번 이상 열어본 해에 수익률이 가장 낮았다
- 남들이 팔 때 오히려 더 사 모은 계좌만 폭락장에서 살아남은 이유
PART 3. 처음으로 배당이 생활비를 넘어선 그달
- 재투자를 3년 유지한 계좌에서 배당이 두 배로 늘어난 순간
- ○○○가 처음으로 통장에 흔적을 남기는 네 가지 장면
- 원금보다 배당이 먼저 눈에 들어오기 시작하는 전환의 신호
- 처음 이 구조를 이해한 사람 열에 여덟이 바로 그때 저지르는 실수
PART 4. 쓰지 않아도 매달 들어오는 돈이 생긴 사람의 뇌는 다르다
- 근로소득만 바라보던 사람이 더 이상 월급날을 기다리지 않게 되는 때
- 배당 수령액이 월 100만원을 넘긴 사람들이 가장 먼저 바꾼 한 가지
- 계좌를 덜 들여다볼수록 잔고가 더 빠르게 불어난 이상한 메커니즘
- 직장을 그만두지 않아도 퇴직 준비가 끝나는 정확한 시점
PART 5. 배당이 월급을 넘어선 날, 인생의 선택지가 달라졌다
- 노동 없이 들어오는 돈이 생기자 직장에서 더 당당해진 이유
- 5년 뒤 가장 크게 갈리는 건 투자 금액이 아니라 이것 하나
- 배당으로 퇴직한 사람들이 다시 월급쟁이로 돌아가지 않는 구조
- '○○○'를 3년 이상 유지한 사람들의 10년 뒤가 완전히 다른 이유

[본보기 D — 직장인 퇴사·독립]
PART 1. 퇴사를 꿈꾸다 40대가 된 직장인의 정확한 실수
- 퇴사 준비를 3년째 하고 있는데 통장이 그대로인 진짜 이유
- "조건이 갖춰지면 나간다"는 그 말이 당신을 평생 붙잡는다
- 먼저 퇴사한 동기가 6개월 만에 다시 돌아온 날 배운 것
- 여기서부터 진짜 이야기, '○○○'가 시작된다
PART 2. 월급이 없어도 버티는 사람과 석 달 만에 무너지는 사람
- 퇴사 첫 달보다 여섯 번째 달이 더 무서운 진짜 이유
- 돈이 아니라 이것 하나가 먼저 무너진 사람이 결국 돌아간다
- "나만의 일"을 찾느라 1년을 날린 사람들의 공통된 함정
- 퇴사 후 6개월 살아남은 사람들이 처음 한 달에 한 일
PART 3. 월 200이 통장에 처음 찍힌 날, 두려움이 사라졌다
- 정확히 73일째, 첫 입금 알림이 울린 그 새벽
- ○○○가 처음으로 작동한다는 걸 알아챈 단 하나의 신호
- 첫 수익을 본 사람이 바로 그 순간 가장 많이 저지르는 실수
- 이 시점을 넘긴 사람과 못 넘긴 사람의 1년 뒤 통장
PART 4. 다시는 월급날을 기다리지 않는 사람이 된다
- 수입이 불규칙해도 흔들리지 않는 뇌가 만들어지는 조건
- 사업자 등록증 한 장이 삶의 선택지를 통째로 바꾼 이유
- 퇴사 1년 차에 가장 먼저 끊은 것과 가장 먼저 시작한 것
- 주변에서 "달라 보인다"는 말을 듣기 시작하는 정확한 시점
PART 5. 월 500을 넘긴 뒤, 그다음 게임은 완전히 다르다
- 돈을 좇던 사람이 돈이 따라오게 만드는 사람이 된 차이
- 1년 만에 사직서를 쓴 사람들의 마지막 3개월이 닮은 이유
- 지금도 퇴사를 못 한 사람과 해낸 사람 사이의 단 하나
- '○○○'로 자립한 사람들이 5년 뒤 가장 잘했다고 꼽은 결정

[본보기에서 훔쳐야 할 6가지]
1. 첫 줄(PART 1)부터 독자의 뜨끔한 현실을 정면으로 때린다 — "직장인 99%가 못 넘기는", "1년 안에 99%가 무너지는".
2. 통념을 뒤집어 "내 탓이 아니었다"는 안도와 충격을 동시에 준다 — "부지런한 사람일수록 더 빨리 나가떨어진다".
3. 결과·장면은 보여주고 방법·이유는 숨긴다 — "47일째 잔고가 늘어난 그 순간"(왜·어떻게는 본문).
4. 5개 PART를 세로로 읽으면 '바닥→깨달음→첫 사건→전환→도약'이 한 편의 이야기로 흐른다.
5. 한 줄도 "~의 방법", "~의 중요성" 같은 교과서 표현이 없다. 전부 사람·장면·숫자·의문이다.
6. ★ "너 지금 이러고 있지?" 직접 저격 — 자청 목차의 가장 강한 무기. 독자가 읽다가 "어? 나 얘기인데?"라고 멈춰야 결제한다. 소제목 20개 중 최소 5개는 독자의 현재 행동을 정확하게 묘사해서 들키게 만들어라. "퇴사 준비를 3년째 하고 있는데 통장이 그대로인 진짜 이유", "공부만 하다 1년을 통째로 날리는 사람들의 5가지 변명", "증권사 앱을 하루 세 번 이상 열어본 해에 수익률이 낮았다" — 독자가 어제 한 행동을 정확히 집어내는 줄이 목차에서 가장 세게 작동한다.

⚠️ 위 A·B·C·D는 '톤·구조의 본보기'다. 문장을 그대로 베끼지 말고, 주어진 [주제]·[저자 노하우]에 맞춰 같은 강도로 완전히 새로 써라. A·B의 문장 패턴을 그대로 옮기면 즉시 탈락이다.

[실전 실패 사례 — 실제로 이렇게 평이하게 나왔다. 같은 실수를 반복하지 마라]
✗ 두 PART가 사실상 같은 말이라 흐름이 안 보인 경우:
   PART 3. 첫 개선 후 매출이 두 배로 뛴 그날
   PART 4. 개선할수록 매출이 따라 올라온다   ← 둘 다 '개선하면 매출 오른다' = 같은 내용. 5막이 무너짐.
   ✓ 역할을 완전히 갈라라 — PART 3은 '첫 성과가 터진 장면', PART 4는 '자동으로 팔리는 사람이 됨':
   PART 3. 사십칠 번째 주문이 찍힌 새벽의 정산 알림
   PART 4. 자는 사이에 휴대폰이 울리는 판매자가 된다

✗ "리뷰 3개가 알려준 결정적 신호" → '결정적 신호'는 금지된 AI 클리셰 + 무슨 신호인지 안 보여 밍밍.
   ✓ "후기 한 줄이 본문 백 페이지를 이기는 이유" (장면 + 의문)
✗ "SRP 사이클이 본격 작동하기 시작하는 4가지 신호" → 시스템 의인화('사이클이 작동') + '신호' 반복.
   ✓ "○○○가 통장에 첫 흔적을 남기는 네 가지 장면"
✗ "매출과 시간이 동시에 따라오기 시작하는 전환점" → 추상적이라 장면이 없음.
   ✓ "일 년 만에 사표를 던진 사람들의 닮은 마지막 한 주"

★ 위 실패작이 평이했던 결정적 이유: 같은 말이 계속 반복됐다 — '전자책' 다섯 번, '개선' 네 번, '매출' 세 번, '가장 먼저'·'결정적'·'신호'·'한 번' 각 두 번, 숫자 3이 네 번. 이 반복이 목차를 가장 평이하게 만든다. 25줄에서 핵심 단어와 문장 끝은 전부 한 번씩만 써라.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🧠 마케팅 뇌과학 8대 트리거 (목차 전체에 골고루 박을 것)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

뇌가 읽기를 멈출 수 없게 만드는 메커니즘. 뇌는 (1) 예측이 깨질 때(예측오류) (2) 빈칸을 보면 채우려 할 때(정보격차) (3) 잃을 것이 보일 때(손실회피) 도파민이 분비돼 "더 알고 싶다"는 충동을 만든다. 아래 트리거는 전부 이 셋을 노린다.

1. **손해회피 + 충격 통계** — "1년 안에 99%가 다시 무너진다", "1억치 강의 들어도 망한다"
2. **인지부조화/통념 박살(예측오류)** — "의지로 버틴 사람일수록 더 크게 무너진다", "노력할수록 가난해진다"
3. **권위 어휘 (과학/임상)** — 뇌, 신경회로, N주 후, 임상, 데이터, 알고리즘 (절대 비유로 남용 금지, 사실 진술로만)
4. **임박감 + 절대성** — "이 90초를 놓치면 며칠 걸린다", "다시는 ~하지 않는다"
5. **정체성 전환 약속** — "회복한 뇌는 다시 무너지지 않는다", "단단해진 사람은 ~한다"
6. **인그룹 사회증명** — "○○를 익힌 사람들의 5년 뒤", "상위 1%만 도달하는"
7. **이중/삼중 보상** — "통장과 인간관계가 함께 변한다", "돈도 사람도 따라온다"
8. **호기심 갭(정보격차)** — "두 달 안에 가장 먼저 끊은 한 가지", "정확히 어디부터 멈추는가"

※ 가장 강한 줄은 위 트리거를 '겹쳐' 쓴다. 예) "10명 중 8명이(사회증명) 배당주를 사도 월 100만원조차 못 받는(손실회피+구체금액) 이유(정보격차)" — 한 줄에 트리거 2~3개가 겹칠수록 결제 충동이 폭발한다.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🧲 호기심 갭(Information Gap) 강화 — 결제를 부르는 가장 강력한 무기
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

조지 로웬스타인의 정보격차 이론: 사람은 자기가 모르는 것이 '뭔지'는 알지만 '내용'은 모를 때 가장 강하게 끌린다. 25줄 중 최소 10줄에 이 갭을 넣어야 목차만 보고 결제한다.

[호기심 갭 6대 공식 - 결과만 보이고 방법/이유는 본문으로 미루기]

1. **숫자 + 미공개 결과** — '정확히/딱'으로 시작해 결과만 보여주고 내용은 숨기기
   ✅ "정확히 47일째, 통장 잔고가 처음으로 불어나기 시작한 순간"
   ✅ "27만원짜리 첫 정산서가 알려준 단 한 가지"
   ❌ "47일 만에 돈을 버는 방법" (방법을 다 보여줘버림 → 결제 안 함)

2. **이미 벌어진 사건 + 원인 숨김** — '왜?'를 유발하는 결과만
   ✅ "3년 차 베테랑이 신입에게 6개월 만에 따라잡힌 단 하나의 이유"
   ✅ "월 1,000을 찍은 사람들이 가장 먼저 끊은 습관 한 가지"

3. **묘하게 구체적인 행동/대상 + 이유 숨김** — 디테일이 호기심을 폭발시킨다
   ✅ "성공한 부업러가 매일 밤 11시에 반드시 끄는 것"
   ✅ "1년 만에 1억 모은 사람들이 절대 안 쓰는 5단어"
   ✅ "월 500 넘긴 사람들 카톡 프로필에서 사라진 한 단어"

4. **반대 결과 미스터리** — 통념과 정반대 결과만 던지고 메커니즘은 본문
   ✅ "더 열심히 할수록 더 가난해진 7년의 비밀"
   ✅ "잠을 늘렸더니 매출이 2배가 된 이상한 메커니즘"

5. **'딱 하나' 절대성** — 수많은 변수 중 단 하나만 보여주기
   ✅ "월 100 / 월 500을 가르는 단 한 줄의 차이"
   ✅ "결국 모든 게 무너지는 사람들의 공통점 단 하나"

6. **시간 압축 미스터리** — 짧은 시간에 큰 일이 일어났는데 그 사이를 숨기기
   ✅ "퇴근 후 90분이 1년 뒤 인생을 갈라놓는다"
   ✅ "주말 4시간이 5년치 월급을 바꾼 그 과정"

[호기심 갭 만들 때 절대 어기지 말 것]
• 답을 같은 줄에 다 보여주지 마라. "왜 ~한가" "어떻게 ~하는가"로 끝나면 본문을 사야 알 수 있게.
• "방법" "비법" "노하우" 같은 말로 끝내면 갭이 닫힘 → 결제 안 함.
• 실행 순서와 디테일(종목·비율·타이밍)은 숨기되, 작동 원리의 '큰 줄기'는 한 조각 드러내라 — 원리까지 다 숨기면 '근거 없는 주장'이 되어 안 팔린다(아래 ⚖️ 섹션 참고).
• 한 줄 안에 "장면 + 의문"이 같이 있어야 호기심이 작동한다.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
⚖️ 호기심 갭의 함정 — '주장만 있고 근거가 없으면' 안 팔린다 (★ 매우 중요)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

방금까지 '방법을 숨기라'고 했다. 하지만 모든 줄을 '결과 + 미스터리'로만 채우면 목차가 '근거 없는 주장 모음'이 되어 매력이 죽는다. 독자는 "오, 이렇게 하니까 되는 거구나"라는 믿음(=작동 원리 감)이 있어야 결제한다. 지금 가장 흔한 실패가 바로 이것 — "달라진다 / 완성된다 / 무너지지 않는다 / 쌓이기 시작한다" 같은 결과 선언만 스무 줄 반복해서 공허해지는 것.

[숨길 것 vs 보여줄 것 — 이 구분이 전부다]
• 숨겨라(본문에서) = 실행 디테일: 정확한 종목·비율·금액·타이밍·순서·도구 설정. "이대로 따라 하면 된다" 수준은 절대 노출 금지.
• 보여줘라(목차에서) = '지렛대 한 조각 + 그게 왜 되는지 원리'. 독자가 큰 흐름은 감 잡되, 혼자 실행하진 못하는 수준.

[나쁜 예 — 주장만, 근거 0 (실제로 이렇게 나왔다)]
✗ "배당금이 통장에 찍히는 순간부터 게임이 완전히 달라진다" → 뭐가 어떤 원리로 달라지는지 지렛대가 안 보임. 게다가 '게임이 달라진다'는 AI 클리셰.
✗ "단단해진 포트폴리오는 폭락장에서도 무너지지 않는다" → 왜 안 무너지는지 근거 0. 공허한 선언.
✗ "현금흐름과 자산 가치가 동시에 따라오기 시작하는 전환점" → 추상적 주장, 작동 원리 없음.
✗ "한 번 흐름을 만든 사람은 다시 원금만 바라보지 않는다" → 무슨 '흐름'인지 레버가 안 보임.

[좋은 예 — 지렛대·원리는 보여주고 디테일은 숨김]
✓ "받은 배당을 쓰지 않고 그대로 다시 사 모은 사람만, 같은 원금으로 더 많은 배당을 받게 된 이유" (지렛대=재투자, 원리=복리가 보임 / 얼마를·어느 종목은 숨김)
✓ "폭락장에 팔지 않고 오히려 더 사 모은 계좌만 이듬해 가장 빨리 회복한 이유" (행동=급락 매수, 인과의 큰 흐름 보임 / 비중·시점은 숨김)
✓ "월급은 일해야 들어오고 배당은 가만히 있어도 들어온다 — 이 둘을 갈라놓은 사람만 자산이 따로 쌓인다" (구조의 논리가 보임)
✓ "수익을 빼 쓰지 않고 3년 내내 재투자만 한 사람이 원금을 두 배로 불린 까닭" (지렛대=재투자 누적, 결과의 근거 보임)

[규칙 — 이 강도로 써라]
- 소제목 20개 중 최소 12개는 '구체적 지렛대(행동·원리) 한 조각'을 드러내라. 막연한 '달라진다 / 완성된다 / 무너지지 않는다 / 따라온다'로 끝나는 순수 주장 줄은 PART당 1개 이내로 제한.
- 20개를 다 읽으면 이 방법의 '큰 골격(무엇을 → 왜 → 대략 어떻게)'이 손에 잡혀야 한다. 단, 그 골격을 '혼자 실행할 수 있을 만큼' 구체적으로 주지는 마라(그건 본문의 몫).
- 컨셉명 약자(예: DIME)의 각 단계가 PART 1~5 흐름에 자연스럽게 드러나, 독자가 "이런 순서·논리로 가는 방법이구나"를 목차만 보고 감 잡게 하라.
- 즉, 호기심(디테일 숨김)과 신뢰(원리 노출)를 한 줄 안에 같이 넣는다 — "[구체적 지렛대]로 [결과]가 일어난 [까닭/순간]"이 가장 강한 공식이다.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🎯 단 하나의 미션 — 목차가 곧 '예고편 스토리'다
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

5개 PART 제목은 독립된 5개의 광고 문구가 아니다. 이어 읽으면 한 사람(독자 자신)의 추락→반전→도약을 그리는 '한 편의 예고편'이어야 한다.
목차를 5초 훑은 사람의 머릿속에 "어? 이거 딱 내 얘기인데, 그래서 마지막에 어떻게 됐는데?"라는 한 문장이 떠올라야 결제한다.

[★ 스토리 흐름 강제 규칙 — 가장 자주 어기는 부분, 반드시 지켜라]
PART 1~5를 '제목만' 세로로 이어 읽었을 때, 끊긴 5개 슬로건이 아니라 시간 순서대로 이어지는 한 사람의 이야기로 읽혀야 한다.

(1) 같은 주인공 — 5개 PART 내내 '같은 한 사람'의 시점이 유지돼야 한다. PART마다 화자가 바뀌면 흐름이 깨진다.
(2) 시간의 화살 — PART가 넘어갈수록 시점이 '과거→현재→미래'로 전진해야 한다. 뒤로 돌아가거나 제자리걸음 금지.
    PART 1 (지금, 바닥) → PART 2 (왜 그랬는지 깨달음) → PART 3 (처음 바뀐 그 사건) → PART 4 (이제 다른 사람이 됨) → PART 5 (1년 뒤, 완전히 다른 인생)
(3) 감정 에스컬레이션 — PART 1은 가장 아프게(좌절·손실), PART 5는 가장 벅차게(도약·보상). PART가 올라갈수록 잃을 공포 또는 얻을 흥분이 한 단계씩 더 커져야 한다. 5개 제목의 '온도'가 평평하면 실패다.
(4) 갈고리 연결 — 각 PART 제목은 바로 다음 PART를 안 보고는 못 배기게 끝나야 한다. PART 1을 읽으면 "그럼 어떻게?"가, PART 3을 읽으면 "그 다음엔?"이 자동으로 떠오르게.

[이어 읽기 예시 — 이렇게 '한 편'으로 읽혀야 한다]
PART 1. 1년 안에 99%가 다시 무너지는 결정적 이유      ← (바닥: 너도 곧 무너진다)
PART 2. 의지로 잡으려는 순간 뇌는 반대로 움직인다       ← (깨달음: 그게 네 탓이 아니었다)
PART 3. 회복의 8할이 결판나는 폭발 직후 90초          ← (사건: 바로 이 순간이 갈림길)
PART 4. 한 번 회복한 뇌는 다시는 무너지지 않는다        ← (전환: 이제 너는 다른 사람)
PART 5. 단단해진 멘탈은 인생을 통째로 바꾼다           ← (도약: 1년 뒤 완전히 다른 인생)
→ 5줄을 이어 읽으면 "무너지던 내가 → 원인을 알고 → 결정적 순간을 잡아 → 다른 사람이 되고 → 인생이 바뀐다"는 한 편의 이야기가 또렷이 보인다. 이 '세로 흐름'이 안 보이면 제목이 아무리 자극적이어도 다시 써라.

한 줄 한 줄은 다음 챕터를 못 참게, 5줄 전체는 사지 않으면 결말이 궁금해 미치게 만들어야 한다.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📐 챕터 제목 형식 (가장 중요)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[형식]
- 9~18자의 자연스러운 한국어 한 문장. 길게 풀어 설명하지 말고 짧고 단단하게 쳐라("노력보다 자리가 먼저다"처럼 짧고 센 게 더 강하다). 단, 짧다고 라벨·단어 나열이 되면 안 됨 — 끝까지 '말이 되는 한 문장'
- 단어 한두 개짜리 라벨 금지, 라벨 분리(— ㅣ :) 금지
- 평서문 또는 단언형. "~이유"로 끝나는 설명체는 한 PART에만 사용
- 명사 엔딩과 동사 엔딩을 챕터별로 섞어라 (5개 모두 명사 엔딩 금지)
- 5개를 이으면 [좌절 → 통념 박살 → 첫 사건 → 안정화 → 도약]의 5막이 '시간 순서'로 또렷이 보여야 한다 (위 🎯 스토리 흐름 강제 규칙 준수)
  · PART 1 = 가장 아픈 바닥(잃을 것), PART 5 = 가장 벅찬 도약(얻을 것). 1→5로 갈수록 감정 온도가 한 단계씩 올라가야 한다. 5개 제목의 세기가 평평하면 실패.
  · 각 PART 제목 끝은 다음 PART를 안 보면 못 배기게(갈고리). 5개가 따로 노는 독립 슬로건이면 즉시 다시 써라.

[좋은 예 - 성공적인 멘탈 관리 비결]
PART 1. 1년 안에 99%가 다시 무너지는 결정적 이유
PART 2. 의지로 잡으려는 순간 뇌는 반대로 움직인다
PART 3. 회복의 8할이 결판나는 폭발 직후 90초
PART 4. 한 번 회복한 뇌는 다시는 무너지지 않는다
PART 5. 단단해진 멘탈은 인생을 통째로 바꾼다

[좋은 예 - 30대 직장인 N잡 월 500]
PART 1. 직장인 99%가 부업 30일을 못 버티는 이유
PART 2. 노력보다 자리가 먼저다
PART 3. 첫 30만원이 통장에 찍힌 그날
PART 4. 새벽 3시에도 매출이 들어온다
PART 5. 월 500 다음, 억대 수익으로 가는 길

[나쁜 예 - 즉시 폐기]
- "발굴", "폭로", "전환" 같은 한두 단어짜리
- "발굴 — 부업의 90%는 첫 단추에서 망한다" (라벨 + 대시)
- "DPS의 첫 관문, 노력 없이도 돈이 따라오는 자리를 찾는 법" (시스템명 라벨화 + 너무 김)
- "이제 무너지는 게 더 이상 사건이 아니다" (사건이 아니다 ← 말이 안 됨)
- "한 번 흔들려도 다음 날엔 흔적도 없다" (AI식 과장)
- "월수도 시스템의 첫 설계" (의미 불명 + 설계라는 설명체 어휘)
- "MDS 파이프라인" (파이프라인이라는 영어 외래어를 시스템 접미사로)
- "주가 -12% 떨어져도 매도 안 하는 뇌 회로가 박혔다" (뇌 회로가 박히다 = 어법 어색, 비유 남용)
- "신경회로가 새로 깔리고 있다는 신호" 류 (한 번까진 OK, 같은 비유 두 번 X)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
💎 책의 시그니처 기법 (위 고정 규칙을 그대로 따른다)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

기법 이름은 위에서 "{method_name}"(영문 약자+한글 접미사)으로 이미 확정됨. 새로 짓지 말 것.
- 한글로 바꾸지 말고, 다른 표기로 바꾸지 말고, 그대로 사용한다.
- 아래 등장 규칙의 ○○○ 자리에는 전부 "{method_name}" 을 넣는다.
- "{method_name}" 을 글자 그대로 쓰고, 약자를 풀어쓰거나 접미사를 덧붙이지 않는다.

[약자의 의미를 목차에 '간접적으로' 녹여라]
이 기법의 약자 풀이: {method_expansion if method_expansion else "(풀이 없음 — 책 흐름으로 자연스럽게 암시)"}
- 약자가 의미 있는 머리글자라면, 각 단계가 PART 1~5 흐름이나 소제목에 자연스럽게 드러나 독자가 목차만 봐도 "이런 단계의 방법론이구나"를 감 잡게 하라.
- "Y: ~" 식 글자별 사전 나열·풀이 직접 받아쓰기 절대 금지. 호기심 자극 제목 안에 의미가 스며들게만.
- 풀이가 없으면 억지로 짜맞추지 말고 평소대로 호기심 중심으로 쓰되, 책 전체가 하나의 방법론을 단계적으로 다룬다는 느낌만 유지.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🎯 컨셉명 등장 규칙 - 3번, 자연스럽게 분산
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

이 컨셉명을 5개 PART 안에 정확히 3곳에 등장시킨다:

⚠️ 컨셉명 글자 수/단어 수를 단정하는 표현 절대 금지 (가장 자주 어기는 실수!)
   ❌ "○○○ 세 글자에서 시작한다" — 컨셉명이 3글자 아니면 거짓말이 됨
   ❌ "○○○ 네 글자만 기억하라" — 글자 수 단정 금지
   ❌ "단 한 단어가 모든 걸 바꾼다" — 단어 수 단정 금지
   ✅ 컨셉명이 'DPS'(3자)든 '단단한 매출 구조'(8자)든 '복리 자산 공식'(7자)이든 모두 자연스럽게 작동하는 문장만 사용

1. **PART 1 마지막 소제목** (도입) — 아래 풀에서 1개 선택, 컨셉명 글자 수와 무관하게 자연스러운 것
   - 예: "결국 모든 답은 '○○○' 안에 있었다"
   - 예: "이 책의 모든 페이지는 '○○○' 하나를 향해 간다"
   - 예: "지금부터 '○○○' 단 하나만 기억하면 된다"
   - 예: "여기서부터 진짜 이야기, '○○○'가 시작된다"
   - 예: "마지막에 도달하는 곳은 결국 '○○○'다"
   - 예: "'○○○'를 만나기 전과 후는 완전히 다른 게임이다"

2. **PART 3 또는 PART 4 안 (1곳)** (작동·전환점)
   - 예: "○○○가 본격 작동하기 시작하는 4가지 신호"
   - 예: "○○○를 처음 적용한 사람들이 가장 먼저 느낀 변화"
   - 예: "○○○가 통장에 처음 흔적을 남기는 순간"
   - 예: "○○○ 한 달 차에 가장 먼저 무너지는 한 가지"

3. **PART 5 마지막 소제목** (확장/사회증명)
   - 예: "○○○를 익힌 사람들의 5년 뒤가 완전히 다른 이유"
   - 예: "○○○로 자리잡은 사람들이 다시는 돌아가지 않는 이유"
   - 예: "○○○ 이후, 1년 만에 가장 크게 달라지는 단 한 가지"

❌ 챕터 제목에는 컨셉명 절대 등장 X (5개 챕터 제목엔 안 들어감)
❌ 매 PART에 박지 말 것. 정확히 3곳.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
✍️ 소제목 톤: 한국 자기계발 베스트셀러 + 마케팅 뇌과학
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

각 PART당 소제목 4개. 모두 다른 패턴 + 명사·동사 엔딩 섞기 (4개 모두 같은 엔딩 금지).

[좋은 패턴 풀 - 매번 다르게]
1. 통계 충격형 — "1년 안에 99%가 다시 무너진다", "한 달 차에 90%가 다시 무너지는 정확한 이유"
2. 인지부조화형 — "의지로 버틴 사람일수록 더 크게 무너진다"
3. 뇌과학 권위형 — "감정이 폭발할 때 뇌는 정확히 어디부터 멈추는가", "신경회로가 새로 깔리고 있다는 4가지 신호"
4. 임박감/손해회피형 — "이 90초를 놓치면 다시 일어서는 데 며칠이 걸린다"
5. 정체성 전환형 — "한 번 회복한 뇌는 다시는 무너지지 않는다"
6. 호기심 갭형 — "단번에 회복한 사람들이 모두 무의식적으로 하는 행동"
7. 이중/삼중 보상형 — "회복 후 6개월 만에 통장과 인간관계가 함께 변한다"
8. 인그룹 사회증명형 — "○○를 익힌 사람들의 5년 뒤가 완전히 다른 이유"

[★ 가장 중요 — 목차만 보고 결제하게 만들어라 (설명문 금지)]
이 목차의 단 하나의 목적: 서점에서 5초 훑은 사람이 "이건 사야 해"라고 결제하게 만드는 것.
정보 전달은 본문이 한다. 목차는 100% 구매 욕구만 자극한다.

지금 가장 흔한 실패 = "설명문처럼 평이함". 아래 셋 중 하나라도 걸리면 그 줄은 죽은 줄이니 다시 써라:
  (1) 정보 전달형("~하는 법", "~의 중요성", "~란 무엇인가") — 교과서 목차
  (2) 실행 디테일(종목·비율·타이밍·순서)까지 다 알려줌 — 본문 살 이유 사라짐 (단, 작동 '원리·지렛대'의 큰 흐름은 오히려 보여줘야 믿고 산다)
  (3) 자극이 없음 — 심장이 안 뛰면 결제 안 한다

[자청식 후킹 8가지 무기 — 소제목 20개 중 최소 14개에 적용 (밍밍한 정보형은 거의 없어야 함)]
1. 통념 정면 박살: "열심히 할수록 가난해지는 이유", "절약이 당신을 평생 가난하게 만든다"
2. 자기기만 적발(뜨끔하게): "공부만 하다 1년을 통째로 날리는 사람들의 5가지 변명"
3. 단정 선언(반박 불가 톤): "결국 전부 이거 하나에서 갈렸다", "답은 처음부터 정해져 있었다"
4. 날 선 대비: "버는 사람과 버는 척하는 사람", "3년 버틴 사람과 3개월에 접은 사람의 통장"
5. 충격 숫자: "10명 중 8명이 1년 안에 무너지는 진짜 이유", "상위 3%만 아는 한 가지"
6. 공포·손실 회피: "지금 이걸 모르면 5년 뒤 똑같은 자리에 있다", "당신이 놓치는 사이 벌어지는 일"
7. 미스터리 갭(정체를 가림): "퇴사 3개월 만에 다시 돌아온 회사원이 깨달은 단 하나"
8. ★ 독자 행동 직접 저격 (가장 강한 무기) — "너 지금 이러고 있지?" 톤으로 독자의 현재 행동·습관·변명을 정확하게 집어내라. 독자가 어제 한 행동을 보여서 "어? 내 얘기잖아"라고 멈춰야 결제한다.
   예: "퇴사 준비를 3년째 하고 있는데 통장은 그대로인 사람의 패턴"
   예: "증권사 앱을 하루 세 번 이상 열어본 해에 수익률이 가장 낮았다"
   예: "조건이 갖춰지면 시작하겠다고 한 사람이 5년째 같은 말을 반복하는 이유"
   예: "유튜브 강의 50개를 들었는데 통장은 그대로인 사람들의 공통점"
   ※ 소제목 20개 중 최소 5개는 이 무기를 써라. 독자의 현재 행동을 실제로 알고 있어야 쓸 수 있으므로, [타겟 독자]와 [독자가 지금 느끼는 통증]을 다시 읽고 그 사람이 오늘 실제로 한 행동·생각·변명 중 가장 뜨끔한 것을 골라라."

[좋은 변환 — 평이 → 자극, 어법은 자연스럽게]
- "배당주 고르는 법" → "10년 배당을 받고도 한 푼도 못 쓴 사람들의 공통점"
- "복리의 중요성" → "같은 돈을 넣었는데 7년 뒤 잔고가 두 배로 갈린 이유"
- "감정 관리가 필요하다" → "계좌를 자주 열어본 해일수록 오히려 수익률이 낮았던 이유"
- "분산 투자를 하자" → "한 종목에 몰빵한 사람이 그해 오히려 덜 잃은 까닭"
- "꾸준함이 답이다" → "3년을 버틴 사람과 3개월에 그만둔 사람의 통장이 비슷했던 이유"

[규칙]
- 소제목 절반 이상에 구체 숫자(금액·기간·비율·인원)를 넣는다.
- 통념을 뒤집거나 자기기만을 찌르는 줄을 PART마다 최소 1개.
- 답을 같은 줄에서 다 말하지 마라. "왜/어떻게/무엇이"의 정체는 본문에 숨긴다.
- 사람을 모욕하지 마라(욕설·인신공격 금지). 통념과 '행동'을 때리되 독자를 적으로 만들지 않는다.
- 자청·역행자 등 특정 작가의 '단어'는 쓰지 말 것. 톤만 가져온다.
- 자극을 위해 억지 비유나 말 안 되는 조합을 만들지 마라. 아래 [자연스러움 원칙]이 항상 우선.

[자연스러움 원칙 - 절대 어기지 말 것]
⚠️ 어법이 1순위다. 호기심보다 "말이 되는가"가 먼저다. 어법이 어색하면 호기심을 줄여서라도 자연스러운 문장으로 바꾼다.
- 모든 소제목/챕터 제목은 한국어 원어민 편집자가 손대지 않고 통과시킬, 문법적으로 완결된 자연스러운 문장이어야 한다.
- 단어를 억지로 조합한 "그럴듯해 보이지만 뜻이 안 통하는" 표현은 즉시 폐기. 특히 은유·비유를 무리하게 끼워 맞추지 마라.
  ❌ "배당 사이 파이프 굵기 차이는 정확히 얼마인가" (파이프 굵기? 배당 사이? — 무슨 말인지 알 수 없음)
  ❌ "수익이 자라는 토양의 산도를 맞추는 법" (억지 비유, 어법 어색)
  ✅ "같은 배당주를 사도 누구는 매달 80만원을 받고 누구는 한 푼도 못 버는 이유" (구체 금액으로 자연스럽고 강력)
- 어법 검사: 주어와 서술어가 자연스럽게 연결되는가? ("회복이 굴러간다" X — 회복은 굴러가지 않음)
- 대비 구조의 짝 맞추기(매우 중요): "A는 ~, B는 ~" 형태로 두 사람/상황을 대비할 때, 양쪽 서술어가 같은 축에서 정반대가 되어야 한다. 받는 것의 반대는 "못 받는 것", 버는 것의 반대는 "못 버는 것"이다. 축이 어긋나면 즉시 폐기.
  ❌ "같은 ETF를 샀는데 누구는 월세처럼 받고 누구는 한 푼도 못 쓴 이유" (받다 ↔ 못 쓰다 — 축이 어긋남)
  ✅ "같은 ETF를 샀는데 누구는 월세처럼 받고 누구는 한 푼도 못 버는 이유" (받다 ↔ 못 벌다 — 같은 축)
- 비유는 '명사'로만 쓰고 '되다/안 되다'의 대상으로 삼지 마라(비유 오용 금지). "월세"는 받는 금액의 비유지, "월세가 된다/안 된다"처럼 술어의 주어가 될 수 없다. 어색하면 비유를 버리고 '구체 금액'으로 직진하라 — 이게 더 강력하다.
  ❌ "10명 중 8명이 배당주를 사도 월세가 안 되는 이유" (월세가 안 된다? — 비유 오용, 말 안 됨)
  ✅ "10명 중 8명이 배당주를 사도 월 100만원조차 못 받는 이유" (구체 금액 — 자연스럽고 더 셈)
  ✅ "배당주 3년을 모아도 매달 통장에 30만원도 안 들어오는 사람들의 공통점" (구체 금액·장면)
- 추상 개념 + 물리·생물 동사 금지: 계좌·통장·재투자·수익·복리·멘탈·습관·시스템 같은 추상/사물에 "켜다/끄다/돌리다/감다/조이다/풀다/꽂다/박다/심다/굴린다/얹는다/멈춘다/자란다/살아난다/숨쉰다/깨어난다" 같은 동사를 붙이지 마라. 동사가 그 대상에 물리적으로 말이 되는지 검사하라.
  ❌ "캘린더에 박은 습관" (일정을 '박는다'? 어색)  ✅ "캘린더에 기재한 습관" / "매일 같은 시간에 반복한 습관" (자연스러움)
  ❌ "통장에 꽂히는 30만원" (꽂히다 = 비유 오용)  ✅ "매달 통장에 들어오는 30만원"
  → 의심되면 그 사물에 자연스럽게 쓰는 동사로 바꾼다(캘린더↔기재하다/적다, 돈↔들어오다/받다, 습관↔반복하다/들이다 등).
  ❌ "재투자를 켠 사람과 끈 사람" / "수익률 화면을 끄니 계좌가 살아났다" / "내버려 두면 계좌가 자란다"
  ✅ "수익을 다시 넣은 사람과 빼서 쓴 사람의 3년 뒤" / "계좌를 덜 들여다본 해에 수익률이 더 높았던 이유" (대비·호기심은 살리되 주어-서술어가 말이 되게)
- 과장 형용사 금지: "흔적도 없다", "통째로", "완전히" (꼭 필요할 때만)
- 추상 X, 구체 O: "많은 사람" → "월급 280만원짜리 7년차 회사원"
- 도구/플랫폼명 적극: 네이버, 카카오, 노션, 카톡, 캘린더, 구글 시트
- 출력 직전, 모든 줄을 소리내어 읽어라. 원어민이 한 번에 이해 못 하거나 "이게 무슨 말이지?" 싶은 줄은 전부 다시 쓴다.

[한 줄 = 한 호흡 — 정보 욱여넣기가 비문과 '노잼'의 가장 큰 원인]
한 소제목에 비교·숫자·기간·원리를 두세 개씩 밀어 넣으면 반드시 문장이 꼬여 비문이 되고, 호기심도 흩어져 끌리지 않는다. 가장 센 핵심 '하나'만 남기고 나머지는 과감히 쳐내라. 짧고 선명한 줄이 더 자연스럽고 훨씬 더 읽고 싶어진다.
✗ "배당률 5%를 찾아 헤맨 3년과 3%로 시작한 1년의 통장이 뒤집힌 순간"
   → 비교(5%·3%)+기간(3년·1년)+결과를 한 줄에 다 욱여넣어 비문. '통장이 뒤집히다'도 어색.
   ✓ "고배당만 좇던 3년보다, 낮은 배당으로 묵힌 1년이 더 많이 남은 이유" (비교를 한 축으로, 의외성으로 호기심)
✗ "계좌를 하루 세 번 확인한 사람이 수익률을 스스로 갉아먹은 증거"
   → '수익률을 갉아먹다'는 어색한 은유, '증거'로 끝나니 뻣뻣하고 김이 샌다.
   ✓ "하루에도 몇 번씩 계좌를 열어본 해에 오히려 수익률이 가장 낮았다" (행동→의외의 결과, 자연스럽고 궁금함)

[은유는 버리고 직진하라 — 어색한 은유는 어법도 자극도 둘 다 죽인다]
'갉아먹다 / 뒤집히다 / 몸집을 불리다 / 굴러가다 / 살아나다 / 깨어나다' 같은 은유를 돈·계좌·수익·복리·멘탈 같은 추상·사물에 붙이면 어색해지고 호기심도 식는다. 은유를 빼고 '구체적 장면 + 숫자 + 의외의 사실'로 직진하면 더 자연스럽고 더 센 호기심이 나온다. (예: "복리가 굴러간다" X → "재투자만 했는데 5년 뒤 받는 배당이 두 배가 됐다" O)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🚫 즉시 폐기 표현
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

특정 작가 고유어 (절대 금지): 자청 / 역행자 / 자의식 해체 / 유전자 역행 / 원시인 / 추월차선 / 아토믹 해빗 / 언카피어블
AI 클리셰: 졸업 / 정체 / 마지막 한 수 / 다른 차원 / 결정적 시그널 / 진짜 게임 / 흔적도 없다 / 사건이 아니다 / 회로가 박혔다
시스템 의인화: "○○가 멈춘 날", "○○ 위에 얹다", "○○를 졸업한", "회복이 굴러가다"
뇌과학 비유 남용: "뇌 회로가 박혔다", "뇌 회로가 새로 깔린다" (전체 목차에 뇌·신경회로는 사실 진술로 1~2회만, 비유 남용 X)
밍밍: 효과적인 / 성공적인 / ~의 모든 것 / ~하는 방법 / 알아야 할 / 의 중요성
유치 비유: 나침반 / 열쇠 / 보물 / 황금 / 마법 / 파이프라인 / 엔진 / 톱니바퀴 / 사이클 / 눈덩이
참고서: 첫걸음 / 완벽가이드 / 핵심정리 / 기초/중급/고급 / 첫 설계 / 첫 셋업
챕터 제목 라벨: "발굴 —", "1단계:", "STEP 1." 같은 분리 형식
의문문 문어체: "왜 ~는 ~하지 못하는가" 식의 한 PART에 1개까지만
콜론(:) — 단 한 번도 쓰지 마라
숫자 중복 금지: 전체 목차에서 같은 숫자(예: 3시간 + 3개월) 두 번 등장 금지

★★ 표현·단어 반복 절대 금지 — 목차 전체(25줄)에서 같은 말은 '딱 한 번'만 ★★
한 번 쓴 단어·표현·문장 끝 패턴은 그걸로 끝이다. 두 번째부터는 무조건 다른 말로 바꾼다. (목차가 평이해 보이는 가장 큰 원인이 '같은 말 반복'이다.)
- 핵심 명사·동사·형용사 반복 금지: '무너지다'를 한 줄에 썼으면 다른 줄에선 '주저앉다 / 흔들리다 / 제자리로 돌아가다' 등 다른 표현으로. ('이유·차이·순간·진짜·결정적·공통점·단 하나' 등 자주 겹치는 말도 전부 한 번씩만)
- 문장 끝 패턴 반복 금지: '~하는 이유'로 끝나는 줄은 전체에서 1개만. 나머지는 '~한 순간 / ~의 차이 / ~하는 사람들 / ~한 그날 / ~인 신호' 등 25줄이 전부 다른 엔딩이 되도록.
- 같은 통계 표현 반복 금지: '99%'를 썼으면 다른 줄은 '열에 아홉 / 대부분 / 상위 3%만' 처럼 다르게.
- 예외: 조사·접속사 같은 기능어, 컨셉명(정확히 3번)과 '○○○' 자리표시는 반복으로 치지 않는다.
- 출력 직전 점검: 25줄의 명사·동사·문장 끝을 쭉 훑어, 두 번 이상 나온 단어·표현이 있으면 한쪽을 반드시 다른 말로 교체한 뒤 출력한다.
숫자 표기 - 부호 금지: "-12%" "+30%" 같은 부호 사용 X. "12% 폭락에도", "30% 상승하면" 식으로

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📝 출력 형식 (이 형식 외 어떤 텍스트도 출력 금지)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[책 시그니처 컨셉]
○○○ + (시스템/구조/공식/알고리즘) | 한 줄 설명

PART 1. [13~18자, 좌절 + 통계 충격]
- [통계 충격 또는 결정적 함정형]
- [인지부조화/통념 박살형]
- [패턴 - 위 8개 중 다른 것]
- [컨셉명 첫 등장: 자연스러운 도입]

PART 2. [13~18자, 통념 박살 + 뇌과학 권위]
- [패턴]
- [패턴 - 다른 것]
- [패턴 - 또 다른 것]
- [패턴 - 또 다른 것]

PART 3. [13~18자, 첫 사건/결정적 순간]
- [패턴]
- [임박감/손해회피형]
- [패턴]
- [컨셉명 등장 가능: "○○○를 처음 적용한 사람들이 가장 먼저 느낀 변화" — PART 4에 넣을 거면 여기는 일반 패턴]

PART 4. [13~18자, 정체성 전환 선언]
- [컨셉명 등장 가능: "○○○가 본격 작동하는 4가지 신호" — PART 3에 안 넣었다면 여기에]
- [뇌과학 권위형 또는 통계형]
- [패턴]
- [패턴]

PART 5. [13~18자, 도약/이중 보상]
- [호기심 갭형]
- [이중/삼중 보상형]
- [패턴 - 또 다른 것]
- [컨셉명 세 번째 등장: 인그룹 사회증명]

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🔍 출력 전 자가 점검 (반드시 통과)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

출력 직전에 5가지 모두 통과해야 한다. 하나라도 No면 다시 써라.

체크 1. (스토리 흐름) 5개 PART 제목만 세로로 이어 읽었을 때, 같은 한 사람의 이야기가 '시간 순서(바닥→깨달음→사건→전환→도약)'로 이어지는가? 화자가 바뀌거나, 순서가 뒤섞이거나, 5줄이 따로 노는 독립 슬로건이면 실패 — 다시 써라.
체크 1-2. (감정 온도) PART 1이 가장 아프고(잃을 공포) PART 5가 가장 벅찬가(얻을 흥분)? 1→5로 세기가 한 단계씩 커지지 않고 평평하면 실패. 그리고 각 PART가 다음 PART를 안 보면 못 배기게 끝나는가?
체크 2. 통념 박살(인지부조화) 패턴이 5개 챕터+20개 소제목 안에 3개 이상 박혀있는가?
체크 3. 구체적 숫자(시간/금액/비율)가 8개 이상 등장하는가?
체크 4. 평이한 표현("~의 방법", "~의 모든 것", "~의 중요성", "효과적인", "성공적인")이 0개인가?
체크 5. 5초간 훑어본 가상 독자가 "이건 안 사면 손해"라고 느낄 만한 손실회피 트리거가 챕터 제목 5개 중 2개 이상에 있는가?
체크 6. (호기심+신뢰 균형) 소제목 20개 중 최소 12개가 '구체적 지렛대(행동·원리) 한 조각'을 드러내 "이렇게 하니까 되는구나"라는 믿음을 주는가? 막연히 '달라진다/완성된다/무너지지 않는다/따라온다'로 끝나는 순수 주장 줄이 PART당 2개 이상이면 실패 — 그 줄에 작동 원리나 지렛대를 한 조각 넣어 다시 써라. (디테일=종목·비율·타이밍은 여전히 숨긴다.)
체크 7. 컨셉명 글자 수를 단정하는 표현("세 글자", "네 글자", "단 한 단어")이 단 하나도 없는가? — 하나라도 있으면 즉시 전체 다시 쓰기.
체크 8. (어법 - 가장 중요) 모든 챕터 제목·소제목을 소리내어 읽었을 때, 원어민이 한 번에 이해되고 어법이 자연스러운가? 아래 셋 중 하나라도 걸리면 그 줄을 자연스럽게 다시 써라.
   (a) 한 줄에 비교·숫자·기간·원리를 두세 개씩 욱여넣어 길고 꼬인 비문("5%를 찾아 헤맨 3년과 3%로 시작한 1년의 통장이 뒤집힌 순간" 류) → 핵심 하나만 남기고 짧게.
   (b) 추상·사물에 어색한 은유 동사를 붙인 줄("수익률을 갉아먹다", "통장이 뒤집히다", "복리가 굴러가다" 류) → 은유 빼고 구체 장면·숫자로 직진.
   (c) "배당 사이 파이프 굵기 차이는…" 같은 억지 조합·뜻 모를 비유.
체크 9. (어법) 추상 개념에 기계·물리 동사를 붙인 줄("재투자를 켠/끈", "수익을 돌린다" 류)이 하나도 없는가? 있으면 즉시 자연스러운 표현으로 교체.
체크 10. (자극) 통념을 정면으로 뒤집거나 자기기만을 찌르는 줄이 PART마다 최소 1개 있는가? 전부 정보 전달형이면 평이한 것 — 다시 써라.
체크 11. (어법-대비) "A는 ~, B는 ~" 대비 줄에서 양쪽 서술어가 같은 축의 정반대인가? "받고 ↔ 못 쓴"처럼 축이 어긋난 줄이 하나라도 있으면 "받고 ↔ 못 받는"처럼 짝을 맞춰 다시 써라.
체크 12. (자극-결정타) 20개 소제목을 다시 훑어, 심장이 안 뛰는(밍밍한) 줄을 전부 골라내라. 그 줄은 [자청식 후킹 7가지 무기] 중 하나를 적용해 더 세게 다시 쓴다. "정보를 알려주는 줄"이 아니라 "안 읽으면 불안한 줄"이 되어야 한다.
체크 13. (비유 오용) "월세가 안 된다", "통장이 못 된다"처럼 비유 명사를 '되다/안 되다'의 대상으로 쓴 줄이 하나라도 있으면, 비유를 버리고 구체 금액("월 100만원조차 못 받는")으로 다시 써라.
체크 14. (뇌과학 겹치기) 20개 소제목 중 최소 6개는 뇌과학 트리거를 2개 이상 겹쳐 썼는가?(예: 사회증명+손실회피+구체금액) 한 줄에 트리거가 하나뿐인 평범한 줄이 절반을 넘으면 평이한 목차다 — 더 겹쳐 써라.
체크 15. (최종 결제 테스트 — 강화) 완성된 목차를 처음부터 끝까지 5초간 훑었다고 상상하라. "이 목차 정도는 나도 알겠다" 싶으면 실패다. 25줄을 한 줄씩 보며 '심장이 뛰는 정도'를 채점하라. 평범한 경쟁 도서 목차에도 들어갈 법한 줄(=교체 가능한 줄)은 전부 죽은 줄이다. 가장 약한 8줄을 골라 [자청식 후킹 7가지 무기] 중 하나로 한 단계 더 세게 다시 써라. "정보를 주는 줄"이 아니라 "안 읽으면 손해라 불안한 줄"이 25줄 중 다수가 될 때까지 반복한다. 단, 세기를 올리되 [자연스러움 원칙]과 어법 검사는 절대 희생하지 않는다 — 자극과 자연스러움이 부딪치면 어법이 이긴다.

체크 16. (반복 금지) 25줄의 핵심 단어·표현·문장 끝 패턴을 훑었을 때, 두 번 이상 나온 게 단 하나도 없는가?(기능어·컨셉명 제외) '이유/차이/결정적/신호/처음/가장 먼저/한 번' 같은 말이나 같은 숫자가 두 줄에 있으면 한쪽을 다른 표현으로 바꿔 다시 써라.
체크 17. (흐름-PART 중복) 5개 PART 제목 중 '사실상 같은 말'을 하는 PART가 둘 있는가?(예: 둘 다 "개선하면 매출 오른다") 있으면 흐름이 무너진 것 — 각 PART에 바닥/깨달음/첫 사건/전환/도약 중 '서로 다른 단 하나의 역할'만 주어 다시 써라.

목차만 출력. 콜론 금지. 매 소제목 다른 패턴. 명사·동사 엔딩 섞기. 같은 단어·표현·문장 끝은 전체에서 한 번만. 5개 PART는 이어 읽으면 서로 다른 5단계의 한 이야기. 각 PART는 정확히 소제목 4개. 컨셉명 글자수 단정 금지. 자가점검 결과는 출력하지 말 것. ★ 어법이 전부다 — 출력 직전 모든 줄을 소리내어 읽어라. 원어민이 한 번에 이해 못 하거나 '이게 무슨 말이지?' 싶은 줄은 자극을 줄이더라도 자연스러운 문장으로 바꿔 출력한다. 어색한 채로 내보내지 마라."""
    return ask_ai(prompt, 0.92, ensure_quality=True)


def generate_content_premium(subtopic, chapter, questions, answers, topic, persona):
    """자기계발 후킹 + 권석천 칼럼 깊이의 몰입형 본문"""
    import random

    qa_pairs = ""
    for i, (q, a) in enumerate(zip(questions, answers), 1):
        if a.strip():
            qa_pairs += f"\n질문{i}: {q}\n답변{i}: {a}\n"

    # 다양한 시작 스타일
    hook_styles = [
        "도발적 질문으로 시작 (예: '왜 99%는 여기서 실패할까요?')",
        "충격적 고백으로 시작 (예: '저도 2년간 완전히 틀리고 있었습니다.')",
        "반전 사실로 시작 (예: '사실 정반대였습니다.')",
        "구체적 숫자로 시작 (예: '정확히 23일 만에 달라졌습니다.')",
        "생생한 에피소드로 시작 (예: '그날 밤 컴퓨터 앞에서 깨달았습니다.')",
        "단호한 선언으로 시작 (예: '핵심부터 말씀드리겠습니다.')",
        "대화체로 시작 (예: '\"진짜요?\" 처음 들었을 때 저도 의심했습니다.')",
        "before/after로 시작 (예: '1년 전의 저는 완전히 다른 상황이었습니다.')",
        "상식 뒤집기로 시작 (예: '열심히 하면 된다? 틀렸습니다.')",
        "사건 장면으로 시작 (예: '2023년 3월의 일이었습니다.')",
        "인용으로 시작 (예: '한 후배가 이런 말을 했습니다.')",
    ]
    current_hook = random.choice(hook_styles)

    # 표 포함 여부 (랜덤하게 30% 확률)
    include_table = random.random() < 0.3

    prompt = f"""당신은 한국 자기계발 베스트셀러 작가입니다. '{subtopic}'에 대해 첫 문장으로 독자를 멈춰 세우고, 권석천 기자 칼럼처럼 정확한 디테일과 통찰로 끝까지 끌고 가는 본문을 씁니다.

[주제]: {topic}
[챕터]: {chapter}
[참고 내용]
{qa_pairs}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🔥 첫 문장이 전부다
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

반드시 이 스타일로 시작:
👉 {current_hook}

(일반론·교훈으로 시작하면 즉시 폐기. 사건/숫자/대사/의문 중 하나로만 시작)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
✍️ 본문 톤: 자기계발 후킹 + 권석천 칼럼 깊이
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

[문체]
- 합쇼체 기본 ("~입니다", "~합니다") + 가끔 구어체 ("~거든요", "~더라고요")
- 현재 시제로 장면을 그리듯
- 짧은 문장과 긴 문장을 교차해 리듬
- 추상보다 구체. "많은 사람" → "월급 280만원짜리 7년차 회사원"

[권석천식 깊이]
- 사실 → 분석 → 통찰 순서로 전개
- 가설 검증식 흐름: "왜 그럴까. 이유를 되짚어봤습니다"
- 사회적 맥락이나 통계, 책/논문 인용을 자연스럽게 섞기
- 결론 강요 X, 독자가 스스로 깨닫게 단서를 깔기

[자기계발식 후킹]
- 첫 문장: 사건/숫자/대사/의문
- 본문 중간: 작은 반전 1회 ("그런데 진짜 흥미로운 건 그 다음이었습니다")
- 마지막 문장: 발견의 결과로서의 통찰 한 줄

[표 포함: {'예' if include_table else '아니오'}]
{'''
📊 본문 중간에 비교표 1개 필수:
<table style="width:100%; border-collapse:collapse; margin:20px 0;">
<tr style="background:#1a1a1a;"><th style="border:1px solid #333;padding:12px;color:#C9A24B;">구분</th><th style="border:1px solid #333;padding:12px;color:#C9A24B;">기존 방식</th><th style="border:1px solid #333;padding:12px;color:#C9A24B;">새로운 방식</th></tr>
<tr><td style="border:1px solid #333;padding:10px;">항목</td><td style="border:1px solid #333;padding:10px;">내용</td><td style="border:1px solid #333;padding:10px;">내용</td></tr>
</table>
''' if include_table else ''}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🎯 구체성을 끝까지
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

- 숫자: "많이" X → "월 340만원, 정확히 47일" O
- 사례: 수강생/지인/현장 (이름은 가끔만)
- 실제 도구명: 네이버, 카카오, 노션, 구글 시트 등

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🚫 절대 금지
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

❌ 자청, 역행자, 자의식 해체, 유전자 역행, 원시인, 추월차선, 아토믹 해빗 (특정 작가/책 고유 표현)
❌ 유치한 비유: 마법, 황금열쇠, 나침반, 로켓, 눈덩이, 톱니바퀴, 파이프라인
❌ AI 어휘: 중요합니다, 따라서, 결론적으로, ~를 통해, 다양한, 효과적인, 진정한
❌ 직접 호칭: 여러분, 당신, 독자님
❌ 형식: 1. 2. 첫째, 글머리 기호, 이모지
❌ 같은 이름 반복 (민준, 지수가 계속 X)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🔥 몰입·충격 어조 (자청식 몰입감 — 특정 작가 고유어는 제외)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

- 도입부터 통념을 정면으로 깬다. 독자가 "어, 내가 알던 거랑 반대네" 하고 멈칫하게.
- 글 중간마다 '작은 충격'을 심어라: 예상과 반대되는 사실, 뒤집히는 데이터, 의외의 결론.
- 한 호흡에 읽히는 리듬 — 짧은 단정문으로 치고, 긴 문장으로 풀고, 다시 끊는다.
- "왜?"를 계속 만들어 다음 문단을 안 읽고는 못 배기게 한다(궁금증 갭).
- 충격 뒤엔 반드시 '그래서 무엇을 어떻게'의 실질이 온다. 겁만 주고 끝내지 않는다.
- 자청·역행자 등 특정 작가의 '단어'는 절대 쓰지 않는다. 몰입감과 태도만 가져온다.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🎓 전문성 (아마추어 글과 가르는 지점)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

- 추상 주장 1개당 근거 1개(데이터·연구·구체 사례·숫자)를 반드시 붙인다.
- 메커니즘을 설명하라. "왜 그렇게 되는가"의 원리와 과정을 단계로 풀어라.
- 바로 따라 할 수 있게: 무엇을·어디서·어떤 순서로·얼마나, 구체적으로.
- 흔한 조언의 한계를 짚고, 더 정확한 기준·예외를 제시한다(깊이의 증거).

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📏 분량: 2400~3000자 (최소 2400자 이상, 전문성 있는 밀도로)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

⚠️ 길이는 사례·데이터·단계별 설명·메커니즘으로 채운다. 같은 말 반복·미사여구로 늘리면 즉시 실패.

'{subtopic}' 본문 작성.
- 시작: {current_hook}
- 권석천 칼럼처럼 사실 → 분석 → 통찰 순서로 인과 추적
- 마지막 한 줄에 발견된 통찰 하나
- {'비교표 1개 포함' if include_table else '순수 텍스트만'}

⛔ 절대 금지: 본문 첫 줄에 소제목('{subtopic}')을 다시 쓰지 마라.
   소제목은 위에 이미 표시되므로, 본문은 곧장 첫 후킹 문장으로 시작한다.
   ❌ 잘못: "{subtopic}\\n\\n그날 새벽 두 시였습니다..." (소제목 반복)
   ✅ 올바름: "그날 새벽 두 시였습니다..." (바로 본문 시작)"""
    return ask_ai(prompt, 0.75, ensure_quality=True)


def format_content_html(content):
    """본문을 HTML 형식으로 변환 (강조 표시 적용)"""
    if not content:
        return ""
    # 「」 → 주황색 볼드
    formatted = re.sub(r'「([^」]+)」', r'<b style="color:#e67e22;">\1</b>', content)
    # ★ → 주황색 볼드 문장
    formatted = re.sub(r'★\s*(.+?)(?=\n|$)', r'<p style="color:#e67e22;font-weight:700;margin:20px 0;font-size:17px;">★ \1</p>', formatted)
    # 문단 구분 (빈 줄) → 문단 간격
    formatted = formatted.replace('\n\n', '</p><p style="color:#000000;margin:25px 0;line-height:2.0;font-size:17px;">')
    # 단일 줄바꿈 제거 (문단 내 연결)
    formatted = formatted.replace('\n', ' ')
    formatted = f'<p style="color:#000000;margin:25px 0;line-height:2.0;font-size:17px;">{formatted}</p>'
    # 빈 <p> 태그 정리
    formatted = re.sub(r'<p[^>]*>\s*</p>', '', formatted)
    return formatted


def generate_questions(subtopic, chapter, topic):
    prompt = f"""'{topic}' 전자책 '{chapter}' 챕터의 '{subtopic}' 작성용 질문 3개:

Q1: [질문]
Q2: [질문]
Q3: [질문]"""
    return ask_ai(prompt, 0.7)


# ==========================================
# 메인 UI
# ==========================================
# 비디오 배경 헤더 — title_bg.webm / title_bg.mp4 를 여러 위치에서 자동 탐색
def _find_header_video(exts=("webm", "mp4")):
    """헤더 배경 영상을 찾는다.
    우선순위: (1) title_bg.webm/mp4  (2) 폴더 안의 아무 .webm/.mp4 파일.
    클라우드 배포 시 업로드한 임의 파일명도 자동으로 잡히게 한다."""
    import glob
    bases = []
    try:
        here = os.path.dirname(os.path.abspath(__file__))
        bases.append(here)
    except Exception:
        pass
    bases += [os.getcwd(), "", os.path.expanduser("~/Desktop"), "/Users/hyunwoo/Desktop"]
    # 중복 제거(순서 유지)
    seen = set(); ordered_bases = []
    for b in bases:
        if b not in seen:
            seen.add(b); ordered_bases.append(b)

    found = {}
    # 1순위: title_bg.<ext>
    for ext in exts:
        for base in ordered_bases:
            p = os.path.join(base, f"title_bg.{ext}") if base else f"title_bg.{ext}"
            if p and os.path.exists(p):
                found[ext] = p
                break
    # 2순위: 폴더 안의 아무 .<ext> 파일 (title_bg가 없을 때)
    for ext in exts:
        if ext in found:
            continue
        for base in ordered_bases:
            pattern = os.path.join(base, f"*.{ext}") if base else f"*.{ext}"
            try:
                matches = sorted(glob.glob(pattern))
            except Exception:
                matches = []
            if matches:
                found[ext] = matches[0]
                break
    return found

_header_videos = _find_header_video()
header_video_webm_b64 = get_video_base64(_header_videos["webm"]) if "webm" in _header_videos else None
header_video_mp4_b64 = get_video_base64(_header_videos["mp4"]) if "mp4" in _header_videos else None
header_video_b64 = header_video_webm_b64 or header_video_mp4_b64  # 둘 중 하나라도 있으면 헤더 표시

# 영상 파일을 못 찾으면, 어디를 찾았는지 화면에 알려준다 (디버그)
if not header_video_b64:
    try:
        _here = os.path.dirname(os.path.abspath(__file__))
    except Exception:
        _here = "(알 수 없음)"
    st.info(
        "🎬 제목 배경 영상을 찾지 못했습니다. "
        "`.mp4` 또는 `.webm` 영상 파일을 이 프로그램과 같은 폴더(저장소)에 올려 주세요. "
        "(파일명은 아무거나 괜찮습니다 — 폴더 안의 첫 영상을 자동으로 사용합니다.)\n\n"
        f"- 프로그램 폴더: `{_here}`\n"
        f"- 현재 작업 폴더: `{os.getcwd()}`"
    )

if header_video_b64:
    st.markdown("""
    <style>
    .writey-brandbar {
        display:flex; align-items:center; justify-content:space-between;
        padding:15px 24px; border-bottom:0.5px solid rgba(255,255,255,0.06);
        margin-bottom:18px;
    }
    .writey-wordmark { font-size:18px; font-weight:500; letter-spacing:0.18em; color:#F5F3EF; }
    .writey-cashtag { font-size:10px; letter-spacing:0.28em; color:#7A776F; margin-left:10px; }
    .writey-author { font-size:11px; letter-spacing:0.1em; color:#7A776F; }
    </style>
    <div class="writey-brandbar">
        <div><span class="writey-wordmark">WRITEY</span><span class="writey-cashtag">CASHMAKER</span></div>
        <span class="writey-author">남현우 작가</span>
    </div>
    """, unsafe_allow_html=True)

    # 영상 히어로는 components.html(iframe)로 렌더 → 큰 base64 data URI도 정상 재생
    _sources = ""
    if header_video_webm_b64:
        _sources += f'<source src="data:video/webm;base64,{header_video_webm_b64}" type="video/webm">'
    if header_video_mp4_b64:
        _sources += f'<source src="data:video/mp4;base64,{header_video_mp4_b64}" type="video/mp4">'
    components.html(f"""
    <!DOCTYPE html><html><head><meta charset="utf-8">
    <style>
        html,body {{ margin:0; padding:0; background:transparent; overflow:hidden; }}
        .hero {{
            position:relative; width:100%; height:320px;
            border-radius:16px; overflow:hidden;
            border:0.5px solid rgba(201,162,75,0.25);
            box-shadow:0 18px 50px rgba(0,0,0,0.45);
            font-family:'Pretendard','Apple SD Gothic Neo',sans-serif;
            background:#0B0B0D;
        }}
        .hero video {{
            position:absolute; top:0; left:0; width:100%; height:100%;
            object-fit:cover; filter:brightness(0.45) saturate(1.15) contrast(1.05);
        }}
        .hero .veil {{
            position:absolute; inset:0;
            background:radial-gradient(ellipse at center, rgba(11,11,13,0.25) 0%, rgba(11,11,13,0.72) 100%);
        }}
        .hero .center {{
            position:absolute; inset:0; display:flex; flex-direction:column;
            justify-content:center; align-items:center; text-align:center;
        }}
        .eyebrow {{ color:#C9A24B; font-size:12px; letter-spacing:0.4em; margin-bottom:16px; font-weight:400; }}
        .title {{ font-size:60px; font-weight:300; letter-spacing:0.14em; color:#FAF8F4; margin:0; }}
        .divider {{ width:36px; height:1px; background:#C9A24B; margin:22px auto 0; }}
        .tagline {{ color:#C8C5BE; font-size:15px; margin-top:18px; font-weight:300; letter-spacing:0.03em; }}
    </style></head>
    <body>
        <div class="hero">
            <video id="bgv" autoplay muted loop playsinline preload="auto">
                {_sources}
            </video>
            <div class="veil"></div>
            <div class="center">
                <div class="eyebrow">AI EBOOK Writer</div>
                <h1 class="title">WRITEY</h1>
                <div class="divider"></div>
                <p class="tagline">6개의 질문에 답하면 AI가 목차부터 본문까지 완성합니다</p>
            </div>
            <div id="vdbg" style="position:absolute;left:10px;bottom:8px;font-size:10px;color:#C9A24B;font-family:monospace;opacity:0.9;background:rgba(0,0,0,0.4);padding:2px 8px;border-radius:6px;"></div>
        </div>
        <script>
        (function() {{
            var v = document.getElementById('bgv');
            var dbg = document.getElementById('vdbg');
            function show(msg) {{ if(dbg){{ dbg.textContent = msg; dbg.style.opacity = '0.9'; }} }}
            if (!v) {{ show('video 태그 없음'); return; }}
            v.muted = true; v.defaultMuted = true;
            var attempts = 0;
            var tryPlay = function() {{
                attempts++;
                var p = v.play();
                if (p && p.then) {{
                    p.then(function() {{ show('재생 중 ✓'); setTimeout(function(){{ dbg.style.opacity='0'; }}, 2000); }})
                     .catch(function(err) {{
                        show('재생차단(' + err.name + ') 재시도' + attempts);
                        if (attempts < 6) setTimeout(tryPlay, 500);
                     }});
                }} else {{ show('재생 시도(구형 브라우저)'); }}
            }};
            v.addEventListener('loadeddata', tryPlay);
            v.addEventListener('canplay', tryPlay);
            v.addEventListener('error', function() {{ show('영상 로드 실패 — 코덱/포맷 확인'); }});
            var src = v.querySelector('source');
            if (src) src.addEventListener('error', function() {{ show('소스 디코딩 실패'); }});
            tryPlay();
            setTimeout(function() {{ if (v.paused && v.readyState < 2) show('영상 데이터 미로딩(파일 확인)'); else if (v.paused) tryPlay(); }}, 1500);
        }})();
        </script>
    </body></html>
    """, height=340)
else:
    st.markdown("""
    <div style="display:flex;align-items:center;justify-content:space-between;padding:15px 24px;border-bottom:0.5px solid rgba(255,255,255,0.06);margin-bottom:30px;">
        <div><span style="font-size:18px;font-weight:500;letter-spacing:0.18em;color:#F5F3EF;">WRITEY</span><span style="font-size:10px;letter-spacing:0.28em;color:#7A776F;margin-left:10px;">CASHMAKER</span></div>
        <span style="font-size:11px;letter-spacing:0.1em;color:#7A776F;">남현우 작가</span>
    </div>
    <div style="text-align:center;padding:56px 20px 44px;margin-bottom:30px;background:#0B0B0D;border-radius:16px;border:0.5px solid rgba(201,162,75,0.18);">
        <div style="font-size:12px;letter-spacing:0.4em;color:#C9A24B;margin-bottom:16px;">AI EBOOK Writer</div>
        <h1 style="font-family:'S-CoreDream','Pretendard',sans-serif;font-size:54px;font-weight:300;letter-spacing:0.14em;color:#FAF8F4;margin:0;">WRITEY</h1>
        <div style="width:36px;height:1px;background:#C9A24B;margin:22px auto 0;"></div>
        <p style="color:#8A8780;font-size:15px;margin-top:18px;font-weight:300;">6개의 질문에 답하면 AI가 목차부터 본문까지 완성합니다</p>
    </div>
    """, unsafe_allow_html=True)

# ==========================================
# 인터뷰 모드 (interview_completed가 False일 때)
# ==========================================
if not st.session_state.get('interview_completed', False):
    st.markdown("""
    <div class="section-title-box">
        <span class="section-step">INTERVIEW</span>
        <h2>나만의 전자책 만들기</h2>
        <p>몇 가지 질문에 답하면 AI가 전자책을 완성해드립니다</p>
    </div>
    """, unsafe_allow_html=True)

    # 인터뷰 진행 상태
    if 'interview_step' not in st.session_state:
        st.session_state['interview_step'] = 1

    step = st.session_state['interview_step']
    total_steps = 6

    # 진행률 표시
    st.progress(step / total_steps)
    st.caption(f"질문 {step} / {total_steps}")

    st.markdown("---")

    # 인터뷰 데이터 임시 저장
    if 'temp_interview' not in st.session_state:
        st.session_state['temp_interview'] = {}

    # ========== STEP 1: 기본 정보 ==========
    if step == 1:
        st.markdown("""
        <div style="background:linear-gradient(135deg, rgba(201,162,75,0.1) 0%, rgba(201,162,75,0.05) 100%);padding:30px;border-radius:15px;border-left:4px solid var(--gold);margin-bottom:30px;">
            <h3 style="color:var(--gold);margin:0 0 10px 0;">👋 먼저 당신에 대해 알려주세요</h3>
            <p style="color:var(--text);margin:0;opacity:0.9;">전자책의 저자로서 기본 정보를 입력해주세요</p>
        </div>
        """, unsafe_allow_html=True)

        with st.form(key="step1_form"):
            author_name = st.text_input("저자명 (필명 가능)", value=st.session_state['temp_interview'].get('author_name', ''), placeholder="예: 김성장, 머니메이커 등")
            field = st.text_input("당신의 전문 분야는?", value=st.session_state['temp_interview'].get('field', ''), placeholder="예: 주식투자, 블로그 수익화, 다이어트, 영어회화 등")

            exp_options = ["선택하세요", "1년 미만", "1~2년", "3~5년", "5~10년", "10년 이상"]
            saved_exp = st.session_state['temp_interview'].get('experience_years', '선택하세요')
            exp_index = exp_options.index(saved_exp) if saved_exp in exp_options else 0
            experience = st.selectbox("이 분야 경험은?", exp_options, index=exp_index)

            col1, col2 = st.columns([1, 1])
            with col2:
                submitted = st.form_submit_button("다음 →", use_container_width=True, type="primary")

            if submitted:
                if not author_name.strip() or not field.strip() or experience == "선택하세요":
                    st.error("모든 항목을 입력해주세요")
                else:
                    st.session_state['temp_interview']['author_name'] = author_name.strip()
                    st.session_state['temp_interview']['field'] = field.strip()
                    st.session_state['temp_interview']['experience_years'] = experience
                    st.session_state['interview_step'] = 2
                    st.rerun()

    # ========== STEP 2: 주제와 노하우 ==========
    elif step == 2:
        st.markdown("""
        <div style="background:linear-gradient(135deg, rgba(201,162,75,0.1) 0%, rgba(201,162,75,0.05) 100%);padding:30px;border-radius:15px;border-left:4px solid var(--gold);margin-bottom:30px;">
            <h3 style="color:var(--gold);margin:0 0 10px 0;">📚 어떤 내용을 담을까요?</h3>
            <p style="color:var(--text);margin:0;opacity:0.9;">당신만의 핵심 노하우를 알려주세요</p>
        </div>
        """, unsafe_allow_html=True)

        col_prev, col_next = st.columns([1, 1])
        with col_prev:
            if st.button("← 이전", key="interview_prev_2", use_container_width=True):
                st.session_state['interview_step'] = 1
                st.rerun()

        with st.form(key="step2_form"):
            topic = st.text_input("전자책 주제", value=st.session_state['temp_interview'].get('topic', ''), placeholder="예: 월 100만원 배당 투자, 하루 1시간 블로그로 월 300 벌기")
            core_method = st.text_area("당신만의 핵심 방법/노하우는?", value=st.session_state['temp_interview'].get('core_method', ''), height=120, placeholder="예: 저는 고배당 ETF를 활용해서 안정적으로 수익을 내는 방법을 알려드립니다. 핵심은 분산투자와 복리의 마법입니다...")
            detailed_knowhow = st.text_area(
                "구체적인 노하우·실전 디테일 (선택, 많을수록 책이 깊어집니다)",
                value=st.session_state['temp_interview'].get('detailed_knowhow', ''),
                height=180,
                placeholder=(
                    "여기에 적은 내용이 목차와 본문에 그대로 녹아 들어갑니다. 구체적일수록 좋습니다.\n\n"
                    "예)\n"
                    "- 실제 수치/사례: '2년간 매달 50만원씩 적립, 3년 차에 월 배당 32만원 돌파'\n"
                    "- 단계별 방법: '1단계 종목 스크리닝 → 2단계 분할매수 → 3단계 배당 재투자'\n"
                    "- 자주 하는 실수: '초보는 고배당률만 보고 사는데, 배당성향 80% 넘으면 위험'\n"
                    "- 나만의 기준/팁: '나는 배당락 2주 전엔 절대 안 산다', '월 단위로 리밸런싱'\n"
                    "- 도구/플랫폼: '네이버 증권으로 배당 캘린더 관리, 구글시트로 수익률 추적'"
                ),
            )

            col1, col2 = st.columns([1, 1])
            with col2:
                submitted = st.form_submit_button("다음 →", use_container_width=True, type="primary")

            if submitted:
                if not topic.strip() or not core_method.strip():
                    st.error("주제와 핵심 방법은 꼭 입력해주세요 (구체적 노하우는 선택)")
                else:
                    st.session_state['temp_interview']['topic'] = topic.strip()
                    st.session_state['temp_interview']['core_method'] = core_method.strip()
                    st.session_state['temp_interview']['detailed_knowhow'] = detailed_knowhow.strip()
                    st.session_state['interview_step'] = 3
                    st.rerun()

    # ========== STEP 3: 타겟 독자 (AI 추천) ==========
    elif step == 3:
        st.markdown("""
        <div style="background:linear-gradient(135deg, rgba(201,162,75,0.1) 0%, rgba(201,162,75,0.05) 100%);padding:30px;border-radius:15px;border-left:4px solid var(--gold);margin-bottom:30px;">
            <h3 style="color:var(--gold);margin:0 0 10px 0;">🎯 누구를 위한 책인가요?</h3>
            <p style="color:var(--text);margin:0;opacity:0.9;">AI가 시장 데이터를 분석해 최적의 타겟을 추천해드립니다</p>
        </div>
        """, unsafe_allow_html=True)

        topic = st.session_state['temp_interview'].get('topic', '')

        # AI 타겟 분석 (캐시)
        if 'ai_target_suggestions' not in st.session_state or st.session_state.get('ai_target_topic') != topic:
            if st.button("🔍 AI 타겟 분석 시작", key="analyze_target", use_container_width=True, type="primary"):
                with st.spinner("시장 데이터 분석 중..."):
                    result = suggest_targets(topic)
                    parsed = parse_json(result)
                    if parsed and parsed.get('personas'):
                        st.session_state['ai_target_suggestions'] = parsed['personas']
                        st.session_state['ai_target_topic'] = topic
                        st.rerun()
                    else:
                        st.error("분석 실패. 다시 시도해주세요.")

        # AI 추천 결과 표시
        if st.session_state.get('ai_target_suggestions'):
            st.markdown("### 📊 AI 추천 타겟")
            personas = st.session_state['ai_target_suggestions']

            selected_idx = st.session_state.get('selected_target_idx', 0)

            for idx, persona in enumerate(personas[:3]):
                is_selected = (idx == selected_idx)
                border_color = "var(--gold)" if is_selected else "var(--line)"
                bg_color = "rgba(201,162,75,0.1)" if is_selected else "rgba(20,20,20,0.5)"

                pain_list = persona.get('pain_points', [])[:3]
                pains_text = " / ".join(pain_list) if pain_list else "고민 분석 중..."

                st.markdown(f"""
                <div style="background:{bg_color};border:1px solid {border_color};border-radius:10px;padding:15px;margin-bottom:10px;">
                    <div style="font-weight:bold;color:var(--gold);margin-bottom:5px;">{persona.get('name', '타겟')}</div>
                    <div style="font-size:13px;color:var(--text2);margin-bottom:8px;">{persona.get('demographics', '')}</div>
                    <div style="font-size:12px;color:var(--text);opacity:0.8;">💭 {pains_text}</div>
                </div>
                """, unsafe_allow_html=True)

                if st.button(f"✓ 이 타겟 선택" if not is_selected else "✓ 선택됨", key=f"select_target_{idx}", use_container_width=True, disabled=is_selected):
                    st.session_state['selected_target_idx'] = idx
                    st.rerun()

            st.markdown("---")

            # 선택된 타겟 정보 자동 입력
            selected_persona = personas[selected_idx] if selected_idx < len(personas) else personas[0]
            default_reader = f"{selected_persona.get('name', '')} ({selected_persona.get('demographics', '')})"
            default_problem = " ".join(selected_persona.get('pain_points', [])[:3])

            st.markdown("##### 선택된 타겟 정보 (수정 가능)")
            target_reader = st.text_input("타겟 독자", value=st.session_state['temp_interview'].get('target_reader', '') or default_reader, key="target_reader_input")
            target_problem = st.text_area("이 독자들의 가장 큰 고민/문제", value=st.session_state['temp_interview'].get('target_problem', '') or default_problem, height=80, key="target_problem_input")

            col1, col2 = st.columns([1, 1])
            with col1:
                if st.button("← 이전", key="interview_prev_3", use_container_width=True):
                    st.session_state['interview_step'] = 2
                    st.rerun()
            with col2:
                if st.button("다음 →", key="interview_next_3", use_container_width=True, type="primary"):
                    if not target_reader or not target_problem:
                        st.error("타겟 독자와 고민을 입력해주세요")
                    else:
                        st.session_state['temp_interview']['target_reader'] = target_reader
                        st.session_state['temp_interview']['target_problem'] = target_problem
                        st.session_state['interview_step'] = 4
                        st.rerun()
        else:
            # AI 분석 전 직접 입력 옵션
            st.markdown("---")
            st.markdown("##### 또는 직접 입력")
            target_reader = st.text_input("타겟 독자", value=st.session_state['temp_interview'].get('target_reader', ''), placeholder="예: 30대 직장인, 투자 초보자")
            target_problem = st.text_area("이 독자들의 가장 큰 고민/문제는?", value=st.session_state['temp_interview'].get('target_problem', ''), height=80, placeholder="예: 월급만으로는 부족하고, 어디서부터 시작해야 할지 모르겠다...")

            col1, col2 = st.columns([1, 1])
            with col1:
                if st.button("← 이전", key="interview_prev_3_manual", use_container_width=True):
                    st.session_state['interview_step'] = 2
                    st.rerun()
            with col2:
                if st.button("다음 →", key="interview_next_3_manual", use_container_width=True, type="primary"):
                    if not target_reader or not target_problem:
                        st.error("타겟 독자와 고민을 입력해주세요")
                    else:
                        st.session_state['temp_interview']['target_reader'] = target_reader
                        st.session_state['temp_interview']['target_problem'] = target_problem
                        st.session_state['interview_step'] = 4
                        st.rerun()

    # ========== STEP 4: 스토리 & 경력 ==========
    elif step == 4:
        st.markdown("""
        <div style="background:linear-gradient(135deg, rgba(201,162,75,0.1) 0%, rgba(201,162,75,0.05) 100%);padding:30px;border-radius:15px;border-left:4px solid var(--gold);margin-bottom:30px;">
            <h3 style="color:var(--gold);margin:0 0 10px 0;">💪 당신의 이야기를 들려주세요</h3>
            <p style="color:var(--text);margin:0;opacity:0.9;">독자들이 공감할 수 있는 진솔한 경험담과 경력</p>
        </div>
        """, unsafe_allow_html=True)

        struggle_story = st.text_area("처음 시작할 때 겪었던 어려움/실패는?", value=st.session_state['temp_interview'].get('struggle_story', ''), height=100, placeholder="예: 처음에는 주식으로 500만원을 잃었습니다. 유튜브 정보만 믿고 투자했다가 큰 손실을 봤죠...")
        breakthrough = st.text_area("어떻게 극복하고 성과를 냈나요?", value=st.session_state['temp_interview'].get('breakthrough', ''), height=100, placeholder="예: 그 후 기본서 10권을 정독하고, 나만의 원칙을 세웠습니다. 1년 후 손실을 모두 만회하고 수익을 내기 시작했습니다...")

        st.markdown("---")
        st.markdown("##### 📌 작가 경력/경험 (선택)")
        author_career = st.text_area("관련 경력이나 자격, 성과가 있다면?", value=st.session_state['temp_interview'].get('author_career', ''), height=100, placeholder="예: 금융회사 7년 근무, 투자 관련 유튜브 구독자 5만명, 월 수익 3천만원 달성, CFA 자격증 보유, 강의 경력 3년...")

        col1, col2 = st.columns([1, 1])
        with col1:
            if st.button("← 이전", key="interview_prev_4", use_container_width=True):
                st.session_state['interview_step'] = 3
                st.rerun()
        with col2:
            if st.button("다음 →", key="interview_next_4", use_container_width=True, type="primary"):
                if not struggle_story or not breakthrough:
                    st.error("어려움/실패와 극복 스토리는 필수입니다")
                else:
                    st.session_state['temp_interview']['struggle_story'] = struggle_story
                    st.session_state['temp_interview']['breakthrough'] = breakthrough
                    st.session_state['temp_interview']['author_career'] = author_career
                    st.session_state['interview_step'] = 5
                    st.rerun()

    # ========== STEP 5: 마무리 ==========
    elif step == 5:
        st.markdown("""
        <div style="background:linear-gradient(135deg, rgba(201,162,75,0.1) 0%, rgba(201,162,75,0.05) 100%);padding:30px;border-radius:15px;border-left:4px solid var(--gold);margin-bottom:30px;">
            <h3 style="color:var(--gold);margin:0 0 10px 0;">✨ 마지막으로!</h3>
            <p style="color:var(--text);margin:0;opacity:0.9;">독자에게 전하고 싶은 메시지</p>
        </div>
        """, unsafe_allow_html=True)

        why_write = st.text_area("왜 이 책을 쓰려고 하나요?", value=st.session_state['temp_interview'].get('why_write', ''), height=80, placeholder="예: 저처럼 헤매는 사람들이 시행착오 없이 바로 성과를 낼 수 있도록 도와주고 싶습니다...")
        final_message = st.text_area("독자에게 마지막으로 전하고 싶은 말", value=st.session_state['temp_interview'].get('final_message', ''), height=80, placeholder="예: 누구나 할 수 있습니다. 포기하지 않으면 반드시 성공합니다...")

        # 입력 내용 미리보기
        st.markdown("---")
        st.markdown("### 📋 입력 내용 확인")

        preview_data = st.session_state['temp_interview']
        st.markdown(f"""
        <div style="background:rgba(20,20,20,0.8);padding:20px;border-radius:10px;border:1px solid var(--line);">
            <p><b>저자:</b> {preview_data.get('author_name', '')}</p>
            <p><b>분야:</b> {preview_data.get('field', '')} ({preview_data.get('experience_years', '')})</p>
            <p><b>주제:</b> {preview_data.get('topic', '')}</p>
            <p><b>타겟:</b> {preview_data.get('target_reader', '')}</p>
        </div>
        """, unsafe_allow_html=True)

        col1, col2 = st.columns([1, 1])
        with col1:
            if st.button("← 이전", key="interview_prev_5", use_container_width=True):
                st.session_state['interview_step'] = 4
                st.rerun()
        with col2:
            if st.button("📋 목차 생성하기", key="interview_generate_outline", use_container_width=True, type="primary"):
                if not get_api_key():
                    st.error("사이드바에서 API 키를 먼저 입력해주세요")
                elif not why_write or not final_message:
                    st.error("모든 항목을 입력해주세요")
                else:
                    st.session_state['temp_interview']['why_write'] = why_write
                    st.session_state['temp_interview']['final_message'] = final_message

                    # 목차만 먼저 생성
                    progress_box = st.empty()
                    interview_data = st.session_state['temp_interview']
                    success = generate_outline_only(interview_data, progress_box)

                    if success:
                        import time
                        time.sleep(1)
                        st.session_state['interview_step'] = 6  # 목차 확인 단계로 이동
                        st.rerun()

    # ========== STEP 6: 목차 확인 및 본문 생성 ==========
    elif step == 6:
        st.markdown("""
        <div style="background:linear-gradient(135deg, rgba(201,162,75,0.1) 0%, rgba(201,162,75,0.05) 100%);padding:30px;border-radius:15px;border-left:4px solid var(--gold);margin-bottom:30px;">
            <h3 style="color:var(--gold);margin:0 0 10px 0;">📋 목차 확인 및 수정</h3>
            <p style="color:var(--text);margin:0;opacity:0.9;">생성된 목차를 확인하고, 직접 수정하거나 AI로 재생성할 수 있습니다</p>
        </div>
        """, unsafe_allow_html=True)

        # 제목 표시
        book_title = st.session_state.get('book_title', '')
        subtitle = st.session_state.get('subtitle', '')
        book_concept = st.session_state.get('book_concept', '')

        if book_title:
            st.markdown(f"""
            <div style="background:rgba(30,30,30,0.9);padding:25px;border-radius:15px;border:0.5px solid var(--gold);margin-bottom:20px;text-align:center;">
                <h2 style="color:var(--gold);margin:0 0 10px 0;font-size:32px;">{book_title}</h2>
                <p style="color:var(--text2);margin:0;font-size:18px;">{subtitle}</p>
            </div>
            """, unsafe_allow_html=True)

        # 컨셉 표시
        if book_concept:
            with st.expander("💡 이 책의 고유 컨셉 보기", expanded=False):
                st.markdown(f"""
                <div style="background:rgba(201,162,75,0.1);padding:20px;border-radius:10px;border-left:3px solid var(--gold);">
                    {book_concept.replace(chr(10), '<br>')}
                </div>
                """, unsafe_allow_html=True)

        st.markdown("---")

        # 목차 표시 및 편집
        outline = st.session_state.get('outline', [])
        chapters = st.session_state.get('chapters', {})

        if outline:
            st.markdown("### 📖 목차 구성")
            st.caption("각 챕터와 소제목을 직접 수정하거나, 🔄 버튼으로 AI가 새로 생성합니다")

            for i, ch in enumerate(outline):
                ch_data = chapters.get(ch, {})
                subtopics = ch_data.get('subtopics', [])

                # 챕터 헤더
                st.markdown(f"""
                <div style="background:linear-gradient(90deg, rgba(201,162,75,0.2) 0%, rgba(30,30,30,0.9) 100%);
                            padding:15px 20px;border-radius:10px;margin:20px 0 10px 0;
                            border-left:4px solid var(--gold);">
                    <span style="color:var(--gold);font-weight:bold;font-size:18px;">PART {i+1}</span>
                </div>
                """, unsafe_allow_html=True)

                # 챕터 제목 편집
                new_ch_name = st.text_input(
                    f"챕터 {i+1} 제목",
                    value=ch,
                    key=f"ch_edit_{i}",
                    label_visibility="collapsed"
                )

                # 챕터 이름 변경 적용
                if new_ch_name != ch and new_ch_name.strip():
                    # 목차에서 이름 변경
                    st.session_state['outline'][i] = new_ch_name.strip()
                    # chapters 딕셔너리에서도 키 변경
                    st.session_state['chapters'][new_ch_name.strip()] = st.session_state['chapters'].pop(ch)
                    st.rerun()

                # 소제목들
                for j, sub in enumerate(subtopics):
                    col1, col2 = st.columns([0.5, 5.5])
                    with col1:
                        st.markdown(f"<div style='color:var(--text2);padding-top:8px;'>•</div>", unsafe_allow_html=True)
                    with col2:
                        new_sub = st.text_input(
                            f"소제목 {j+1}",
                            value=sub,
                            key=f"sub_edit_{i}_{j}",
                            label_visibility="collapsed"
                        )
                        # 소제목 변경 적용
                        if new_sub != sub and new_sub.strip():
                            st.session_state['chapters'][ch]['subtopics'][j] = new_sub.strip()
                            # subtopic_data도 업데이트
                            old_data = st.session_state['chapters'][ch]['subtopic_data'].pop(sub, {'questions': [], 'answers': [], 'content': ''})
                            st.session_state['chapters'][ch]['subtopic_data'][new_sub.strip()] = old_data
                            st.rerun()

            st.markdown("---")

        # 하단 버튼
        col1, col2 = st.columns([1, 1])
        with col1:
            if st.button("← 이전", key="interview_prev_6", use_container_width=True):
                st.session_state['interview_step'] = 5
                st.rerun()
        with col2:
            if st.button("✍️ 본문 생성하기", key="generate_body", use_container_width=True, type="primary"):
                progress_box = st.empty()
                interview_data = st.session_state.get('interview_data', st.session_state['temp_interview'])
                success = generate_body_from_outline(interview_data, progress_box)

                if success:
                    import time
                    time.sleep(1)
                    st.session_state['current_page'] = 7  # 최종 출력 페이지로 이동
                    st.rerun()

    st.stop()

# ==========================================
# 여기서부터 기존 페이지 로직 (인터뷰 완료 후)
# ==========================================

# 페이지 네비게이션 (간소화: 4단계)
simple_pages = ["주제", "목차", "본문", "완성"]
page_mapping = [0, 4, 5, 7]  # 실제 페이지 인덱스
current = st.session_state['current_page']

# 현재 페이지가 간소화된 네비게이션의 어디에 해당하는지
def get_simple_index(current_page):
    if current_page <= 0:
        return 0
    elif current_page <= 4:
        return 1
    elif current_page <= 5:
        return 2
    else:
        return 3

simple_current = get_simple_index(current)

# 프리미엄 네비게이션 바 (4단계)
st.markdown('<div class="premium-nav-container">', unsafe_allow_html=True)
cols = st.columns(4)
for i, (col, page) in enumerate(zip(cols, simple_pages)):
    with col:
        if i == simple_current:
            st.markdown(f'<div class="nav-item active">{i+1}. {page}</div>', unsafe_allow_html=True)
        else:
            if st.button(f"{i+1}. {page}", key=f"nav_{i}", use_container_width=True):
                st.session_state['current_page'] = page_mapping[i]
                st.rerun()
st.markdown('</div>', unsafe_allow_html=True)

st.markdown('<div class="section-divider"></div>', unsafe_allow_html=True)

# API 키가 없으면 환영 화면 표시
if not get_api_key():
    st.markdown("""
    <div style="background:linear-gradient(135deg, rgba(201,162,75,0.2) 0%, rgba(30,30,30,0.98) 100%);
                border:0.5px solid rgba(201,162,75,0.4);border-radius:25px;padding:50px 40px;text-align:center;margin:20px 0;">
        <p style="font-size:60px;margin:0 0 20px 0;">👋</p>
        <h2 style="color:var(--gold);font-size:42px;margin-bottom:20px;font-weight:700;">환영합니다!</h2>
        <p style="color:var(--text);font-size:24px;margin-bottom:10px;line-height:1.8;">
            AI가 전자책을 대신 써주는 프로그램입니다
        </p>
        <p style="color:var(--text2);font-size:20px;">
            처음 한 번만 설정하면 바로 사용할 수 있어요
        </p>
    </div>
    """, unsafe_allow_html=True)

    # 큰 안내 박스
    st.markdown("""
    <div style="background:#1a1a2e;border:3px solid #e74c3c;padding:30px;border-radius:20px;margin:30px 0;">
        <p style="font-size:28px;margin:0;line-height:1.6;color:#fff;text-align:center;">
            🔑 <b style="color:#e74c3c;">첫 번째 할 일</b><br><br>
            <span style="font-size:24px;">👈 왼쪽에 <span style="color:#C9A24B;font-weight:700;">"API 키"</span>를 넣어야 해요</span>
        </p>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
    <p style="text-align:center;font-size:32px;color:var(--gold);margin:40px 0 30px 0;font-weight:700;">
        📖 딱 3단계만 하면 끝!
    </p>
    """, unsafe_allow_html=True)

    # STEP 1 - Anthropic 가입
    st.markdown("""
    <div style="background:linear-gradient(135deg, #7c3aed 0%, #5b21b6 100%);padding:30px;border-radius:20px;margin:20px 0;">
        <p style="font-size:28px;margin:0 0 15px 0;color:#fff;font-weight:800;text-align:center;">
            1️⃣ Anthropic 회원가입
        </p>
        <div style="background:rgba(255,255,255,0.15);padding:15px;border-radius:10px;">
            <p style="font-size:16px;margin:0;color:#fff;line-height:1.8;">
                1. 아래 버튼 클릭<br>
                2. <b>"Sign up"</b> 클릭<br>
                3. Google 계정으로 가입 (가장 쉬움)
            </p>
        </div>
    </div>
    """, unsafe_allow_html=True)

    st.link_button("👆 Anthropic 가입하기", "https://console.anthropic.com/", use_container_width=True, type="primary")

    st.markdown("<p style='height:20px;'></p>", unsafe_allow_html=True)

    # STEP 2 - 결제 등록
    st.markdown("""
    <div style="background:linear-gradient(135deg, #e74c3c 0%, #c0392b 100%);padding:30px;border-radius:20px;margin:20px 0;">
        <p style="font-size:28px;margin:0 0 15px 0;color:#fff;font-weight:800;text-align:center;">
            2️⃣ 결제 수단 & 크레딧 충전
        </p>
        <div style="background:rgba(255,255,255,0.15);padding:15px;border-radius:10px;">
            <p style="font-size:16px;margin:0;color:#fff;line-height:1.8;">
                1. 로그인 후 왼쪽 메뉴 <b>"Settings"</b> 클릭<br>
                2. <b>"Billing"</b> 클릭<br>
                3. <b>"Add payment method"</b>로 카드 등록<br>
                4. <b>"Add credits"</b>로 $5~10 충전
            </p>
        </div>
    </div>
    """, unsafe_allow_html=True)

    st.link_button("👆 Billing 페이지 바로가기", "https://console.anthropic.com/settings/billing", use_container_width=True, type="primary")

    st.markdown("<p style='height:20px;'></p>", unsafe_allow_html=True)

    # STEP 3 - 키 받기
    st.markdown("""
    <div style="background:linear-gradient(135deg, #27ae60 0%, #2ecc71 100%);padding:30px;border-radius:20px;margin:20px 0;">
        <p style="font-size:28px;margin:0 0 15px 0;color:#fff;font-weight:800;text-align:center;">
            3️⃣ API 키 발급
        </p>
        <div style="background:rgba(255,255,255,0.15);padding:15px;border-radius:10px;">
            <p style="font-size:16px;margin:0;color:#fff;line-height:1.8;">
                1. 왼쪽 메뉴 <b>"API Keys"</b> 클릭<br>
                2. <b>"Create Key"</b> 버튼 클릭<br>
                3. 이름 입력 (예: ebook) → <b>"Create Key"</b><br>
                4. 생성된 키 <b>복사</b> (sk-ant-api03-... 형식)
            </p>
        </div>
    </div>
    """, unsafe_allow_html=True)

    st.link_button("👆 API Keys 페이지 바로가기", "https://console.anthropic.com/settings/keys", use_container_width=True, type="primary")

    # 마무리
    st.markdown("""
    <div style="background:linear-gradient(135deg, #C9A24B 0%, #A8852F 100%);padding:30px;border-radius:20px;margin:40px 0;text-align:center;">
        <p style="font-size:28px;margin:0 0 10px 0;color:#000;font-weight:800;">
            👈 복사한 키를 왼쪽 사이드바에 붙여넣기
        </p>
        <p style="font-size:16px;margin:0;color:#000;">
            💰 비용: 전자책 1권 약 200~500원 (Claude Sonnet 4 기준)
        </p>
    </div>
    """, unsafe_allow_html=True)

    st.warning("⚠️ API 키는 생성 시 한 번만 보여줍니다. 꼭 복사해두세요!")

    st.markdown("---")

    # 도움말
    st.markdown("""
    <p style="text-align:center;font-size:20px;color:var(--text2);margin:20px 0;">
        😕 어려우시면 유튜브 영상을 보세요
    </p>
    """, unsafe_allow_html=True)

    st.link_button("📺 Claude API 키 발급 방법 (유튜브)", "https://www.youtube.com/results?search_query=anthropic+claude+api+key+발급", use_container_width=True)

    st.markdown("""
    <div style="background:rgba(100,100,100,0.2);padding:20px;border-radius:15px;margin:30px 0;text-align:center;">
        <p style="font-size:18px;margin:0;color:var(--text2);">
            💡 <b>팁:</b> 키는 한 번만 넣으면 저장돼요. 다음부터는 바로 시작!
        </p>
    </div>
    """, unsafe_allow_html=True)

    st.stop()  # API 키 없으면 여기서 멈춤

# ==========================================
# PAGE 0: 주제 & 시장분석
# ==========================================
if current == 0:
    st.markdown("""
    <div class="section-title-box">
        <span class="section-step">STEP 01</span>
        <h2>주제 선정 & 시장 분석</h2>
        <p>AI가 전자책의 성공 가능성을 분석합니다</p>
    </div>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns([1, 1], gap="large")

    with col1:
        st.markdown("### 주제 입력")

        topic = st.text_input("어떤 주제로 전자책을 쓸까요?", value=st.session_state['topic'], placeholder="예: 주식 배당으로 월 100만원", key="p0_topic", label_visibility="collapsed")
        if topic != st.session_state['topic']:
            st.session_state['topic'] = topic
            st.session_state['score_details'] = None

        # 빠른 제작 버튼 (자동 모드)
        st.markdown("""
        <div style="background:linear-gradient(135deg, #C9A24B 0%, #A8852F 100%);padding:20px;border-radius:15px;margin:20px 0;text-align:center;">
            <p style="font-size:14px;margin:0 0 5px 0;color:#000;opacity:0.8;">⚡ 클릭 한 번으로</p>
            <p style="font-size:20px;margin:0;color:#000;font-weight:800;">목차 + 본문 자동 완성</p>
        </div>
        """, unsafe_allow_html=True)

        if st.button("🚀 빠른 제작 시작", use_container_width=True, key="p0_auto", type="primary"):
            if not topic:
                st.error("주제를 먼저 입력해주세요")
            elif not get_api_key():
                st.error("사이드바에서 API 키를 입력해주세요")
            else:
                progress_box = st.empty()
                success = auto_generate_all(topic, progress_box)
                if success:
                    import time
                    time.sleep(1)
                    st.session_state['current_page'] = 5  # 본문 페이지로 이동
                    st.rerun()

        st.markdown("---")
        st.caption("또는 시장 분석부터 단계별로 진행:")

        if st.button("📊 시장 분석 먼저 하기", use_container_width=True, key="p0_analyze"):
            if not topic:
                st.error("주제를 입력해주세요")
            elif not get_api_key():
                st.error("사이드바에서 API 키를 입력해주세요")
            else:
                with st.spinner("AI가 시장을 분석하고 있습니다..."):
                    result = analyze_market_deep(topic)
                    parsed = parse_json(result)
                    if parsed:
                        st.session_state['score_details'] = parsed
                        st.rerun()

    with col2:
        if st.session_state.get('score_details'):
            d = st.session_state['score_details']
            score = d.get('total_score', 0)
            verdict = d.get('verdict', '')
            v_class = "verdict-go" if "추천" in verdict else ("verdict-wait" if "보류" in verdict else "verdict-no")

            st.markdown(f"""
            <div class="score-card">
                <div class="score-number">{score}</div>
                <div style="font-size:14px;color:var(--text-dim);margin-top:8px;">종합 점수</div>
                <div style="margin-top:24px;"><span class="{v_class}">{verdict}</span></div>
            </div>
            """, unsafe_allow_html=True)

            st.markdown(f"""
            <div class="result-card" style="margin-top:20px;">
                <div style="font-size:13px;color:var(--text-dim);margin-bottom:8px;">AI 분석 요약</div>
                <div style="font-size:15px;color:var(--text-bright);line-height:1.7;">{d.get('verdict_reason', '')}</div>
            </div>
            """, unsafe_allow_html=True)

            sd = d.get('search_data', {})
            if sd:
                st.markdown(f"""
                <div class="data-card" style="margin-top:16px;">
                    <b>검색 데이터</b><br><br>
                    <div style="display:grid;grid-template-columns:1fr 1fr;gap:12px;">
                        <div>• 네이버: <b>{sd.get('naver_monthly', 'N/A')}</b></div>
                        <div>• 구글: <b>{sd.get('google_monthly', 'N/A')}</b></div>
                        <div>• 블로그: <b>{sd.get('naver_blog_posts', 'N/A')}</b></div>
                        <div>• 유튜브: <b>{sd.get('youtube_videos', 'N/A')}</b></div>
                    </div>
                </div>
                """, unsafe_allow_html=True)

            ms = d.get('market_size', {})
            comp = d.get('competition', {})

            c1, c2 = st.columns(2)
            with c1:
                st.markdown(f'<div class="stat-box"><div class="stat-value">{ms.get("level", "")}</div><div class="stat-label">시장 규모 ({ms.get("score", 0)}점)</div></div>', unsafe_allow_html=True)
            with c2:
                st.markdown(f'<div class="stat-box"><div class="stat-value">{comp.get("level", "")}</div><div class="stat-label">경쟁 강도 ({comp.get("score", 0)}점)</div></div>', unsafe_allow_html=True)

            if comp.get('your_opportunity'):
                st.success(f"**차별화 기회:** {comp.get('your_opportunity', '')}")

            # 경쟁 도서 검색 - 주제 키워드로 직접 검색
            current_topic = st.session_state.get('topic', '')
            if current_topic:
                st.markdown("""
                <div style="margin-top:35px;">
                    <div style="display:flex;align-items:center;gap:12px;margin-bottom:25px;">
                        <div style="width:50px;height:50px;background:linear-gradient(135deg, var(--gold) 0%, var(--gold-dark) 100%);border-radius:12px;display:flex;align-items:center;justify-content:center;">
                            <span style="font-size:26px;">🔍</span>
                        </div>
                        <div>
                            <h4 style="color:var(--gold);margin:0;font-size:22px;font-weight:600;">경쟁 도서 직접 확인하기</h4>
                            <p style="color:var(--text2);margin:4px 0 0 0;font-size:14px;">각 플랫폼에서 이 주제의 책들을 살펴보세요</p>
                        </div>
                    </div>
                </div>
                """, unsafe_allow_html=True)

                # 플랫폼별 검색 URL 매핑
                platforms = [
                    {
                        'name': '크몽 전자책',
                        'icon': '📘',
                        'url': 'https://kmong.com/search?c=ebook&q=',
                        'desc': '전자책/PDF 마켓',
                        'gradient': 'linear-gradient(135deg, #667eea 0%, #764ba2 100%)'
                    },
                    {
                        'name': '리디북스',
                        'icon': '📗',
                        'url': 'https://ridibooks.com/search?q=',
                        'desc': '국내 최대 전자책',
                        'gradient': 'linear-gradient(135deg, #11998e 0%, #38ef7d 100%)'
                    },
                    {
                        'name': 'YES24',
                        'icon': '📙',
                        'url': 'https://www.yes24.com/Product/Search?domain=BOOK&query=',
                        'desc': '종합 서점',
                        'gradient': 'linear-gradient(135deg, #f093fb 0%, #f5576c 100%)'
                    },
                    {
                        'name': '교보문고',
                        'icon': '📕',
                        'url': 'https://search.kyobobook.co.kr/search?keyword=',
                        'desc': '국내 대표 서점',
                        'gradient': 'linear-gradient(135deg, #4facfe 0%, #00f2fe 100%)'
                    },
                    {
                        'name': '클래스101',
                        'icon': '🎓',
                        'url': 'https://class101.net/search?query=',
                        'desc': '온라인 클래스',
                        'gradient': 'linear-gradient(135deg, #fa709a 0%, #fee140 100%)'
                    },
                    {
                        'name': '탈잉',
                        'icon': '👨‍🏫',
                        'url': 'https://taling.me/search?query=',
                        'desc': '재능 마켓',
                        'gradient': 'linear-gradient(135deg, #a8edea 0%, #fed6e3 100%)'
                    }
                ]

                search_query = urllib.parse.quote(current_topic)

                cols = st.columns(3)
                for idx, platform in enumerate(platforms):
                    with cols[idx % 3]:
                        search_url = platform['url'] + search_query
                        st.markdown(f"""
                        <a href="{search_url}" target="_blank" style="text-decoration:none;display:block;margin-bottom:15px;">
                            <div style="background:rgba(25,25,25,0.9);border:1px solid rgba(201,162,75,0.3);border-radius:16px;overflow:hidden;transition:all 0.3s ease;">
                                <div style="height:80px;background:{platform['gradient']};display:flex;align-items:center;justify-content:center;">
                                    <span style="font-size:40px;">{platform['icon']}</span>
                                </div>
                                <div style="padding:18px;text-align:center;">
                                    <div style="font-size:17px;color:var(--text);font-weight:700;margin-bottom:6px;">
                                        {platform['name']}
                                    </div>
                                    <div style="font-size:13px;color:var(--text2);margin-bottom:12px;">
                                        {platform['desc']}
                                    </div>
                                    <div style="background:linear-gradient(135deg, var(--gold) 0%, var(--gold-dark) 100%);color:var(--dark);padding:10px 16px;border-radius:8px;font-size:13px;font-weight:700;">
                                        🔍 "{current_topic[:15]}{'...' if len(current_topic) > 15 else ''}" 검색
                                    </div>
                                </div>
                            </div>
                        </a>
                        """, unsafe_allow_html=True)
        else:
            st.markdown("""
            <div class="empty-state">
                <div class="empty-state-text">
                    주제를 입력하고 <b>AI 시장 분석</b>을 시작하세요<br>
                    검색량, 경쟁 강도, 수익 가능성을 분석합니다
                </div>
            </div>
            """, unsafe_allow_html=True)

    st.markdown('<div class="next-section"></div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1, 2, 1])
    with c2:
        if st.button("다음 단계로 타겟 설정", key="p0_next", use_container_width=True):
            go_next()
            st.rerun()


# ==========================================
# PAGE 1: 타겟 & 컨셉
# ==========================================
elif current == 1:
    st.markdown("""
    <div class="section-title-box">
        <span class="section-step">STEP 02</span>
        <h2>타겟 설정 & 제목 생성</h2>
        <p>구매할 사람을 정하고 끌리는 제목을 만듭니다</p>
    </div>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns([1, 1])

    with col1:
        st.markdown("### 타겟 선정")

        if st.button("AI 타겟 추천", key="p1_target"):
            if st.session_state['topic'] and get_api_key():
                with st.spinner("분석 중..."):
                    result = suggest_targets(st.session_state['topic'])
                    parsed = parse_json(result)
                    if parsed:
                        st.session_state['suggested_targets'] = parsed
                        st.rerun()

        if st.session_state.get('suggested_targets'):
            personas = st.session_state['suggested_targets'].get('personas', [])[:3]

            for i, p in enumerate(personas):
                target_name = p.get('name', '')
                target_demo = p.get('demographics', '')
                target_needs = p.get('needs', '')
                target_pains = p.get('pain_points', [])

                st.markdown(f"""<div class="data-card">
                    <b>{html.escape(str(target_name))}</b><br>
                    <small>{html.escape(str(target_demo))}</small><br>
                    <small style="color:var(--gold);">{html.escape(str(target_needs))}</small>
                </div>""", unsafe_allow_html=True)

                if st.button(f"이 타겟 선택", key=f"sel_target_{i}", use_container_width=True):
                    selected_target = f"{target_name} - {target_demo}"
                    st.session_state['target_persona'] = selected_target
                    st.session_state['p1_persona'] = selected_target
                    st.session_state['pain_points'] = ", ".join(target_pains[:5])
                    st.session_state['suggested_targets'] = None
                    st.rerun()

        st.markdown("---")
        st.markdown("### 선택된 타겟")
        persona = st.text_area("타겟:", value=st.session_state.get('target_persona', ''), height=60, key="p1_persona", placeholder="AI 추천에서 선택하거나 직접 입력")
        st.session_state['target_persona'] = persona

        if st.button("고민 심층 분석", key="p1_analyze", use_container_width=True):
            if not persona:
                st.error("타겟을 먼저 입력해주세요")
            elif not get_api_key():
                st.error("API 키를 입력해주세요")
            else:
                with st.spinner("심층 분석 중..."):
                    r = analyze_pains_deep(st.session_state['topic'], persona)
                    parsed = parse_json(r)
                    if parsed:
                        st.session_state['analyzed_pains'] = parsed
                        surface = parsed.get('surface_pains', {}).get('pains', [])
                        hidden = parsed.get('hidden_pains', {}).get('pains', [])
                        st.session_state['pain_points'] = ", ".join((surface + hidden)[:6])
                        st.rerun()
                    else:
                        st.error("분석 실패. 다시 시도해주세요")

        if st.session_state.get('analyzed_pains'):
            p = st.session_state['analyzed_pains']
            st.markdown("**표면적 고민**")
            for pain in p.get('surface_pains', {}).get('pains', []):
                st.write(f"• {pain}")
            st.markdown("**숨겨진 진짜 고민**")
            for pain in p.get('hidden_pains', {}).get('pains', []):
                st.write(f"• {pain}")
            if p.get('marketing_hook'):
                st.info(f"**마케팅 훅:** {p.get('marketing_hook', '')}")

    with col2:
        st.markdown("### 베스트셀러급 제목 생성")

        # 선택된 제목이 있으면 상단에 확정 표시
        if st.session_state.get('book_title'):
            st.markdown(f"""
            <div style="background:linear-gradient(135deg,#10b981,#059669);padding:16px 20px;border-radius:12px;margin-bottom:20px;">
                <div style="color:white;font-size:12px;margin-bottom:6px;">✓ 확정된 제목</div>
                <div style="color:white;font-size:20px;font-weight:700;">{html.escape(st.session_state.get('book_title', ''))}</div>
                <div style="color:rgba(255,255,255,0.85);font-size:14px;margin-top:4px;">{html.escape(st.session_state.get('subtitle', ''))}</div>
            </div>
            """, unsafe_allow_html=True)

        pain_points = st.text_area("독자의 고민:", value=st.session_state['pain_points'], height=60, key="p1_pains")
        st.session_state['pain_points'] = pain_points

        if st.button("베스트셀러 제목 생성", key="p1_title"):
            if st.session_state['topic']:
                with st.spinner("베스트셀러 패턴 분석 중..."):
                    r = generate_titles_bestseller(st.session_state['topic'], st.session_state['target_persona'], st.session_state['pain_points'])
                    parsed = parse_json(r)
                    if parsed:
                        st.session_state['generated_titles'] = parsed
                        st.rerun()

        if st.session_state.get('generated_titles'):
            titles_list = st.session_state['generated_titles'].get('titles', [])[:5]
            for i, t in enumerate(titles_list):
                title_val = t.get('title', '')
                subtitle_val = t.get('subtitle', '')
                concept_val = t.get('concept', '')

                st.markdown(f"""
                <div class="title-card">
                    <div class="title-main">{html.escape(title_val)}</div>
                    <div class="title-sub">{html.escape(subtitle_val)}</div>
                    <div style="font-size:11px;color:var(--gold);margin-top:12px;letter-spacing:2px;">{html.escape(concept_val)}</div>
                </div>
                """, unsafe_allow_html=True)

                if st.button(f"✓ 이 제목으로 확정", key=f"sel_title_{i}", use_container_width=True):
                    st.session_state['book_title'] = title_val
                    st.session_state['subtitle'] = subtitle_val
                    st.toast(f"'{title_val}' 제목이 확정되었습니다!")
                    st.rerun()

        # 직접 입력 옵션
        st.markdown("---")
        st.markdown("#### 또는 직접 입력")
        manual_title = st.text_input("제목 입력", key="manual_title_v3")
        manual_subtitle = st.text_input("부제 입력", key="manual_subtitle_v3")
        if st.button("✓ 직접 입력한 제목으로 확정", key="manual_confirm_v3", use_container_width=True):
            if manual_title:
                st.session_state['book_title'] = manual_title
                st.session_state['subtitle'] = manual_subtitle if manual_subtitle else ''
                st.toast(f"'{manual_title}' 제목이 확정되었습니다!")
                st.rerun()

    st.markdown('<div class="next-section"></div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1, 1, 1])
    with c1:
        if st.button("이전", key="p1_prev", use_container_width=True):
            go_prev()
            st.rerun()
    with c3:
        if st.button("다음 경쟁분석", key="p1_next", use_container_width=True):
            go_next()
            st.rerun()


# ==========================================
# PAGE 2: 경쟁도서 분석
# ==========================================
elif current == 2:
    st.markdown("""
    <div class="section-title-box">
        <span class="section-step">STEP 03</span>
        <h2>경쟁 도서 분석</h2>
        <p>기존 도서의 부정 리뷰를 분석해서 숨은 니즈를 찾습니다</p>
    </div>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns([1, 1])

    with col1:
        st.markdown("### 부정 리뷰 분석")

        if st.button("경쟁 도서 분석하기", use_container_width=True, key="p2_analyze"):
            if not st.session_state['topic']:
                st.error("주제를 먼저 입력해주세요")
            elif not get_api_key():
                st.error("API 키를 입력해주세요")
            else:
                with st.spinner("경쟁 도서 분석 중..."):
                    result = analyze_competitor_reviews(st.session_state['topic'])
                    parsed = parse_json(result)
                    if parsed:
                        st.session_state['review_analysis'] = parsed
                        concepts = parsed.get('concept_suggestions', [])
                        st.session_state['market_gaps'] = [c.get('concept', '') for c in concepts]
                        st.rerun()

        if st.session_state.get('review_analysis'):
            a = st.session_state['review_analysis']
            scope = a.get('analysis_scope', {})
            col_s1, col_s2 = st.columns(2)
            with col_s1:
                st.markdown(f'<div class="stat-box"><div class="stat-value">{scope.get("books_analyzed", "N/A")}</div><div class="stat-label">분석 도서</div></div>', unsafe_allow_html=True)
            with col_s2:
                st.markdown(f'<div class="stat-box"><div class="stat-value">{scope.get("negative_reviews", "N/A")}</div><div class="stat-label">부정 리뷰</div></div>', unsafe_allow_html=True)

    with col2:
        st.markdown("### 분석 결과")

        if st.session_state.get('review_analysis'):
            a = st.session_state['review_analysis']

            patterns = a.get('negative_patterns', [])
            if patterns:
                st.markdown("#### 독자 불만 패턴")
                for i, p in enumerate(patterns[:3], 1):
                    st.markdown(f"""<div class="data-card">
                        <b>{i}. {p.get('pattern', '')} ({p.get('frequency', '')})</b>
                    </div>""", unsafe_allow_html=True)
                    for rev in p.get('example_reviews', []):
                        st.caption(f'"{rev}"')
                    st.info(f"**숨겨진 니즈:** {p.get('hidden_need', '')}")
                    st.success(f"**해결책:** {p.get('solution', '')}")

            concepts = a.get('concept_suggestions', [])
            if concepts:
                st.markdown("#### 차별화 컨셉")
                for c in concepts[:2]:
                    st.markdown(f"""
                    <div class="info-card">
                        <b>「{html.escape(c.get('concept', ''))}」</b><br>
                        <span style="color:rgba(255,255,255,0.7);">{html.escape(c.get('why_works', ''))}</span>
                    </div>
                    """, unsafe_allow_html=True)
        else:
            st.markdown('<div style="text-align:center;padding:60px;background:rgba(255,255,255,0.03);border-radius:16px;border:1px solid rgba(201,162,75,0.15);"><p style="color:rgba(255,255,255,0.5);">분석 버튼을 눌러주세요</p></div>', unsafe_allow_html=True)

    st.markdown('<div class="next-section"></div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1, 1, 1])
    with c1:
        if st.button("이전", key="p2_prev", use_container_width=True):
            go_prev()
            st.rerun()
    with c3:
        if st.button("다음 학습", key="p2_next", use_container_width=True):
            go_next()
            st.rerun()


# ==========================================
# PAGE 3: 학습 & 리서치
# ==========================================
elif current == 3:
    st.markdown("""
    <div class="section-title-box">
        <span class="section-step">STEP 04</span>
        <h2>학습 & 리서치</h2>
        <p>베스트셀러 분석, 트렌드 파악, 핵심 인사이트를 수집합니다</p>
    </div>
    """, unsafe_allow_html=True)

    # 탭으로 구분
    tab1, tab2, tab3 = st.tabs(["레퍼런스 추천", "트렌드 분석", "경쟁서 분석"])

    # ========== 탭1: 레퍼런스 추천 & 아이디어 ==========
    with tab1:
        topic = st.session_state.get('topic', '')

        col1, col2 = st.columns([1, 1])

        with col1:
            st.markdown("### 레퍼런스 자동 추천")
            st.markdown('<p style="color:var(--text2);font-size:13px;">주제에 맞는 참고 자료를 AI가 자동으로 추천합니다</p>', unsafe_allow_html=True)

            if not topic:
                st.warning("먼저 시장분석 페이지에서 주제를 입력해주세요")
            else:
                st.markdown(f'<p style="color:var(--accent);font-size:14px;margin:10px 0;">현재 주제: <b>{html.escape(topic)}</b></p>', unsafe_allow_html=True)

                ref_category = st.selectbox("추천 카테고리", ["베스트셀러 도서", "핵심 개념/이론", "성공 사례", "전문가 인사이트"], key="ref_cat")

                if st.button("레퍼런스 추천받기", use_container_width=True, key="auto_ref_btn"):
                    if not get_api_key():
                        st.error("API 키를 입력해주세요")
                    else:
                        with st.spinner("관련 레퍼런스 심층 분석 중..."):
                            prompt = f"""'{topic}' 주제로 전자책을 쓰려고 합니다.
'{ref_category}' 카테고리에서 참고할 만한 자료 3개를 추천해주세요.

중요: 마치 이 책/자료를 직접 읽은 것처럼 아주 상세하게 설명해주세요.

각 추천 자료에 대해 다음을 포함해주세요:
1. 제목과 저자
2. 책/자료의 핵심 메시지 (10문장 이상으로 상세히)
3. 주요 챕터/섹션별 핵심 내용
4. 저자의 핵심 주장과 근거
5. 실제 사례나 스토리
6. 전자책에 활용할 수 있는 구체적 인사이트

중요: 책의 모든 주요 챕터를 빠짐없이 요약해주세요. 일부만 하지 말고 전체 목차를 다 포함해주세요.

JSON 형식으로 응답:
{{
    "recommendations": [
        {{
            "title": "자료 제목",
            "author": "저자/출처",
            "core_message": "이 책의 핵심 메시지와 주장을 10문장 이상으로 상세하게 설명",
            "chapters": [
                {{"name": "1장 제목", "summary": "이 챕터의 핵심 내용 2-3문장"}},
                {{"name": "2장 제목", "summary": "이 챕터의 핵심 내용 2-3문장"}},
                {{"name": "3장 제목", "summary": "이 챕터의 핵심 내용 2-3문장"}},
                {{"name": "4장 제목", "summary": "이 챕터의 핵심 내용 2-3문장"}},
                {{"name": "5장 제목", "summary": "이 챕터의 핵심 내용 2-3문장"}},
                {{"name": "6장 제목", "summary": "이 챕터의 핵심 내용 2-3문장"}},
                {{"name": "7장 제목", "summary": "이 챕터의 핵심 내용 2-3문장"}},
                {{"name": "8장 제목", "summary": "이 챕터의 핵심 내용 2-3문장"}}
            ],
            "key_arguments": ["저자의 핵심 주장1과 근거", "핵심 주장2와 근거", "핵심 주장3과 근거"],
            "real_examples": ["책에 나온 실제 사례/스토리 1", "사례 2", "사례 3"],
            "key_insights": ["전자책에 활용할 인사이트 1", "인사이트 2", "인사이트 3", "인사이트 4", "인사이트 5"],
            "application": "내 전자책에 구체적으로 활용하는 방법 (3문장 이상)"
        }}
    ]
}}"""
                            result = ask_ai(prompt, 0.8)
                            parsed = parse_json(result)
                            if parsed and parsed.get('recommendations'):
                                st.session_state['recommended_refs'] = parsed['recommendations']
                                st.rerun()
                            else:
                                st.error("추천 생성 실패")

                # 추천된 레퍼런스 표시
                if st.session_state.get('recommended_refs'):
                    st.markdown("---")
                    st.markdown("#### 추천 레퍼런스")

                    for i, ref in enumerate(st.session_state['recommended_refs']):
                        st.markdown(f"""<div class="data-card">
                            <b>{html.escape(str(ref.get('title', '')))}</b>
                            <br><small style="color:var(--text2);">{html.escape(str(ref.get('author', '')))}</small>
                        </div>""", unsafe_allow_html=True)

                        # 핵심 메시지
                        if ref.get('core_message'):
                            st.markdown("**핵심 메시지**")
                            st.markdown(f'<p style="color:#e0e0e0;line-height:1.8;">{html.escape(str(ref.get("core_message", "")))}</p>', unsafe_allow_html=True)

                        # 챕터 요약
                        if ref.get('chapters'):
                            st.markdown("**챕터별 요약**")
                            for ch in ref.get('chapters', []):
                                st.markdown(f"""<div style="background:rgba(255,255,255,0.05);padding:12px 16px;margin:8px 0;border-left:3px solid var(--gold);">
                                    <b style="color:var(--gold);">{html.escape(str(ch.get('name', '')))}</b><br>
                                    <span style="color:#e0e0e0;">{html.escape(str(ch.get('summary', '')))}</span>
                                </div>""", unsafe_allow_html=True)

                        # 핵심 주장
                        if ref.get('key_arguments'):
                            st.markdown("**저자의 핵심 주장**")
                            for arg in ref.get('key_arguments', []):
                                st.info(arg)

                        # 실제 사례
                        if ref.get('real_examples'):
                            st.markdown("**실제 사례**")
                            for ex in ref.get('real_examples', []):
                                st.markdown(f'<p style="color:#e0e0e0;">- {html.escape(str(ex))}</p>', unsafe_allow_html=True)

                        # 활용 인사이트
                        if ref.get('key_insights'):
                            st.markdown("**활용 인사이트**")
                            for insight in ref.get('key_insights', []):
                                st.success(insight)

                        # 적용 방법
                        if ref.get('application'):
                            st.markdown("**내 책에 적용하는 방법**")
                            st.markdown(f'<p style="color:#e0e0e0;line-height:1.8;">{html.escape(str(ref.get("application", "")))}</p>', unsafe_allow_html=True)

                        if st.button("이 자료 저장하기", key=f"save_ref_{i}", use_container_width=True):
                            ref_item = {
                                'title': ref.get('title', ''),
                                'source': ref.get('author', ''),
                                'core_message': ref.get('core_message', ''),
                                'chapters': ref.get('chapters', []),
                                'key_arguments': ref.get('key_arguments', []),
                                'real_examples': ref.get('real_examples', []),
                                'key_insights': ref.get('key_insights', []),
                                'application': ref.get('application', ''),
                                'type': 'recommended',
                                'added_at': datetime.now().strftime('%Y-%m-%d %H:%M')
                            }
                            st.session_state['knowledge_hub'].append(ref_item)
                            st.success("저장 완료")
                            st.rerun()

                        st.markdown("---")

        with col2:
            st.markdown("### 저장된 자료 & 아이디어 도출")
            hub = st.session_state.get('knowledge_hub', [])

            if hub:
                st.caption(f"총 {len(hub)}개 자료 저장됨")

                for i, item in enumerate(hub):
                    title = item.get('title', item.get('main_topic', item.get('source', f'자료 {i+1}')))

                    st.markdown(f"""<div class="data-card">
                        <b>{html.escape(str(title))}</b><br>
                        <small>{html.escape(str(item.get('source', '')))} | {item.get('added_at', '')}</small>
                    </div>""", unsafe_allow_html=True)

                    # 핵심 메시지 전체 표시
                    if item.get('core_message'):
                        st.write(item['core_message'])

                    # 핵심 인사이트 표시
                    if item.get('key_insights'):
                        st.markdown("**핵심 인사이트:**")
                        for insight in item.get('key_insights', [])[:3]:
                            st.success(insight)

                    # 적용 방법 표시
                    if item.get('application'):
                        st.info(f"적용법: {item['application']}")

                    col_a, col_b = st.columns([1, 1])
                    with col_a:
                        if st.button("상세보기", key=f"view_ref_{i}"):
                            st.session_state[f'show_detail_{i}'] = not st.session_state.get(f'show_detail_{i}', False)
                            st.rerun()
                    with col_b:
                        if st.button("삭제", key=f"del_ref_{i}"):
                            st.session_state['knowledge_hub'].pop(i)
                            st.rerun()

                    # 상세 보기 토글
                    if st.session_state.get(f'show_detail_{i}', False):
                        if item.get('chapters'):
                            st.markdown("**챕터 요약:**")
                            for ch in item.get('chapters', []):
                                st.markdown(f"""<div style="background:rgba(255,255,255,0.05);padding:10px 14px;margin:6px 0;border-left:2px solid var(--gold);">
                                    <b style="color:var(--gold);">{html.escape(str(ch.get('name', '')))}</b><br>
                                    <span style="color:#e0e0e0;font-size:14px;">{html.escape(str(ch.get('summary', '')))}</span>
                                </div>""", unsafe_allow_html=True)
                        if item.get('key_arguments'):
                            st.markdown("**핵심 주장:**")
                            for arg in item.get('key_arguments', []):
                                st.markdown(f'<p style="color:#e0e0e0;">- {html.escape(str(arg))}</p>', unsafe_allow_html=True)
                        if item.get('real_examples'):
                            st.markdown("**실제 사례:**")
                            for ex in item.get('real_examples', []):
                                st.markdown(f'<p style="color:#e0e0e0;">- {html.escape(str(ex))}</p>', unsafe_allow_html=True)

                    st.markdown("---")

                st.markdown("---")
                st.markdown("#### 아이디어 도출")
                st.markdown('<p style="color:var(--text2);font-size:13px;">수집된 자료를 바탕으로 전자책 아이디어를 생성합니다</p>', unsafe_allow_html=True)

                if st.button("아이디어 생성하기", use_container_width=True, key="ideate_btn"):
                    if not get_api_key():
                        st.error("API 키를 입력해주세요")
                    else:
                        with st.spinner("아이디어 생성 중..."):
                            hub_summary = ""
                            for item in hub[:5]:
                                hub_summary += f"\n[{item.get('title', '')}]\n"
                                if item.get('core_message'):
                                    hub_summary += f"핵심 메시지: {item.get('core_message', '')}\n"
                                if item.get('key_insights'):
                                    hub_summary += f"인사이트: {', '.join(item.get('key_insights', []))}\n"
                                if item.get('application'):
                                    hub_summary += f"적용법: {item.get('application', '')}\n"

                            prompt = f"""다음 수집된 자료들을 철저히 분석하여 '{topic}' 주제의 전자책 아이디어를 도출해주세요:

수집된 자료:
{hub_summary}

위 자료들의 공통점, 차이점, 빈틈을 분석하고 다음을 포함해서 아이디어를 생성해주세요:
1. 기존 책들과 확실히 다른 차별화된 콘셉트
2. 독자의 문제를 해결하는 독특한 관점
3. 구체적인 목차/콘텐츠 구성 아이디어
4. 타겟 독자에게 강하게 어필할 포인트

JSON 형식으로 응답:
{{
    "main_concept": "핵심 콘셉트 한 문장 (경쟁작과 어떻게 다른지 명확히)",
    "unique_angles": ["독특한 관점 1 (왜 이 관점이 효과적인지 설명)", "관점 2", "관점 3"],
    "content_ideas": ["챕터 아이디어 1", "챕터 아이디어 2", "챕터 아이디어 3", "챕터 아이디어 4", "챕터 아이디어 5"],
    "appeal_points": ["어필 포인트 1", "포인트 2", "포인트 3"],
    "title_suggestions": ["제목 제안 1 (부제 포함)", "제목 제안 2 (부제 포함)", "제목 제안 3 (부제 포함)"],
    "differentiation": "경쟁작 대비 구체적인 차별화 전략 (3문장 이상)"
}}"""
                            result = ask_ai(prompt, 0.9)
                            parsed = parse_json(result)
                            if parsed:
                                st.session_state['generated_ideas'] = parsed
                                st.rerun()
                            else:
                                st.error("아이디어 생성 실패")

                # 생성된 아이디어 표시
                if st.session_state.get('generated_ideas'):
                    ideas = st.session_state['generated_ideas']

                    st.markdown(f"""<div class="summary-hub">
                        <b>핵심 콘셉트</b><br>
                        {html.escape(str(ideas.get('main_concept', '')))}
                    </div>""", unsafe_allow_html=True)

                    if ideas.get('unique_angles'):
                        st.markdown("**독특한 관점**")
                        for angle in ideas.get('unique_angles', []):
                            st.info(angle)

                    if ideas.get('title_suggestions'):
                        st.markdown("**제목 제안**")
                        for title in ideas.get('title_suggestions', []):
                            st.success(title)

                    if ideas.get('content_ideas'):
                        st.markdown("**콘텐츠 아이디어**")
                        for idea in ideas.get('content_ideas', []):
                            st.write(f"- {idea}")

                    if ideas.get('differentiation'):
                        st.markdown(f"""<div class="data-card">
                            <b>차별화 전략</b><br>
                            <small>{html.escape(str(ideas.get('differentiation', '')))}</small>
                        </div>""", unsafe_allow_html=True)
            else:
                st.markdown('<div style="text-align:center;padding:60px 20px;background:var(--card);border:1px solid var(--line);"><p style="color:var(--text2);">추천받은 레퍼런스를 저장하면<br>아이디어를 도출할 수 있습니다</p></div>', unsafe_allow_html=True)

    # ========== 탭2: 트렌드 분석 ==========
    with tab2:
        st.markdown("### 시장 트렌드 분석")
        st.markdown('<p style="color:var(--text2);">현재 인기 있는 전자책 트렌드와 키워드를 파악합니다</p>', unsafe_allow_html=True)

        col1, col2 = st.columns([1, 1])

        with col1:
            st.markdown("#### 트렌드 키워드 분석")
            trend_topic = st.text_input("분석할 분야", key="trend_topic", placeholder="예: 재테크, 자기계발, 다이어트...")

            if st.button("트렌드 분석", use_container_width=True, key="trend_btn"):
                if not trend_topic:
                    st.error("분야를 입력하세요")
                elif not get_api_key():
                    st.error("API 키를 입력해주세요")
                else:
                    with st.spinner("트렌드 분석 중..."):
                        prompt = f"""'{trend_topic}' 분야의 전자책 시장 트렌드를 분석해주세요.

JSON 형식으로 응답:
{{
    "hot_keywords": ["인기 키워드 1", "키워드 2", "키워드 3", "키워드 4", "키워드 5"],
    "rising_topics": ["떠오르는 주제 1", "주제 2", "주제 3"],
    "reader_needs": ["독자가 원하는 것 1", "원하는 것 2", "원하는 것 3"],
    "content_gaps": ["시장에서 부족한 콘텐츠 1", "부족한 콘텐츠 2"],
    "recommended_angles": ["추천 접근 방식 1", "접근 방식 2", "접근 방식 3"],
    "avoid": ["피해야 할 것 1", "피해야 할 것 2"]
}}"""
                        result = ask_ai(prompt, 0.8)
                        parsed = parse_json(result)
                        if parsed:
                            st.session_state['trend_analysis'] = parsed
                            st.rerun()

            if st.session_state.get('trend_analysis'):
                ta = st.session_state['trend_analysis']
                if ta.get('hot_keywords'):
                    st.write("**인기 키워드**")
                    st.write(" | ".join(ta.get('hot_keywords', [])))
                if ta.get('rising_topics'):
                    st.write("**떠오르는 주제**")
                    for t in ta.get('rising_topics', []):
                        st.write(f"- {t}")

        with col2:
            st.markdown("#### 독자 니즈")
            if st.session_state.get('trend_analysis'):
                ta = st.session_state['trend_analysis']
                if ta.get('reader_needs'):
                    st.write("**독자가 원하는 것**")
                    for n in ta.get('reader_needs', []):
                        st.info(n)
                if ta.get('content_gaps'):
                    st.write("**시장 빈틈**")
                    for g in ta.get('content_gaps', []):
                        st.success(g)
                if ta.get('recommended_angles'):
                    st.write("**추천 접근법**")
                    for r in ta.get('recommended_angles', []):
                        st.write(f"- {r}")
            else:
                st.markdown('<div style="text-align:center;padding:60px;background:var(--card);border:1px solid var(--line);"><p style="color:var(--text2);">분야를 입력하고<br>트렌드 분석을 시작하세요</p></div>', unsafe_allow_html=True)

    # ========== 탭3: 경쟁서 분석 ==========
    with tab3:
        st.markdown("### 경쟁 도서 분석")
        st.markdown('<p style="color:var(--text2);">경쟁 전자책의 목차, 리뷰, 강점을 분석합니다</p>', unsafe_allow_html=True)

        col1, col2 = st.columns([1, 1])

        with col1:
            st.markdown("#### 경쟁서 정보 입력")
            comp_title = st.text_input("책 제목", key="comp_title", placeholder="예: 돈의 심리학")
            comp_toc = st.text_area("목차 (복사/붙여넣기)", height=150, key="comp_toc", placeholder="1장. 제목\n2장. 제목\n...")
            comp_reviews = st.text_area("대표 리뷰 (선택)", height=100, key="comp_reviews", placeholder="인상적인 리뷰를 붙여넣으세요...")

            if st.button("경쟁서 분석", use_container_width=True, key="comp_btn"):
                if not comp_title or not comp_toc:
                    st.error("제목과 목차를 입력하세요")
                elif not get_api_key():
                    st.error("API 키를 입력해주세요")
                else:
                    with st.spinner("분석 중..."):
                        prompt = f"""다음 경쟁 도서를 분석해주세요:

제목: {comp_title}
목차:
{comp_toc}

리뷰: {comp_reviews if comp_reviews else '없음'}

JSON 형식으로 응답:
{{
    "book_summary": "이 책의 핵심 콘셉트",
    "target_audience": "예상 타겟 독자",
    "strengths": ["강점 1", "강점 2", "강점 3"],
    "weaknesses": ["약점/빈틈 1", "약점 2"],
    "unique_selling_point": "이 책만의 차별점",
    "improvement_opportunities": ["내 책에서 더 잘할 수 있는 것 1", "기회 2", "기회 3"],
    "key_chapters": ["핵심 챕터 1", "챕터 2"],
    "content_structure": "콘텐츠 구성 방식"
}}"""
                        result = ask_ai(prompt, 0.7)
                        parsed = parse_json(result)
                        if parsed:
                            parsed['title'] = comp_title
                            parsed['added_at'] = datetime.now().strftime('%Y-%m-%d %H:%M')
                            if 'competitor_analysis' not in st.session_state:
                                st.session_state['competitor_analysis'] = []
                            st.session_state['competitor_analysis'].append(parsed)
                            st.success("분석 완료")
                            st.rerun()

        with col2:
            st.markdown("#### 분석 결과")
            comps = st.session_state.get('competitor_analysis', [])

            if comps:
                for i, comp in enumerate(comps):
                    st.markdown(f"""<div class="data-card">
                        <b>{html.escape(str(comp.get('title', f'경쟁서 {i+1}')))}</b>
                        <br><small>{comp.get('added_at', '')}</small>
                    </div>""", unsafe_allow_html=True)
                    st.caption(comp.get('book_summary', ''))

                    if comp.get('improvement_opportunities'):
                        for o in comp.get('improvement_opportunities', [])[:2]:
                            st.success(f"차별화: {o}")

                    if st.button("삭제", key=f"del_comp_{i}"):
                        st.session_state['competitor_analysis'].pop(i)
                        st.rerun()
            else:
                st.markdown('<div style="text-align:center;padding:60px;background:var(--card);border:1px solid var(--line);"><p style="color:var(--text2);">경쟁 도서 정보를 입력하고<br>분석해보세요</p></div>', unsafe_allow_html=True)

    st.markdown('<div class="next-section"></div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1, 1, 1])
    with c1:
        if st.button("이전", key="p3_prev", use_container_width=True):
            go_prev()
            st.rerun()
    with c3:
        if st.button("다음", key="p3_next", use_container_width=True):
            go_next()
            st.rerun()


# ==========================================
# PAGE 4: 목차 설계
# ==========================================
elif current == 4:
    st.markdown("""
    <div class="section-title-box">
        <span class="section-step">STEP 05</span>
        <h2>목차 설계</h2>
        <p>독자의 호기심을 자극하는 목차를 만듭니다</p>
    </div>
    """, unsafe_allow_html=True)

    if st.session_state.get('market_gaps'):
        st.success(f"{len(st.session_state['market_gaps'])}개 차별화 포인트 반영")

    col1, col2 = st.columns([1, 1])

    with col1:
        st.markdown("### 목차 생성")

        st.markdown("""
        <div class="info-card">
            <b>🔥 목차 작성 팁</b><br><br>
            • 설명하지 말고 <b>궁금하게</b><br>
            • 구체적 <b>숫자 + 결과</b> 보여주기<br>
            • <b>실패담/고백</b>으로 공감 얻기<br>
            • "99%가 모르는" <b>비밀</b> 암시<br>
            • <b>반전</b>이 있을 것 같은 느낌<br><br>
            <span style="color:var(--gold);">❌ "시간관리의 중요성"</span><br>
            <span style="color:#50c878;">✓ "20대에 이걸 몰라서 5년 날렸다"</span>
        </div>
        """, unsafe_allow_html=True)

        if st.button("목차 생성하기", use_container_width=True, key="p4_outline_btn"):
            if not st.session_state.get('topic'):
                st.error("주제를 입력하세요")
            elif not get_api_key():
                st.error("API 키를 입력해주세요")
            else:
                with st.spinner("목차 생성 중..."):
                    result = generate_outline(
                        st.session_state['topic'],
                        st.session_state.get('target_persona', ''),
                        st.session_state.get('pain_points', ''),
                        st.session_state.get('market_gaps', [])
                    )

                    if result:
                        lines = result.split('\n')
                        chapters = []
                        current_ch = None
                        subtopics = {}

                        for line in lines:
                            orig_line = line
                            line = line.strip()
                            if not line:
                                continue

                            # 마크다운 정리 (먼저 정리한 후 검사)
                            clean_line = re.sub(r'^[#\*\s]+', '', line).strip()
                            clean_line = clean_line.replace('**', '').replace('*', '').strip()

                            # PART 또는 챕터 형식 인식 (더 유연하게)
                            is_chapter = False

                            # PART 형식 (다양한 변형)
                            if re.search(r'PART\s*\d+', clean_line, re.IGNORECASE):
                                is_chapter = True
                            # 파트 형식 (한글)
                            elif re.search(r'파트\s*\d+', clean_line):
                                is_chapter = True
                            # Chapter 형식
                            elif re.search(r'(Chapter|챕터)\s*\d+', clean_line, re.IGNORECASE):
                                is_chapter = True
                            # "1. 제목" 형식
                            elif re.match(r'^\d+[\.\)]\s*.+', clean_line) and not clean_line.startswith('-'):
                                is_chapter = True
                            # 숫자로 시작하는 제목 (예: "1 첫번째 파트")
                            elif re.match(r'^\d+\s+[가-힣A-Za-z]', clean_line):
                                is_chapter = True

                            if is_chapter:
                                name = clean_line
                                if name and len(name) > 3:
                                    current_ch = name
                                    chapters.append(current_ch)
                                    subtopics[current_ch] = []

                            # 소제목 - 다양한 형식 지원
                            elif current_ch:
                                is_subtopic = False
                                st_name = ""

                                # "-" 또는 "•" 또는 "·" 로 시작
                                if re.match(r'^\s*[\-\•\·]\s*', line):
                                    is_subtopic = True
                                    st_name = re.sub(r'^[\s\-\•\·]+', '', line).strip()
                                # 들여쓰기 된 내용
                                elif orig_line.startswith('  ') or orig_line.startswith('\t'):
                                    if not any(x in line.upper() for x in ['PART', 'CHAPTER', '파트']):
                                        is_subtopic = True
                                        st_name = line.strip().lstrip('-•· ')
                                # "  1)" 또는 "  a)" 형식
                                elif re.match(r'^\s+[\da-z][\)\.]', orig_line):
                                    is_subtopic = True
                                    st_name = re.sub(r'^[\s\da-z\)\.\-]+', '', line).strip()

                                if is_subtopic:
                                    st_name = st_name.replace('**', '').replace('*', '').replace('#', '').strip()
                                    # 소제목이 충분히 길고 유효한 경우만 추가
                                    if st_name and len(st_name) > 3 and not re.match(r'^(PART|파트|Chapter|챕터)', st_name, re.IGNORECASE):
                                        subtopics[current_ch].append(st_name)

                        if chapters:
                            st.session_state['outline'] = chapters
                            st.session_state['chapters'] = {}
                            for ch in chapters:
                                st.session_state['chapters'][ch] = {
                                    'subtopics': subtopics.get(ch, []),
                                    'subtopic_data': {s: {'questions': [], 'answers': [], 'content': ''} for s in subtopics.get(ch, [])}
                                }
                            st.success(f"{len(chapters)}개 챕터 생성!")
                            st.rerun()
                        else:
                            st.error("목차 생성 실패. 다시 시도해주세요.")
                    else:
                        st.error("AI 응답 없음. 다시 시도해주세요.")

    with col2:
        st.markdown("### 현재 목차")

        if st.session_state.get('outline'):
            # 수정 모드 토글
            if 'edit_outline_mode' not in st.session_state:
                st.session_state['edit_outline_mode'] = False

            col_view, col_edit = st.columns([1, 1])
            with col_view:
                if st.button("👁 보기 모드", use_container_width=True, disabled=not st.session_state['edit_outline_mode']):
                    st.session_state['edit_outline_mode'] = False
                    st.rerun()
            with col_edit:
                if st.button("✏️ 수정 모드", use_container_width=True, disabled=st.session_state['edit_outline_mode']):
                    st.session_state['edit_outline_mode'] = True
                    st.rerun()

            st.markdown("---")

            if st.session_state['edit_outline_mode']:
                # 수정 모드
                st.markdown('<p style="color:var(--gold);font-size:14px;">📝 제목을 직접 수정할 수 있습니다</p>', unsafe_allow_html=True)

                updated_outline = []
                updated_chapters = {}

                for ch_idx, ch in enumerate(st.session_state['outline']):
                    # 챕터 제목 수정
                    new_ch_title = st.text_input(
                        f"PART {ch_idx + 1}",
                        value=ch,
                        key=f"edit_ch_{ch_idx}"
                    )
                    updated_outline.append(new_ch_title)
                    updated_chapters[new_ch_title] = {'subtopics': [], 'subtopic_data': {}}

                    # 소제목 수정
                    subtopics = st.session_state['chapters'].get(ch, {}).get('subtopics', [])
                    new_subtopics = []
                    for st_idx, st_name in enumerate(subtopics):
                        new_st = st.text_input(
                            f"  └ 소제목 {st_idx + 1}",
                            value=st_name,
                            key=f"edit_st_{ch_idx}_{st_idx}",
                            label_visibility="collapsed"
                        )
                        if new_st.strip():
                            new_subtopics.append(new_st)
                            # 기존 데이터 유지
                            old_data = st.session_state['chapters'].get(ch, {}).get('subtopic_data', {}).get(st_name, {'questions': [], 'answers': [], 'content': ''})
                            updated_chapters[new_ch_title]['subtopic_data'][new_st] = old_data

                    updated_chapters[new_ch_title]['subtopics'] = new_subtopics
                    st.markdown("---")

                # 저장 버튼
                if st.button("💾 수정 내용 저장", use_container_width=True, type="primary"):
                    st.session_state['outline'] = updated_outline
                    st.session_state['chapters'] = updated_chapters
                    st.session_state['edit_outline_mode'] = False
                    st.success("목차가 수정되었습니다!")
                    st.rerun()

            else:
                # 보기 모드 - 예쁘게 표시
                for ch_idx, ch in enumerate(st.session_state['outline']):
                    st.markdown(f"""
                    <div style="background:linear-gradient(135deg, rgba(201,162,75,0.15) 0%, rgba(201,162,75,0.05) 100%);
                                padding:16px 20px;border-radius:12px;margin-bottom:8px;border-left:4px solid var(--gold);">
                        <span style="color:var(--gold);font-size:13px;font-weight:600;">PART {ch_idx + 1}</span>
                        <p style="color:var(--text);font-size:17px;font-weight:600;margin:8px 0 0 0;">{ch}</p>
                    </div>
                    """, unsafe_allow_html=True)

                    subtopics = st.session_state['chapters'].get(ch, {}).get('subtopics', [])
                    for st_idx, st_name in enumerate(subtopics):
                        st.markdown(f"""
                        <div style="padding:10px 20px 10px 35px;color:var(--text);font-size:15px;">
                            <span style="color:var(--gold);margin-right:8px;">•</span>{st_name}
                        </div>
                        """, unsafe_allow_html=True)

                    st.markdown("<div style='margin-bottom:20px;'></div>", unsafe_allow_html=True)

        else:
            st.markdown('<div style="text-align:center;padding:60px;background:rgba(255,255,255,0.03);border-radius:16px;border:1px solid rgba(201,162,75,0.15);"><p style="color:rgba(255,255,255,0.5);">목차를 생성해주세요</p></div>', unsafe_allow_html=True)

    st.markdown('<div class="next-section"></div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1, 1, 1])
    with c1:
        if st.button("이전", key="p4_prev", use_container_width=True):
            go_prev()
            st.rerun()
    with c3:
        if st.button("다음 본문", key="p4_next", use_container_width=True):
            go_next()
            st.rerun()


# ==========================================
# PAGE 5: 본문 작성
# ==========================================
elif current == 5:
    st.markdown("""
    <div class="section-title-box">
        <span class="section-step">STEP 06</span>
        <h2>본문 작성</h2>
        <p>AI가 각 챕터의 콘텐츠를 작성합니다</p>
    </div>
    """, unsafe_allow_html=True)

    if not st.session_state.get('outline'):
        st.warning("먼저 목차를 설계하세요")
    else:
        col_sel1, col_sel2 = st.columns([1, 1])
        with col_sel1:
            selected_ch = st.selectbox("챕터", st.session_state['outline'], key="p5_chapter")

        # 선택된 챕터가 있고 chapters에 존재하는지 확인
        if selected_ch and selected_ch in st.session_state.get('chapters', {}):
            ch_data = st.session_state['chapters'][selected_ch]
            subtopics_list = ch_data.get('subtopics', [])

            # 소제목이 있는 경우에만 선택박스 표시
            selected_st = None
            if subtopics_list:
                with col_sel2:
                    selected_st = st.selectbox("소제목", subtopics_list, key="p5_subtopic")

            # 진행률 표시
            completed = sum(1 for s in subtopics_list if ch_data.get('subtopic_data', {}).get(s, {}).get('content'))
            total = len(subtopics_list)
            if total > 0:
                st.progress(completed / total)
                st.caption(f"{completed}/{total} 완료")

            # 소제목이 선택된 경우에만 편집 UI 표시
            if selected_st:
                # subtopic_data 초기화 확인
                if 'subtopic_data' not in ch_data:
                    ch_data['subtopic_data'] = {}
                if selected_st not in ch_data['subtopic_data']:
                    ch_data['subtopic_data'][selected_st] = {'questions': [], 'answers': [], 'content': ''}

                st_data = ch_data['subtopic_data'][selected_st]

                col1, col2 = st.columns([1, 1])

                # 버튼 키를 위한 고유 식별자
                st_key = f"{selected_ch}_{selected_st}".replace(" ", "_")

                with col1:
                    st.markdown("### 인터뷰")
                    if st.button("질문 생성", key=f"gen_q_{st_key}"):
                        if not get_api_key():
                            st.error("API 키를 입력해주세요")
                        else:
                            with st.spinner("생성 중..."):
                                q_text = generate_questions(selected_st, selected_ch, st.session_state['topic'])
                                if q_text:
                                    questions = re.findall(r'Q\d+:\s*(.+)', q_text)
                                    if not questions:
                                        questions = [q.strip() for q in q_text.split('\n') if '?' in q][:3]
                                    if questions:
                                        st_data['questions'] = questions
                                        st_data['answers'] = [''] * len(questions)
                                        st.rerun()
                                    else:
                                        st.error("질문 생성에 실패했습니다")

                    if st_data.get('questions'):
                        for i, q in enumerate(st_data['questions']):
                            st.markdown(f"**Q{i+1}.** {q}")
                            # answers 리스트 크기 확인
                            while len(st_data.get('answers', [])) <= i:
                                st_data['answers'].append('')
                            st_data['answers'][i] = st.text_area(f"A{i+1}", value=st_data['answers'][i], height=80, key=f"ans_{st_key}_{i}", label_visibility="collapsed")

                with col2:
                    st.markdown("### 본문")
                    has_ans = st_data.get('questions') and any(a.strip() for a in st_data.get('answers', []))

                    if has_ans:
                        if st.button("본문 생성", key=f"gen_content_{st_key}", use_container_width=True, type="primary"):
                            if not get_api_key():
                                st.error("API 키를 입력해주세요")
                            else:
                                with st.spinner("본문 작성 중... (1~2분 소요)"):
                                    content = generate_content_premium(selected_st, selected_ch, st_data['questions'], st_data['answers'], st.session_state['topic'], st.session_state['target_persona'])
                                    if content:
                                        st_data['content'] = content
                                        st.success("본문 생성 완료!")
                                        st.rerun()
                                    else:
                                        st.error("본문 생성에 실패했습니다. 다시 시도해주세요.")
                    else:
                        st.info("왼쪽에서 질문에 답변을 입력하면 본문을 생성할 수 있습니다")

                    # 본문 표시
                    current_content = st_data.get('content', '')
                    if current_content:
                        # HTML 형식으로 변환하여 표시
                        formatted_html = format_content_html(current_content)
                        st.markdown(f"""
                        <style>
                        .content-preview-box {{
                            background:#ffffff !important;
                            padding:25px 30px;
                            border-radius:12px;
                            border:1px solid rgba(201,162,75,0.3);
                            margin:15px 0;
                            font-family:'S-CoreDream', sans-serif !important;
                            font-size:17px;
                            max-height:500px;
                            overflow-y:auto;
                        }}
                        .content-preview-box,
                        .content-preview-box p,
                        .content-preview-box span,
                        .content-preview-box div {{
                            color:#000000 !important;
                            -webkit-text-fill-color:#000000 !important;
                        }}
                        .content-preview-box b[style*="color:#e67e22"],
                        .content-preview-box p[style*="color:#e67e22"] {{
                            color:#e67e22 !important;
                            -webkit-text-fill-color:#e67e22 !important;
                        }}
                        </style>
                        <div class="content-preview-box">
                            {formatted_html}
                        </div>
                        """, unsafe_allow_html=True)
                        st.caption(f"📝 {len(current_content.replace(' ', '').replace(chr(10), '')):,}자")

                        # 이미지 추가 기능
                        st.markdown("---")
                        st.markdown("**📷 이미지 추가**")
                        uploaded_img = st.file_uploader("이미지 업로드", type=['png', 'jpg', 'jpeg'], key=f"img_{st_key}", label_visibility="collapsed")
                        if uploaded_img:
                            # 이미지 저장
                            if 'images' not in st_data:
                                st_data['images'] = []
                            img_b64 = base64.b64encode(uploaded_img.read()).decode()
                            st_data['images'].append({'name': uploaded_img.name, 'data': img_b64})
                            st.success(f"이미지 '{uploaded_img.name}' 추가됨!")
                            st.rerun()

                        # 추가된 이미지 표시
                        if st_data.get('images'):
                            st.caption(f"추가된 이미지: {len(st_data['images'])}개")
                            for idx, img in enumerate(st_data['images']):
                                col_img, col_del = st.columns([4, 1])
                                with col_img:
                                    st.image(f"data:image/png;base64,{img['data']}", caption=img['name'], width=200)
                                with col_del:
                                    if st.button("삭제", key=f"del_img_{st_key}_{idx}"):
                                        st_data['images'].pop(idx)
                                        st.rerun()

                        # 수정 기능
                        st.markdown("---")
                        with st.expander("✏️ 본문 직접 수정"):
                            st.caption("「중요단어」 → 주황색 강조 | ★ 문장 → 핵심 강조")
                            edited = st.text_area("본문 편집", value=current_content, height=400, key=f"content_{st_key}", label_visibility="collapsed")
                            if edited != current_content:
                                st_data['content'] = edited
                                st.rerun()
                    else:
                        st.markdown('<div style="text-align:center;padding:80px 20px;background:rgba(255,255,255,0.03);border-radius:12px;border:1px dashed rgba(201,162,75,0.3);"><p style="color:var(--text2);font-size:16px;">본문이 아직 없습니다<br>질문에 답변 후 "본문 생성" 버튼을 누르세요</p></div>', unsafe_allow_html=True)
            else:
                st.info("이 챕터에는 소제목이 없습니다. 목차를 다시 생성해주세요.")

        st.markdown("---")
        st.markdown("### 전체 본문")
        full_content = get_full_content()
        if full_content:
            char_count = len(full_content.replace(' ', '').replace('\n', ''))
            est_pages = estimate_docx_pages(
                st.session_state.get('chapters', {}),
                st.session_state.get('outline', [])
            )
            st.success(f"총 {char_count:,}자 | 약 {est_pages}페이지 (WORD A5 기준)")

    st.markdown('<div class="next-section"></div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1, 1, 1])
    with c1:
        if st.button("이전", key="p5_prev", use_container_width=True):
            go_prev()
            st.rerun()
    with c3:
        if st.button("다음 출력", key="p5_next", use_container_width=True):
            go_next()
            st.rerun()


# ==========================================
# PAGE 6: 표지 디자인
# ==========================================
elif current == 6:
    st.markdown("""
    <div class="section-title-box">
        <span class="section-step">STEP 07</span>
        <h2>표지 디자인</h2>
        <p>전문 디자인 툴로 고품질 표지를 만드세요</p>
    </div>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns([1, 1])

    with col1:
        st.markdown("### 표지 정보 정리")

        # 이전 페이지에서 설정한 제목/부제 자동 연동
        saved_title = st.session_state.get('book_title', '')
        saved_subtitle = st.session_state.get('subtitle', '')

        cover_title = st.text_input("표지 제목", value=saved_title, key="cover_title", placeholder="예: 돈의 속성")
        cover_subtitle = st.text_input("부제목", value=saved_subtitle, key="cover_subtitle", placeholder="예: 당신이 모르는 부의 법칙")
        cover_author = st.text_input("저자명", key="cover_author", placeholder="예: 홍길동")

        st.markdown("---")
        st.markdown("### AI 표지 스타일 추천")

        if st.button("내 주제에 맞는 표지 스타일 추천받기", use_container_width=True, key="ai_cover_suggest"):
            topic = st.session_state.get('topic', '')
            if not topic:
                st.error("시장분석 페이지에서 주제를 먼저 입력해주세요")
            elif not get_api_key():
                st.error("API 키를 입력해주세요")
            else:
                with st.spinner("베스트셀러 표지 분석 중..."):
                    prompt = f"""'{topic}' 주제의 전자책 표지 디자인을 추천해주세요.

이 분야의 실제 베스트셀러 책 표지를 분석해서 추천해주세요.

JSON 형식으로 응답:
{{
    "recommended_style": "추천 스타일명",
    "color_scheme": "추천 색상 조합 (예: 검정 배경 + 금색 텍스트)",
    "design_concept": "디자인 콘셉트 설명 (2문장)",
    "typography_tip": "타이포그래피 팁 (폰트 스타일, 크기 등)",
    "reference_books": ["참고할 베스트셀러 표지 1", "표지 2", "표지 3"],
    "canva_search_keyword": "Canva에서 검색할 키워드 (영문)"
}}"""
                    result = ask_ai(prompt, 0.7)
                    parsed = parse_json(result)
                    if parsed:
                        st.session_state['cover_suggestion'] = parsed
                        st.rerun()

        if st.session_state.get('cover_suggestion'):
            sug = st.session_state['cover_suggestion']
            st.markdown(f"""<div class="data-card">
                <b>추천 스타일: {html.escape(str(sug.get('recommended_style', '')))}</b><br>
                <small>색상: {html.escape(str(sug.get('color_scheme', '')))}</small>
            </div>""", unsafe_allow_html=True)
            st.write(sug.get('design_concept', ''))
            if sug.get('typography_tip'):
                st.info(f"💡 타이포그래피 팁: {sug.get('typography_tip', '')}")
            if sug.get('reference_books'):
                st.markdown("**참고 베스트셀러:**")
                for book in sug.get('reference_books', []):
                    st.caption(f"- {book}")
            if sug.get('canva_search_keyword'):
                st.session_state['canva_keyword'] = sug.get('canva_search_keyword', '')

    with col2:
        st.markdown("### 표지 미리보기")

        _label_to_id = {v: k for k, v in COVER_TEMPLATES.items()}
        _choice = st.radio(
            "표지 스타일",
            list(_label_to_id.keys()),
            key="cover_template_choice",
        )
        _tmpl = _label_to_id[_choice]

        _svg = build_cover_svg(
            _tmpl,
            cover_title or "제목을 입력하세요",
            cover_subtitle,
            cover_author,
        )
        st.markdown(
            '<div style="max-width:320px;margin:12px auto 18px;border-radius:8px;'
            'overflow:hidden;box-shadow:0 24px 70px rgba(0,0,0,0.55);">'
            f'{_svg}</div>',
            unsafe_allow_html=True,
        )
        st.download_button(
            "표지 다운로드 (SVG · 고해상도 벡터)",
            _svg,
            file_name=f"{(cover_title or 'cover')}_cover.svg",
            mime="image/svg+xml",
            use_container_width=True,
            key="cover_svg_dl",
        )
        st.caption(
            "SVG는 무손실 벡터 파일입니다. 브라우저에서 열어 캡처하거나, "
            "Canva·Figma·미리캔버스에 올려 PNG/JPG로 내보낼 수 있습니다."
        )

        st.markdown("---")
        st.markdown("### Canva로 표지 만들기")

        st.markdown("""
        <div class="data-card">
            <p style="font-size:16px;margin-bottom:15px;">
                <b>Canva</b>는 전문 디자이너 수준의 표지를 무료로 만들 수 있는 온라인 툴입니다.
            </p>
            <p style="color:var(--text2);font-size:14px;">
                ✓ 수천 개의 프로 템플릿<br>
                ✓ 드래그 앤 드롭 편집<br>
                ✓ 무료 이미지/아이콘<br>
                ✓ 한글 폰트 지원
            </p>
        </div>
        """, unsafe_allow_html=True)

        st.markdown("---")

        # Canva 검색 키워드 설정
        canva_keyword = st.session_state.get('canva_keyword', 'book cover')

        # Canva 책 표지 템플릿 링크
        canva_url = f"https://www.canva.com/templates/?query={canva_keyword}%20book%20cover"

        st.markdown(f"""
        <a href="{canva_url}" target="_blank" style="
            display:block;
            background:linear-gradient(135deg,#7c3aed,#6366f1);
            color:white;
            padding:18px 24px;
            border-radius:12px;
            text-decoration:none;
            text-align:center;
            font-size:18px;
            font-weight:600;
            margin-bottom:15px;
            transition:transform 0.2s;
        ">
            🎨 Canva에서 표지 만들기
        </a>
        """, unsafe_allow_html=True)

        st.markdown("---")
        st.markdown("### 사용 방법")
        st.markdown("""
        1. **Canva 열기** - 위 버튼 클릭 (무료 가입)
        2. **템플릿 선택** - 마음에 드는 디자인 클릭
        3. **텍스트 수정** - 제목, 부제목, 저자명 입력
        4. **다운로드** - PNG 또는 PDF로 저장
        """)

        st.markdown("---")
        st.markdown("### 추천 검색어")

        search_keywords = [
            "ebook cover", "book cover minimalist",
            "book cover gold", "book cover business",
            "korean book cover", "self help book cover"
        ]

        cols = st.columns(2)
        for i, kw in enumerate(search_keywords):
            with cols[i % 2]:
                if st.button(kw, key=f"canva_kw_{i}", use_container_width=True):
                    st.session_state['canva_keyword'] = kw
                    st.rerun()

        st.markdown("---")

        # 복사할 텍스트
        if cover_title or cover_subtitle or cover_author:
            st.markdown("### 복사할 텍스트")
            copy_text = f"제목: {cover_title}\n부제목: {cover_subtitle}\n저자: {cover_author}"
            st.code(copy_text, language=None)

    st.markdown('<div class="next-section"></div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1, 1, 1])
    with c1:
        if st.button("이전", key="p6_prev", use_container_width=True):
            go_prev()
            st.rerun()
    with c3:
        if st.button("다음", key="p6_next", use_container_width=True):
            go_next()
            st.rerun()


# ==========================================
# PAGE 7: 최종 출력
# ==========================================
elif current == 7:
    st.markdown("""
    <div class="section-title-box">
        <span class="section-step">STEP 08</span>
        <h2>최종 출력</h2>
        <p>완성된 전자책을 다운로드하세요</p>
    </div>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns([1.5, 1])

    with col1:
        st.markdown("### 다운로드")

        final_title = st.text_input("제목", value=st.session_state.get('book_title', ''), key="p6_title")
        final_subtitle = st.text_input("부제", value=st.session_state.get('subtitle', ''), key="p6_subtitle")

        full = f"{final_title}\n{final_subtitle}\n\n{'='*50}\n\n"
        for ch in st.session_state.get('outline', []):
            if ch in st.session_state.get('chapters', {}):
                ch_data = st.session_state['chapters'][ch]
                ch_content = ""
                for s in ch_data.get('subtopics', []):
                    c = ch_data.get('subtopic_data', {}).get(s, {}).get('content', '')
                    if c:
                        ch_content += f"\n\n【{s}】\n\n{clean_content(c)}"
                if ch_content:
                    full += f"\n\n{ch}\n{'-'*40}{ch_content}\n"

        st.markdown("**미리보기**")
        st.text_area("전체 내용", value=full, height=300, disabled=True, key="p7_preview")

        # 저자명 가져오기
        author_name = st.session_state.get('author_name', '') or st.session_state.get('interview_data', {}).get('author_name', '')

        # 다운로드 버튼 3개
        st.markdown("### 📥 다운로드")

        c1, c2, c3 = st.columns(3)
        with c1:
            st.download_button("📄 TXT", full, file_name=f"{final_title or 'ebook'}.txt", use_container_width=True, key="p7_txt")
        with c2:
            # HTML 내보내기 - 특수문자 이스케이프 처리
            escaped_title = html.escape(final_title)
            escaped_content = html.escape(full).replace('\n', '<br>')
            html_content = f"""<!DOCTYPE html>
<html lang="ko">
<head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{escaped_title}</title>
    <style>
        body {{
            max-width: 800px;
            margin: 0 auto;
            padding: 60px 40px;
            font-family: 'Pretendard', -apple-system, BlinkMacSystemFont, sans-serif;
            line-height: 2;
            background: #fafafa;
            color: #333;
        }}
        h1 {{ font-size: 32px; color: #1a1a2e; margin-bottom: 10px; }}
        h2 {{ font-size: 14px; color: #888; font-weight: normal; }}
    </style>
</head>
<body>
{escaped_content}
</body>
</html>"""
            st.download_button("🌐 HTML", html_content, file_name=f"{final_title or 'ebook'}.html", use_container_width=True, key="p7_html")

        with c3:
            # DOCX 다운로드
            if DOCX_AVAILABLE:
                docx_data, docx_error = create_ebook_docx(
                    final_title,
                    final_subtitle,
                    author_name,
                    st.session_state.get('chapters', {}),
                    st.session_state.get('outline', []),
                    st.session_state.get('interview_data', {})
                )
                if docx_data:
                    st.download_button(
                        "📘 WORD",
                        docx_data,
                        file_name=f"{final_title or 'ebook'}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        key="p7_docx"
                    )
                else:
                    st.button("📘 WORD", disabled=True, use_container_width=True, key="p7_docx_disabled")
                    if docx_error:
                        st.caption(f"⚠️ {docx_error[:30]}")
            else:
                st.button("📘 WORD", disabled=True, use_container_width=True, key="p7_docx_na")
                st.caption("pip install python-docx")

        total = len(full.replace(' ', '').replace('\n', ''))
        if total > 0:
            est_pages = estimate_docx_pages(
                st.session_state.get('chapters', {}),
                st.session_state.get('outline', [])
            )
            st.success(f"총 {total:,}자 | 약 {est_pages}페이지 (WORD A5 기준)")
            st.caption("표지·판권·프롤로그·에필로그·저자소개와 소제목별 페이지 나눔까지 반영한 추정치입니다. 워드의 글꼴·자동 줄나눔에 따라 1~2페이지 오차가 있을 수 있습니다.")

    with col2:
        st.markdown("### 현황")
        total_st = sum(len(ch.get('subtopics', [])) for ch in st.session_state.get('chapters', {}).values())
        done = sum(1 for ch in st.session_state.get('chapters', {}).values() for s in ch.get('subtopic_data', {}).values() if s.get('content'))

        if total_st > 0:
            st.progress(done / total_st)
            st.write(f"**완료:** {done}/{total_st}")

    st.markdown('<div class="next-section"></div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1, 2, 1])
    with c1:
        if st.button("이전", key="p7_prev", use_container_width=True):
            go_prev()
            st.rerun()


st.markdown("""
<div style="
    text-align: center;
    padding: 30px 20px;
    margin-top: 50px;
    border-top: 1px solid rgba(201,162,75,0.3);
    color: #ffffff !important;
    font-size: 16px;
    letter-spacing: 2px;
    background: rgba(0,0,0,0.3);
">
    <span style="color: #C9A24B;">CASHMAKER</span> | 제작: <span style="color: #ffffff;">남현우 작가</span>
</div>
""", unsafe_allow_html=True)
